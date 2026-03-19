"""Tests for Phase 6.2 — DOCX Renderer + Blob Store.

Covers:
  1. BlobStore — put/get/stream/delete/exists, path traversal prevention
  2. DOCX Renderer — placeholder filling, metadata normalization,
     deterministic hashing, full lifecycle, failure handling
  3. Router — execute render, download artifact, error cases
  4. Runner — get_artifact, get_template_version_with_template
"""

import hashlib
import io
import json
import os
import shutil
import tempfile
import pytest
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch
from uuid import UUID, uuid4

from docx import Document

from shadow_service import docx_factory_runner as runner
from shadow_service import sql_docx_factory as sql


# =============================================================================
# Helpers
# =============================================================================

NOW = datetime(2026, 2, 6, 12, 0, 0, tzinfo=timezone.utc)


class FakeRecord(dict):
    """Mimics asyncpg.Record: dict(row), row["key"], row.get(k)."""
    def __getattr__(self, name):
        try:
            return self[name]
        except KeyError:
            raise AttributeError(name)


def _make_test_docx(placeholders: dict[str, str] | None = None) -> bytes:
    """Create a minimal DOCX with optional {{ placeholder }} text."""
    doc = Document()
    if placeholders:
        for key, default_text in placeholders.items():
            doc.add_paragraph(f"{{{{ {key} }}}}")
    else:
        doc.add_paragraph("Hello {{ drug_name }}, protocol {{ protocol_id }}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_render_row(**overrides) -> FakeRecord:
    base = {
        "id": uuid4(),
        "program_id": uuid4(),
        "template_version_id": uuid4(),
        "inputs_json": {"drug_name": "Axolotinib", "protocol_id": "PROTO-001"},
        "status": "queued",
        "started_at": None,
        "completed_at": None,
        "error": None,
        "created_at": NOW,
    }
    base.update(overrides)
    return FakeRecord(base)


def _make_artifact_row(**overrides) -> FakeRecord:
    base = {
        "id": uuid4(),
        "render_id": uuid4(),
        "file_type": "docx",
        "storage_key": "renders/abc/output.docx",
        "sha256": "b" * 64,
        "size_bytes": 54321,
        "created_at": NOW,
    }
    base.update(overrides)
    return FakeRecord(base)


def _make_tv_with_template_row(**overrides) -> FakeRecord:
    base = {
        "version_id": uuid4(),
        "template_id": uuid4(),
        "version": 1,
        "storage_key": "templates/test-v1.docx",
        "sha256": "a" * 64,
        "program_id": uuid4(),
        "template_name": "IND Module 2.5",
        "doc_type": "ind_clinical_overview",
    }
    base.update(overrides)
    return FakeRecord(base)


# =============================================================================
# 1. BlobStore — local filesystem implementation
# =============================================================================

class TestLocalBlobStore:
    """Test LocalBlobStore with a real temp directory."""

    @pytest.fixture(autouse=True)
    def _setup_store(self, tmp_path):
        from shadow_service.blob_store import LocalBlobStore, reset_blob_store
        self.root = tmp_path / "test_blobs"
        self.store = LocalBlobStore(root=self.root)
        yield
        reset_blob_store()

    @pytest.mark.asyncio
    async def test_put_and_get_bytes(self):
        data = b"hello regulatory world"
        result = await self.store.put_bytes("test/file.bin", data)

        assert result["key"] == "test/file.bin"
        assert result["size_bytes"] == len(data)
        assert result["sha256"] == hashlib.sha256(data).hexdigest()

        retrieved = await self.store.get_bytes("test/file.bin")
        assert retrieved == data

    @pytest.mark.asyncio
    async def test_get_bytes_not_found(self):
        with pytest.raises(FileNotFoundError):
            await self.store.get_bytes("nonexistent/file.bin")

    @pytest.mark.asyncio
    async def test_get_stream(self):
        data = b"stream me"
        await self.store.put_bytes("streamable.bin", data)
        stream = await self.store.get_stream("streamable.bin")
        assert stream.read() == data

    @pytest.mark.asyncio
    async def test_exists(self):
        assert not await self.store.exists("nope.bin")
        await self.store.put_bytes("yep.bin", b"x")
        assert await self.store.exists("yep.bin")

    @pytest.mark.asyncio
    async def test_delete(self):
        await self.store.put_bytes("deleteme.bin", b"gone")
        assert await self.store.delete("deleteme.bin")
        assert not await self.store.exists("deleteme.bin")

    @pytest.mark.asyncio
    async def test_delete_nonexistent(self):
        assert not await self.store.delete("ghost.bin")

    @pytest.mark.asyncio
    async def test_path_traversal_blocked(self):
        with pytest.raises(ValueError, match="path traversal"):
            await self.store.put_bytes("../../etc/passwd", b"nope")

    @pytest.mark.asyncio
    async def test_put_creates_subdirectories(self):
        await self.store.put_bytes("deep/nested/path/file.bin", b"deep")
        assert await self.store.exists("deep/nested/path/file.bin")

    @pytest.mark.asyncio
    async def test_sha256_correctness(self):
        data = b"deterministic content"
        result = await self.store.put_bytes("hash-test.bin", data)
        expected = hashlib.sha256(data).hexdigest()
        assert result["sha256"] == expected


class TestBlobStoreFactory:
    """Test get_blob_store factory and singleton."""

    def setup_method(self):
        from shadow_service.blob_store import reset_blob_store
        reset_blob_store()

    def teardown_method(self):
        from shadow_service.blob_store import reset_blob_store
        reset_blob_store()

    def test_default_creates_local_store(self):
        from shadow_service.blob_store import get_blob_store, LocalBlobStore
        store = get_blob_store()
        assert isinstance(store, LocalBlobStore)

    def test_singleton_returns_same_instance(self):
        from shadow_service.blob_store import get_blob_store
        s1 = get_blob_store()
        s2 = get_blob_store()
        assert s1 is s2

    def test_unknown_type_raises(self):
        from shadow_service.blob_store import get_blob_store, reset_blob_store
        reset_blob_store()
        with patch.dict(os.environ, {"BLOB_STORE_TYPE": "s3"}):
            with pytest.raises(ValueError, match="Unknown blob store type"):
                get_blob_store()

    def test_reset_clears_singleton(self):
        from shadow_service.blob_store import get_blob_store, reset_blob_store
        s1 = get_blob_store()
        reset_blob_store()
        s2 = get_blob_store()
        assert s1 is not s2


# =============================================================================
# 2. DOCX Renderer — unit tests
# =============================================================================

class TestPlaceholderFilling:
    """Test _fill_paragraph and _fill_document logic."""

    def test_fill_simple_placeholder(self):
        from shadow_service.docx_renderer import _fill_document
        doc = Document()
        doc.add_paragraph("Patient: {{ drug_name }}")
        _fill_document(doc, {"drug_name": "Axolotinib"})
        text = doc.paragraphs[0].text
        assert "Axolotinib" in text
        assert "{{" not in text

    def test_fill_multiple_placeholders(self):
        from shadow_service.docx_renderer import _fill_document
        doc = Document()
        doc.add_paragraph("Drug: {{ drug_name }}, Protocol: {{ protocol_id }}")
        _fill_document(doc, {"drug_name": "TestDrug", "protocol_id": "P-001"})
        text = doc.paragraphs[0].text
        assert "TestDrug" in text
        assert "P-001" in text

    def test_fill_preserves_unknown_placeholder(self):
        from shadow_service.docx_renderer import _fill_document
        doc = Document()
        doc.add_paragraph("{{ unknown_key }}")
        _fill_document(doc, {"other_key": "value"})
        assert "{{ unknown_key }}" in doc.paragraphs[0].text

    def test_fill_table_cells(self):
        from shadow_service.docx_renderer import _fill_document
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].text = "{{ field }}"
        _fill_document(doc, {"field": "Filled!"})
        assert "Filled!" in table.rows[0].cells[0].text

    def test_fill_empty_values(self):
        from shadow_service.docx_renderer import _fill_document
        doc = Document()
        doc.add_paragraph("{{ drug_name }}")
        _fill_document(doc, {"drug_name": ""})
        text = doc.paragraphs[0].text
        assert "{{" not in text


class TestMetadataNormalization:
    """Test _normalize_core_properties for deterministic hashing."""

    def test_core_properties_are_pinned(self):
        from shadow_service.docx_renderer import _normalize_core_properties, _EPOCH
        doc = Document()
        doc.core_properties.author = "Random Author"
        doc.core_properties.last_modified_by = "Some User"
        _normalize_core_properties(doc)
        # python-docx stores datetimes as UTC-aware, so compare without tz
        assert doc.core_properties.created.replace(tzinfo=None) == _EPOCH
        assert doc.core_properties.modified.replace(tzinfo=None) == _EPOCH
        assert doc.core_properties.author == "Concept2Cure.RI"
        assert doc.core_properties.last_modified_by == "Concept2Cure.RI"
        assert doc.core_properties.revision == 1


class TestDeterministicHashing:
    """Verify that same inputs + same template → same SHA-256."""

    def test_same_inputs_produce_same_hash(self):
        from shadow_service.docx_renderer import (
            _fill_document, _normalize_core_properties,
        )

        def _render_to_bytes(template_bytes: bytes, values: dict) -> bytes:
            doc = Document(io.BytesIO(template_bytes))
            _fill_document(doc, values)
            _normalize_core_properties(doc)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        template = _make_test_docx()
        values = {"drug_name": "Axolotinib", "protocol_id": "PROTO-001"}

        output1 = _render_to_bytes(template, values)
        output2 = _render_to_bytes(template, values)

        hash1 = hashlib.sha256(output1).hexdigest()
        hash2 = hashlib.sha256(output2).hexdigest()
        assert hash1 == hash2, "Deterministic hashing failed: same inputs → different SHA-256"

    def test_different_inputs_produce_different_hash(self):
        from shadow_service.docx_renderer import (
            _fill_document, _normalize_core_properties,
        )

        def _render_to_bytes(template_bytes: bytes, values: dict) -> bytes:
            doc = Document(io.BytesIO(template_bytes))
            _fill_document(doc, values)
            _normalize_core_properties(doc)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        template = _make_test_docx()
        values_a = {"drug_name": "DrugA", "protocol_id": "P-A"}
        values_b = {"drug_name": "DrugB", "protocol_id": "P-B"}

        hash_a = hashlib.sha256(_render_to_bytes(template, values_a)).hexdigest()
        hash_b = hashlib.sha256(_render_to_bytes(template, values_b)).hexdigest()
        assert hash_a != hash_b


# =============================================================================
# 3. Renderer — full lifecycle (mocked DB + blob store)
# =============================================================================

class TestRenderDocumentLifecycle:
    """Test render_document orchestrator with mocked dependencies."""

    @pytest.mark.asyncio
    async def test_successful_render(self):
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        tv_id = uuid4()
        program_id = uuid4()
        template_bytes = _make_test_docx()
        sha256_template = hashlib.sha256(template_bytes).hexdigest()

        render_row = _make_render_row(
            id=render_id,
            program_id=program_id,
            template_version_id=tv_id,
            status="queued",
        )
        tv_row = _make_tv_with_template_row(
            version_id=tv_id,
            storage_key="templates/test.docx",
            sha256=sha256_template,
            program_id=program_id,
        )
        artifact_row = _make_artifact_row(render_id=render_id)

        # Mock DB connection for bypass-RLS render lookup
        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=render_row)

        # Mock blob store
        mock_store = AsyncMock()
        mock_store.get_bytes = AsyncMock(return_value=template_bytes)
        mock_store.put_bytes = AsyncMock(return_value={
            "sha256": "x" * 64,
            "size_bytes": 12345,
            "key": f"renders/{render_id}/output.docx",
        })

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.runner") as mock_runner, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            mock_runner.transition_render_running = AsyncMock(
                return_value=FakeRecord({**dict(render_row), "status": "running", "started_at": NOW})
            )
            mock_runner.get_template_version_with_template = AsyncMock(return_value=tv_row)
            mock_runner.create_artifact = AsyncMock(return_value=artifact_row)
            mock_runner.transition_render_completed = AsyncMock(
                return_value=FakeRecord({**dict(render_row), "status": "completed", "completed_at": NOW})
            )

            result = await render_document(render_id)

            assert result["id"] == artifact_row["id"]
            mock_runner.transition_render_running.assert_called_once_with(render_id)
            mock_runner.transition_render_completed.assert_called_once_with(render_id)
            mock_runner.create_artifact.assert_called_once()
            mock_store.get_bytes.assert_called_once_with("templates/test.docx")
            mock_store.put_bytes.assert_called_once()

    @pytest.mark.asyncio
    async def test_render_not_queued_raises(self):
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        render_row = _make_render_row(id=render_id, status="running")

        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=render_row)

        mock_store = AsyncMock()

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.runner") as mock_runner, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()
            mock_runner.transition_render_failed = AsyncMock(
                return_value=_make_render_row(status="failed")
            )

            with pytest.raises(ValueError, match="expected 'queued'"):
                await render_document(render_id)

    @pytest.mark.asyncio
    async def test_render_not_found_raises(self):
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=None)

        mock_store = AsyncMock()

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            with pytest.raises(ValueError, match="not found"):
                await render_document(render_id)

    @pytest.mark.asyncio
    async def test_render_failure_transitions_to_failed(self):
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        tv_id = uuid4()
        render_row = _make_render_row(id=render_id, template_version_id=tv_id, status="queued")

        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=render_row)

        mock_store = AsyncMock()
        mock_store.get_bytes = AsyncMock(side_effect=FileNotFoundError("Template missing"))

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.runner") as mock_runner, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()
            mock_runner.transition_render_running = AsyncMock(
                return_value=FakeRecord({**dict(render_row), "status": "running"})
            )
            mock_runner.get_template_version_with_template = AsyncMock(
                return_value=_make_tv_with_template_row(version_id=tv_id)
            )
            mock_runner.transition_render_failed = AsyncMock(
                return_value=_make_render_row(status="failed")
            )

            with pytest.raises(FileNotFoundError):
                await render_document(render_id)

            mock_runner.transition_render_failed.assert_called_once()
            call_args = mock_runner.transition_render_failed.call_args
            assert call_args[0][0] == render_id

    @pytest.mark.asyncio
    async def test_render_cross_program_rejected(self):
        """render_document rejects when program_id doesn't match the render's."""
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        actual_program = uuid4()
        claimed_program = uuid4()

        render_row = _make_render_row(
            id=render_id, program_id=actual_program, status="queued",
        )

        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=render_row)

        mock_store = AsyncMock()

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.runner") as mock_runner, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()
            mock_runner.transition_render_failed = AsyncMock(
                return_value=_make_render_row(status="failed")
            )

            with pytest.raises(ValueError, match="not found"):
                await render_document(render_id, program_id=claimed_program)

    @pytest.mark.asyncio
    async def test_render_same_program_allowed(self):
        """render_document proceeds when program_id matches the render's."""
        from shadow_service.docx_renderer import render_document

        render_id = uuid4()
        program_id = uuid4()
        tv_id = uuid4()
        template_bytes = _make_test_docx()

        render_row = _make_render_row(
            id=render_id, program_id=program_id,
            template_version_id=tv_id, status="queued",
        )
        tv_row = _make_tv_with_template_row(
            version_id=tv_id, program_id=program_id,
        )
        artifact_row = _make_artifact_row(render_id=render_id)

        mock_conn = AsyncMock()
        mock_conn.execute = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=render_row)

        mock_store = AsyncMock()
        mock_store.get_bytes = AsyncMock(return_value=template_bytes)
        mock_store.put_bytes = AsyncMock(return_value={
            "sha256": "x" * 64, "size_bytes": 12345,
            "key": f"renders/{render_id}/output.docx",
        })

        with patch("shadow_service.docx_renderer.db") as mock_db, \
             patch("shadow_service.docx_renderer.runner") as mock_runner, \
             patch("shadow_service.docx_renderer.get_blob_store", return_value=mock_store):

            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()
            mock_runner.transition_render_running = AsyncMock(
                return_value=FakeRecord({**dict(render_row), "status": "running"})
            )
            mock_runner.get_template_version_with_template = AsyncMock(return_value=tv_row)
            mock_runner.create_artifact = AsyncMock(return_value=artifact_row)
            mock_runner.transition_render_completed = AsyncMock(
                return_value=FakeRecord({**dict(render_row), "status": "completed"})
            )

            result = await render_document(render_id, program_id=program_id)
            assert result["id"] == artifact_row["id"]


# =============================================================================
# 4. Runner — new functions
# =============================================================================

class TestRunnerGetArtifact:
    """Test get_artifact function."""

    @pytest.mark.asyncio
    async def test_get_artifact_found(self):
        artifact_row = _make_artifact_row()
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=artifact_row)

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_artifact(artifact_row["id"])
            assert result["id"] == artifact_row["id"]
            assert result["file_type"] == "docx"

    @pytest.mark.asyncio
    async def test_get_artifact_not_found(self):
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=None)

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_artifact(uuid4())
            assert result is None


class TestRunnerGetArtifactForProgram:
    """Test get_artifact_for_program function (cross-program safe)."""

    @pytest.mark.asyncio
    async def test_returns_artifact_when_program_matches(self):
        artifact_row = _make_artifact_row()
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=artifact_row)

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_artifact_for_program(artifact_row["id"], uuid4())
            assert result["id"] == artifact_row["id"]

    @pytest.mark.asyncio
    async def test_returns_none_when_program_mismatch(self):
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=None)  # JOIN finds nothing

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_artifact_for_program(uuid4(), uuid4())
            assert result is None

    @pytest.mark.asyncio
    async def test_uses_correct_sql_query(self):
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=None)

        aid = uuid4()
        pid = uuid4()

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            await runner.get_artifact_for_program(aid, pid)
            mock_conn.fetchrow.assert_called_once_with(
                sql.SELECT_ARTIFACT_BY_ID_WITH_PROGRAM, aid, pid,
            )


class TestRunnerGetTemplateVersionWithTemplate:
    """Test get_template_version_with_template function."""

    @pytest.mark.asyncio
    async def test_get_tv_with_template_found(self):
        tv_row = _make_tv_with_template_row()
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=tv_row)

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_template_version_with_template(tv_row["version_id"])
            assert result["version_id"] == tv_row["version_id"]
            assert result["template_name"] == "IND Module 2.5"
            assert result["doc_type"] == "ind_clinical_overview"

    @pytest.mark.asyncio
    async def test_get_tv_with_template_not_found(self):
        mock_conn = AsyncMock()
        mock_conn.fetchrow = AsyncMock(return_value=None)

        with patch("shadow_service.docx_factory_runner.db") as mock_db:
            mock_db.acquire_connection = AsyncMock(return_value=mock_conn)
            mock_db.release_connection = AsyncMock()

            result = await runner.get_template_version_with_template(uuid4())
            assert result is None


# =============================================================================
# 5. Router — execute render + download artifact
# =============================================================================

class TestRouterExecuteRender:
    """Test POST /docx/renders/{render_id}/execute endpoint."""

    @pytest.fixture
    def _client(self):
        """Create a test client with the router mounted."""
        os.environ["REVIEW_ADMIN_TOKEN"] = "test-token-62"
        from fastapi import FastAPI
        from fastapi.testclient import TestClient
        from shadow_service.router_docx_factory import router

        app = FastAPI()
        app.include_router(router)
        self.client = TestClient(app)
        self.headers = {"X-Admin-Token": "test-token-62"}
        self.program_id = uuid4()
        yield
        os.environ.pop("REVIEW_ADMIN_TOKEN", None)

    def test_execute_render_success(self, _client):
        artifact = _make_artifact_row()
        with patch("shadow_service.router_docx_factory.docx_renderer") as mock_renderer:
            mock_renderer.render_document = AsyncMock(return_value=dict(artifact))
            resp = self.client.post(
                f"/docx/renders/{uuid4()}/execute",
                params={"program_id": str(self.program_id)},
                headers=self.headers,
            )
            assert resp.status_code == 200
            data = resp.json()
            assert data["file_type"] == "docx"
            # Verify program_id was forwarded
            call_kwargs = mock_renderer.render_document.call_args
            assert call_kwargs.kwargs["program_id"] == self.program_id

    def test_execute_render_bad_request(self, _client):
        with patch("shadow_service.router_docx_factory.docx_renderer") as mock_renderer:
            mock_renderer.render_document = AsyncMock(
                side_effect=ValueError("Render xyz is 'running', expected 'queued'")
            )
            resp = self.client.post(
                f"/docx/renders/{uuid4()}/execute",
                params={"program_id": str(self.program_id)},
                headers=self.headers,
            )
            assert resp.status_code == 400
            assert "expected 'queued'" in resp.json()["detail"]

    def test_execute_render_no_auth(self, _client):
        resp = self.client.post(
            f"/docx/renders/{uuid4()}/execute",
            params={"program_id": str(self.program_id)},
        )
        assert resp.status_code == 403

    def test_execute_render_requires_program_id(self, _client):
        """422 when program_id query param is missing."""
        resp = self.client.post(
            f"/docx/renders/{uuid4()}/execute",
            headers=self.headers,
        )
        assert resp.status_code == 422

    def test_execute_render_cross_program_rejected(self, _client):
        """Render belonging to program B is rejected when program A is claimed."""
        with patch("shadow_service.router_docx_factory.docx_renderer") as mock_renderer:
            mock_renderer.render_document = AsyncMock(
                side_effect=ValueError("Render abc not found")
            )
            resp = self.client.post(
                f"/docx/renders/{uuid4()}/execute",
                params={"program_id": str(uuid4())},
                headers=self.headers,
            )
            assert resp.status_code == 400
            assert "not found" in resp.json()["detail"]


class TestRouterDownloadArtifact:
    """Test GET /docx/artifacts/{artifact_id}/download endpoint."""

    @pytest.fixture
    def _client(self):
        os.environ["REVIEW_ADMIN_TOKEN"] = "test-token-62"
        from fastapi import FastAPI
        from fastapi.testclient import TestClient
        from shadow_service.router_docx_factory import router

        app = FastAPI()
        app.include_router(router)
        self.client = TestClient(app)
        self.headers = {"X-Admin-Token": "test-token-62"}
        self.program_id = uuid4()
        yield
        os.environ.pop("REVIEW_ADMIN_TOKEN", None)

    def test_download_artifact_success(self, _client):
        artifact = _make_artifact_row()
        docx_bytes = _make_test_docx()

        with patch("shadow_service.router_docx_factory.runner") as mock_runner, \
             patch("shadow_service.router_docx_factory.get_blob_store") as mock_get_store:

            mock_runner.get_artifact_for_program = AsyncMock(return_value=dict(artifact))
            mock_store = AsyncMock()
            mock_store.get_stream = AsyncMock(return_value=io.BytesIO(docx_bytes))
            mock_get_store.return_value = mock_store

            resp = self.client.get(
                f"/docx/artifacts/{artifact['id']}/download",
                params={"program_id": str(self.program_id)},
                headers=self.headers,
            )
            assert resp.status_code == 200
            assert "attachment" in resp.headers["content-disposition"]
            assert resp.headers["x-artifact-sha256"] == artifact["sha256"]

    def test_download_artifact_not_found(self, _client):
        with patch("shadow_service.router_docx_factory.runner") as mock_runner:
            mock_runner.get_artifact_for_program = AsyncMock(return_value=None)
            resp = self.client.get(
                f"/docx/artifacts/{uuid4()}/download",
                params={"program_id": str(self.program_id)},
                headers=self.headers,
            )
            assert resp.status_code == 404

    def test_download_artifact_file_missing(self, _client):
        artifact = _make_artifact_row()
        with patch("shadow_service.router_docx_factory.runner") as mock_runner, \
             patch("shadow_service.router_docx_factory.get_blob_store") as mock_get_store:

            mock_runner.get_artifact_for_program = AsyncMock(return_value=dict(artifact))
            mock_store = AsyncMock()
            mock_store.get_stream = AsyncMock(side_effect=FileNotFoundError("gone"))
            mock_get_store.return_value = mock_store

            resp = self.client.get(
                f"/docx/artifacts/{artifact['id']}/download",
                params={"program_id": str(self.program_id)},
                headers=self.headers,
            )
            assert resp.status_code == 404
            assert "not found in storage" in resp.json()["detail"]

    def test_download_artifact_no_auth(self, _client):
        resp = self.client.get(
            f"/docx/artifacts/{uuid4()}/download",
            params={"program_id": str(self.program_id)},
        )
        assert resp.status_code == 403

    def test_download_artifact_requires_program_id(self, _client):
        """422 when program_id query param is missing."""
        resp = self.client.get(
            f"/docx/artifacts/{uuid4()}/download",
            headers=self.headers,
        )
        assert resp.status_code == 422

    def test_download_artifact_cross_program_rejected(self, _client):
        """Artifact belonging to a different program returns 404."""
        with patch("shadow_service.router_docx_factory.runner") as mock_runner:
            # Simulate cross-program: get_artifact_for_program returns None
            mock_runner.get_artifact_for_program = AsyncMock(return_value=None)
            resp = self.client.get(
                f"/docx/artifacts/{uuid4()}/download",
                params={"program_id": str(uuid4())},
                headers=self.headers,
            )
            assert resp.status_code == 404


# =============================================================================
# 6. SQL queries — structural tests
# =============================================================================

class TestNewSQLQueries:
    """Verify the new SQL queries have correct structure."""

    def test_select_artifact_by_id_targets_documents_schema(self):
        assert "documents.artifacts" in sql.SELECT_ARTIFACT_BY_ID
        assert "$1" in sql.SELECT_ARTIFACT_BY_ID

    def test_select_artifact_by_id_returns_expected_columns(self):
        for col in ["id", "render_id", "file_type", "storage_key", "sha256", "size_bytes"]:
            assert col in sql.SELECT_ARTIFACT_BY_ID

    def test_select_artifact_by_id_with_program_has_join(self):
        query = sql.SELECT_ARTIFACT_BY_ID_WITH_PROGRAM
        assert "JOIN" in query
        assert "documents.artifacts" in query
        assert "documents.renders" in query
        assert "program_id" in query
        assert "$1" in query
        assert "$2" in query

    def test_select_tv_with_template_has_join(self):
        query = sql.SELECT_TEMPLATE_VERSION_WITH_TEMPLATE
        assert "JOIN" in query
        assert "documents.template_versions" in query
        assert "documents.templates" in query

    def test_select_tv_with_template_returns_expected_columns(self):
        query = sql.SELECT_TEMPLATE_VERSION_WITH_TEMPLATE
        for col in ["version_id", "template_id", "version", "storage_key",
                     "sha256", "program_id", "template_name", "doc_type"]:
            assert col in query
