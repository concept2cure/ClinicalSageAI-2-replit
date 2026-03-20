"""Phase 6.2 — DOCX Renderer: Template filling + deterministic hashing.

Orchestrates the full render lifecycle:
  1. Look up render record → get template_version_id, inputs_json
  2. Transition queued → running
  3. Load template bytes from blob store via storage_key
  4. Fill template placeholders with inputs_json using python-docx
  5. Normalize DOCX metadata (strip nondeterministic fields)
  6. Compute SHA-256 of produced bytes
  7. Store artifact in blob store
  8. Register artifact metadata in DB
  9. Transition running → completed
  10. On error: transition running → failed

Deterministic hashing: same inputs + same template bytes → same SHA-256.
"""

import hashlib
import io
import logging
import re
from datetime import datetime
from typing import Any
from uuid import UUID

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from . import docx_factory_runner as runner
from . import db
from . import sql_docx_factory as sql
from .blob_store import get_blob_store

logger = logging.getLogger(__name__)

# Sentinel date used to normalize core-properties for deterministic output.
_EPOCH = datetime(2000, 1, 1, 0, 0, 0)

# Regex for Jinja2-style placeholders: {{ variable_name }}
_PLACEHOLDER_RE = re.compile(r"\{\{\s*(\w+)\s*\}\}")


# =============================================================================
# Placeholder filling
# =============================================================================

def _fill_paragraph(paragraph, values: dict[str, str]) -> None:
    """Replace {{ key }} placeholders in a paragraph's runs.

    Handles the common case where python-docx splits a single
    placeholder across multiple runs.  Strategy: join all run texts,
    do a single regex replace, then rewrite the runs.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not _PLACEHOLDER_RE.search(full_text):
        return

    new_text = _PLACEHOLDER_RE.sub(
        lambda m: str(values.get(m.group(1), m.group(0))),
        full_text,
    )

    if new_text == full_text:
        return  # no change

    # Preserve first run's formatting, clear the rest
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""


def _fill_document(doc: Document, values: dict[str, str]) -> None:
    """Fill all placeholders in the document body + headers/footers."""
    # Body paragraphs
    for para in doc.paragraphs:
        _fill_paragraph(para, values)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _fill_paragraph(para, values)

    # Headers and footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is None:
                continue
            for para in header_footer.paragraphs:
                _fill_paragraph(para, values)


# =============================================================================
# Metadata normalization (deterministic hashing)
# =============================================================================

def _normalize_core_properties(doc: Document) -> None:
    """Strip/normalize nondeterministic metadata from DOCX core properties.

    python-docx sets `created`, `modified`, `last_modified_by`, `revision`
    on save.  We pin these to known sentinel values so the same input data
    always produces the same output bytes.
    """
    cp = doc.core_properties
    cp.created = _EPOCH
    cp.modified = _EPOCH
    cp.last_modified_by = "Concept2Cure.RI"
    cp.revision = 1
    cp.author = "Concept2Cure.RI"
    # Clear other volatile fields
    cp.last_printed = _EPOCH
    cp.category = cp.category or ""
    cp.comments = cp.comments or ""


# =============================================================================
# Orchestrator
# =============================================================================

async def render_document(render_id: UUID, *, program_id: UUID | None = None) -> dict[str, Any]:
    """Execute a full render: template fill → artifact storage → DB update.

    Args:
        render_id: The render to execute.
        program_id: If provided, enforces that the render belongs to this
            program.  Returns a ValueError on mismatch (used by the router
            to prevent cross-program access).

    Returns the artifact record on success.
    Raises on failure (caller should handle / the router catches).
    """
    store = get_blob_store()
    render = None

    try:
        # 1. Look up the render record
        #    Bypass RLS — render_document is a system-level operation

        conn = await db.acquire_connection()
        try:
            # Bypass RLS — render_document is a system-level operation
            await conn.execute(sql.SET_BYPASS_RLS)
            row = await conn.fetchrow(sql.SELECT_RENDER_BY_ID, render_id)
            if not row:
                raise ValueError(f"Render {render_id} not found")
            render = dict(row)

            # Cross-program ownership check
            if program_id is not None and render["program_id"] != program_id:
                raise ValueError(f"Render {render_id} not found")
        finally:
            await db.release_connection(conn)

        if render["status"] != "queued":
            raise ValueError(
                f"Render {render_id} is '{render['status']}', expected 'queued'"
            )

        # 2. Transition queued → running
        await runner.transition_render_running(render_id)
        logger.info("Render %s: running", render_id)

        # 3. Resolve template version + parent template
        tv = await runner.get_template_version_with_template(
            render["template_version_id"],
        )
        if not tv:
            raise ValueError(
                f"Template version {render['template_version_id']} not found"
            )

        # 4. Load template bytes from blob store
        template_bytes = await store.get_bytes(tv["storage_key"])

        # 5. Fill template with inputs_json using python-docx
        doc = Document(io.BytesIO(template_bytes))
        inputs = render.get("inputs_json", {})
        if isinstance(inputs, str):
            import json
            inputs = json.loads(inputs)

        # Flatten all values to strings for placeholder replacement
        str_values = {k: str(v) for k, v in inputs.items()}
        _fill_document(doc, str_values)

        # 6. Normalize metadata for deterministic hashing
        _normalize_core_properties(doc)

        # 7. Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        rendered_bytes = buf.getvalue()

        # 8. Compute SHA-256
        sha256 = hashlib.sha256(rendered_bytes).hexdigest()

        # 9. Store artifact in blob store
        artifact_key = f"renders/{render_id}/{sha256}.docx"
        blob_meta = await store.put_bytes(artifact_key, rendered_bytes)

        # 10. Register artifact metadata in DB
        artifact = await runner.create_artifact(
            render_id=render_id,
            file_type="docx",
            storage_key=artifact_key,
            sha256=sha256,
            size_bytes=blob_meta["size_bytes"],
        )

        # 11. Transition running → completed
        await runner.transition_render_completed(render_id)
        logger.info(
            "Render %s: completed (artifact=%s, sha256=%s)",
            render_id, artifact["id"], sha256[:12],
        )

        return artifact

    except Exception as exc:
        logger.exception("Render %s: failed — %s", render_id, exc)
        try:
            await runner.transition_render_failed(render_id, str(exc)[:1000])
        except Exception:
            logger.exception("Failed to record render failure for %s", render_id)
        raise
