"""Evidence Cell DOCX Renderer — Phase 6.6.C.

Renders SE matrix comparison rows into python-docx tables with:
- Green highlight: evidence_ids exist (linked, verified)
- Yellow highlight: missing evidence_ids (needs citation)
- Red markers: SIGNIFICANT / NOT_EQUIVALENT diffs
- Invisible bookmarks: EV_<id> for Defense Packet extraction

Provides both:
  1. render_se_matrix_table() — structured render plan (JSON-friendly)
  2. render_to_docx_table()  — actual python-docx Table object

Integrates with DOCX Factory render lifecycle.
"""
from __future__ import annotations

import logging
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table, _Cell

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Color constants (RGB hex for python-docx)
# ─────────────────────────────────────────────────────────────────────────────
COLOR_GREEN = "C6EFCE"   # Evidence present
COLOR_YELLOW = "FFEB9C"  # Missing evidence
COLOR_RED = "FFC7CE"     # Not equivalent / critical
COLOR_WHITE = "FFFFFF"   # Equivalent / no issue
COLOR_LIGHT_GRAY = "F2F2F2"  # Alternating row
COLOR_HEADER_BG = "002060"

# Diff flag display labels
DIFF_LABELS = {
    "EQUIVALENT": "≡ Equivalent",
    "DISCUSSION_REQUIRED": "⚠ Discussion Required",
    "NOT_EQUIVALENT": "✗ Not Equivalent",
    "TOXIC": "☠ Toxic",
    "PENDING": "⏳ Pending",
}

# Severity badge labels
SEVERITY_BADGES = {
    "none":     "—",
    "low":      "Low",
    "medium":   "⚠ Medium",
    "high":     "▲ High",
    "critical": "◆ Critical",
}


def _apply_cell_shading(cell: _Cell, hex_color: str) -> None:
    """Apply background fill to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _apply_cell_borders(cell: _Cell, color: str = "BFBFBF") -> None:
    """Apply thin borders to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        el.set(qn("w:space"), "0")
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _insert_bookmark(paragraph, name: str) -> None:
    """Insert an invisible bookmark at the end of a paragraph."""
    bid = str(abs(hash(name)) % 10000)
    bs = OxmlElement("w:bookmarkStart")
    bs.set(qn("w:id"), bid)
    bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd")
    be.set(qn("w:id"), bid)
    paragraph._element.append(bs)
    paragraph._element.append(be)


def get_cell_highlight(
    equivalence_status: str,
    has_evidence: bool,
    diff_severity: str = "none",
) -> str:
    """Determine cell background color based on evidence + diff status.

    Returns hex RGB color string.
    """
    if equivalence_status == "NOT_EQUIVALENT":
        return COLOR_RED
    if equivalence_status == "TOXIC":
        return COLOR_RED
    if diff_severity == "critical":
        return COLOR_RED

    if equivalence_status in ("DISCUSSION_REQUIRED",):
        if has_evidence:
            return COLOR_GREEN
        return COLOR_YELLOW

    if equivalence_status == "EQUIVALENT":
        return COLOR_GREEN if has_evidence else COLOR_WHITE

    # PENDING or unknown
    return COLOR_YELLOW if not has_evidence else COLOR_WHITE


def _status_column_color(eq_status: str) -> str:
    """Color for the diff-flag / SE determination column."""
    if eq_status in ("NOT_EQUIVALENT", "TOXIC"):
        return COLOR_RED
    if eq_status == "DISCUSSION_REQUIRED":
        return COLOR_YELLOW
    if eq_status == "EQUIVALENT":
        return COLOR_GREEN
    return COLOR_LIGHT_GRAY


# ═══════════════════════════════════════════════════════════════════════════════
# 1.  Structured render plan (JSON-friendly, no python-docx dependency)
# ═══════════════════════════════════════════════════════════════════════════════

def render_se_matrix_table(
    comparison_rows: list[dict[str, Any]],
    subject_device_name: str = "Subject Device",
    predicate_device_name: str = "Predicate Device",
    predicate_k_number: str = "",
) -> dict[str, Any]:
    """Produce a rendering instruction set for the SE matrix table.

    This is the JSON-friendly representation that can be sent to a frontend
    or consumed by the DOCX Factory runner.

    Returns:
        {
            "table_header": [...],
            "table_rows": [...],
            "bookmarks": [...],
            "color_map": {...},
        }
    """
    header = [
        "Characteristic",
        f"Subject: {subject_device_name}",
        f"Predicate: {predicate_device_name} ({predicate_k_number})",
        "Diff Flag",
        "Discussion",
        "Evidence Status",
    ]

    rows: list[dict[str, Any]] = []
    bookmarks: list[dict[str, str]] = []
    color_cells: list[dict[str, Any]] = []

    for idx, row in enumerate(comparison_rows):
        subj_val = row.get("subject_value", {})
        pred_val = row.get("predicate_value", {})

        # Extract values
        subj_text = subj_val.get("value", "N/A") if isinstance(subj_val, dict) else str(subj_val)
        pred_text = pred_val.get("value", "N/A") if isinstance(pred_val, dict) else str(pred_val)

        subj_evidence = subj_val.get("evidence_ids", []) if isinstance(subj_val, dict) else []
        pred_evidence = pred_val.get("evidence_ids", []) if isinstance(pred_val, dict) else []

        subj_has_evidence = bool(subj_evidence)
        pred_has_evidence = bool(pred_evidence)

        eq_status = row.get("equivalence_status", "PENDING")
        diff_sev = row.get("diff_severity", "none")
        discussion = row.get("discussion_text", "")
        requires_citation = row.get("requires_citation", False)

        # Build evidence status text
        if subj_has_evidence and pred_has_evidence:
            evidence_text = "✓ Linked"
        elif subj_has_evidence:
            evidence_text = "⚠ Predicate evidence missing"
        elif pred_has_evidence:
            evidence_text = "⚠ Subject evidence missing"
        else:
            evidence_text = "✗ No evidence linked"

        # Determine colors
        subj_color = get_cell_highlight(eq_status, subj_has_evidence, diff_sev)
        pred_color = get_cell_highlight(eq_status, pred_has_evidence, diff_sev)
        status_color = get_cell_highlight(eq_status, subj_has_evidence and pred_has_evidence, diff_sev)

        table_row = {
            "index": idx,
            "characteristic": row.get("characteristic", ""),
            "category": row.get("category", ""),
            "subject_text": subj_text,
            "predicate_text": pred_text,
            "diff_flag": DIFF_LABELS.get(eq_status, eq_status),
            "discussion": discussion,
            "evidence_status": evidence_text,
            "diff_severity": diff_sev,
            "severity_badge": SEVERITY_BADGES.get(diff_sev, diff_sev),
        }
        rows.append(table_row)

        # Color map
        color_cells.append({
            "row": idx,
            "subject_bg": subj_color,
            "predicate_bg": pred_color,
            "status_bg": status_color,
            "diff_flag_bg": COLOR_RED if eq_status in ("NOT_EQUIVALENT", "TOXIC") else (
                COLOR_YELLOW if eq_status == "DISCUSSION_REQUIRED" else COLOR_GREEN
            ),
        })

        # Bookmarks for defense packet extraction
        for ev_id in subj_evidence:
            bookmarks.append({
                "name": f"EV_{ev_id}",
                "row": idx,
                "column": "subject",
                "evidence_id": ev_id,
            })
        for ev_id in pred_evidence:
            bookmarks.append({
                "name": f"EV_{ev_id}",
                "row": idx,
                "column": "predicate",
                "evidence_id": ev_id,
            })

    return {
        "table_header": header,
        "table_rows": rows,
        "bookmarks": bookmarks,
        "color_map": color_cells,
        "total_rows": len(rows),
        "equivalent_count": sum(1 for r in comparison_rows if r.get("equivalence_status") == "EQUIVALENT"),
        "discussion_count": sum(1 for r in comparison_rows if r.get("equivalence_status") == "DISCUSSION_REQUIRED"),
        "not_equivalent_count": sum(1 for r in comparison_rows if r.get("equivalence_status") == "NOT_EQUIVALENT"),
    }


# ═══════════════════════════════════════════════════════════════════════════════
# 2.  Real python-docx table renderer
# ═══════════════════════════════════════════════════════════════════════════════

def render_to_docx_table(
    doc: Document,
    comparison_rows: list[dict[str, Any]],
    subject_device_name: str = "Subject Device",
    predicate_device_name: str = "Predicate Device",
    predicate_k_number: str = "",
) -> Table:
    """Render an actual python-docx Table into the provided Document.

    Creates a formatted SE matrix table with:
    - Color-coded cells (green/yellow/red)
    - Evidence ID annotations in small blue text
    - EV_ bookmarks for defense packet extraction
    - Severity badges

    Args:
        doc: python-docx Document to add the table to
        comparison_rows: From SEMatrixGenerator.generate_payload()["comparison_rows"]
        subject_device_name: Name for subject column header
        predicate_device_name: Name for predicate column header
        predicate_k_number: K-number for predicate column header

    Returns:
        The created python-docx Table object
    """
    headers = [
        "#",
        "Characteristic",
        f"Subject: {subject_device_name}",
        f"Predicate: {predicate_device_name}\n({predicate_k_number})",
        "SE Determination",
        "Severity",
    ]

    num_rows = 1 + len(comparison_rows)
    table = doc.add_table(rows=num_rows, cols=len(headers))
    table.autofit = True

    # ── Header row ────────────────────────────────────────────────
    for col_idx, text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = text
        _apply_cell_shading(cell, COLOR_HEADER_BG)
        _apply_cell_borders(cell, "002060")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ── Data rows ─────────────────────────────────────────────────
    for row_idx, row_data in enumerate(comparison_rows):
        tr = table.rows[row_idx + 1]

        subj_val = row_data.get("subject_value", {})
        pred_val = row_data.get("predicate_value", {})
        eq_status = row_data.get("equivalence_status", "PENDING")
        diff_sev = row_data.get("diff_severity", "none")

        subj_text = subj_val.get("value", "N/A") if isinstance(subj_val, dict) else str(subj_val)
        pred_text = pred_val.get("value", "N/A") if isinstance(pred_val, dict) else str(pred_val)
        subj_ev = subj_val.get("evidence_ids", []) if isinstance(subj_val, dict) else []
        pred_ev = pred_val.get("evidence_ids", []) if isinstance(pred_val, dict) else []

        # Col 0: row number
        tr.cells[0].text = str(row_data.get("sort_order", row_idx + 1))
        tr.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Col 1: characteristic
        tr.cells[1].text = row_data.get("characteristic", "")
        for run in tr.cells[1].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

        # Col 2: subject value + evidence bookmarks
        _render_value_cell(tr.cells[2], subj_text, subj_ev, eq_status, diff_sev)

        # Col 3: predicate value + evidence bookmarks
        _render_value_cell(tr.cells[3], pred_text, pred_ev, eq_status, diff_sev)

        # Col 4: SE Determination
        tr.cells[4].text = DIFF_LABELS.get(eq_status, eq_status)
        tr.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _apply_cell_shading(tr.cells[4], _status_column_color(eq_status))
        for run in tr.cells[4].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

        # Col 5: severity
        tr.cells[5].text = SEVERITY_BADGES.get(diff_sev, diff_sev)
        tr.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if diff_sev in ("critical", "high"):
            _apply_cell_shading(tr.cells[5], COLOR_RED)
        elif diff_sev == "medium":
            _apply_cell_shading(tr.cells[5], COLOR_YELLOW)

        # Borders + vertical centering for all cells
        for cell in tr.cells:
            _apply_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Alternating row stripe on non-colored columns
        if row_idx % 2 == 1:
            for ci in (0, 1):
                _apply_cell_shading(tr.cells[ci], COLOR_LIGHT_GRAY)

    return table


def _render_value_cell(
    cell: _Cell,
    text: str,
    evidence_ids: list[str],
    eq_status: str,
    diff_severity: str,
) -> None:
    """Render a subject/predicate value cell with evidence annotations."""
    p = cell.paragraphs[0]
    p.clear()

    # Main value text
    run = p.add_run(text)
    run.font.size = Pt(9)

    # Evidence IDs in small blue text
    if evidence_ids:
        ev_run = p.add_run(f"\n[{', '.join(evidence_ids)}]")
        ev_run.font.size = Pt(7)
        ev_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
        # Insert bookmarks for each evidence ID
        for ev_id in evidence_ids:
            _insert_bookmark(p, f"EV_{ev_id}")

    # Background color
    has_evidence = bool(evidence_ids)
    _apply_cell_shading(cell, get_cell_highlight(eq_status, has_evidence, diff_severity))
