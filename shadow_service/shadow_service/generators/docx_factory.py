"""SE Matrix DOCX Factory — Phase 6.6.C.

Produces a professional FDA 510(k) Substantial Equivalence comparison
document using python-docx.  Every cell is evidence-linked with
green/yellow/red confidence highlighting and EV_ bookmarks.

Document structure:
  1.  Cover page (subject device, predicate, date, readiness score)
  2.  Executive summary (equiv counts, readiness gauge)
  3.  SE comparison matrix table (9 characteristics × 6 columns)
  4.  Per-characteristic discussion notes
  5.  Reviewer questions appendix
  6.  Toxicity warnings appendix
  7.  Evidence traceability index
  8.  Footer: manifest hash + generation metadata

Usage:
    factory = SEMatrixDocxFactory()
    buf = factory.render(se_payload, reviewer_questions, toxicity_warnings)
    # buf is io.BytesIO ready for streaming
"""
from __future__ import annotations

import io
import logging
from datetime import datetime, timezone
from typing import Any, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.table import _Cell

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Color palette (RGB hex)
# ─────────────────────────────────────────────────────────────────────────────
COLORS = {
    "green":      "C6EFCE",
    "dark_green":  "006100",
    "yellow":     "FFEB9C",
    "dark_yellow": "9C6500",
    "red":        "FFC7CE",
    "dark_red":    "9C0006",
    "white":      "FFFFFF",
    "light_gray": "F2F2F2",
    "dark_gray":  "404040",
    "header_bg":  "002060",
    "header_fg":  "FFFFFF",
    "accent":     "0070C0",
    "border":     "BFBFBF",
}

# Diff flag display labels
EQUIV_LABELS = {
    "EQUIVALENT":           "≡ Equivalent",
    "DISCUSSION_REQUIRED":  "⚠ Discussion Required",
    "NOT_EQUIVALENT":       "✗ Not Equivalent",
    "TOXIC":                "☠ Toxic",
    "PENDING":              "⏳ Pending",
}

SEVERITY_LABELS = {
    "none":     "—",
    "low":      "Low",
    "medium":   "Medium",
    "high":     "High",
    "critical": "⚠ Critical",
}


def _set_cell_shading(cell: _Cell, color_hex: str) -> None:
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_borders(cell: _Cell, color: str = "BFBFBF", size: str = "4") -> None:
    """Apply borders to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)
        element.set(qn("w:space"), "0")
        tc_borders.append(element)
    tc_pr.append(tc_borders)


def _set_cell_width(cell: _Cell, width_cm: float) -> None:
    """Set cell width."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))  # cm to twips
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _add_bookmark(paragraph, bookmark_name: str) -> None:
    """Insert an invisible bookmark at the end of a paragraph."""
    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(abs(hash(bookmark_name)) % 10000))
    bookmark_start.set(qn("w:name"), bookmark_name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(abs(hash(bookmark_name)) % 10000))

    paragraph._element.append(bookmark_start)
    paragraph._element.append(bookmark_end)


def _cell_color_for_status(
    eq_status: str,
    has_evidence: bool,
    diff_severity: str = "none",
) -> str:
    """Determine background color for a matrix cell."""
    if eq_status in ("NOT_EQUIVALENT", "TOXIC"):
        return COLORS["red"]
    if diff_severity == "critical":
        return COLORS["red"]
    if eq_status == "DISCUSSION_REQUIRED":
        return COLORS["yellow"] if not has_evidence else COLORS["green"]
    if eq_status == "EQUIVALENT":
        return COLORS["green"] if has_evidence else COLORS["white"]
    return COLORS["yellow"] if not has_evidence else COLORS["white"]


def _status_color(eq_status: str) -> str:
    """Color for the diff flag column."""
    if eq_status in ("NOT_EQUIVALENT", "TOXIC"):
        return COLORS["red"]
    if eq_status == "DISCUSSION_REQUIRED":
        return COLORS["yellow"]
    if eq_status == "EQUIVALENT":
        return COLORS["green"]
    return COLORS["light_gray"]


class SEMatrixDocxFactory:
    """Renders a complete 510(k) SE Matrix DOCX from structured payload."""

    def render(
        self,
        se_payload: dict[str, Any],
        reviewer_questions: list[dict[str, Any]] | None = None,
        toxicity_warnings: list[dict[str, Any]] | None = None,
        manifest_hash: str = "",
    ) -> io.BytesIO:
        """Produce the DOCX document.

        Args:
            se_payload: Output of SEMatrixGenerator.generate_payload()
            reviewer_questions: From /reviewer-questions endpoint
            toxicity_warnings: From toxic-detail or defense-preview
            manifest_hash: SHA-256 from defense manifest for footer

        Returns:
            BytesIO buffer containing the DOCX file
        """
        reviewer_questions = reviewer_questions or []
        toxicity_warnings = toxicity_warnings or []

        doc = Document()
        self._setup_styles(doc)

        # 1. Cover page
        self._render_cover_page(doc, se_payload)

        # 2. Executive summary
        self._render_executive_summary(doc, se_payload)

        # 3. SE comparison matrix table
        self._render_matrix_table(doc, se_payload)

        # 4. Per-characteristic discussion
        self._render_discussions(doc, se_payload)

        # 5. Reviewer questions appendix
        if reviewer_questions:
            self._render_reviewer_questions(doc, reviewer_questions)

        # 6. Toxicity warnings appendix
        if toxicity_warnings:
            self._render_toxicity_warnings(doc, toxicity_warnings)

        # 7. Evidence traceability index
        self._render_evidence_index(doc, se_payload)

        # 8. Footer metadata
        self._render_footer(doc, se_payload, manifest_hash)

        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # ─────────────────────────────────────────────────────────────────
    # Style setup
    # ─────────────────────────────────────────────────────────────────

    def _setup_styles(self, doc: Document) -> None:
        """Configure document defaults and styles."""
        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x40, 0x40, 0x40)

        # Narrow margins for table-heavy document
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            # Landscape for wide table
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)

    # ─────────────────────────────────────────────────────────────────
    # 1. Cover page
    # ─────────────────────────────────────────────────────────────────

    def _render_cover_page(self, doc: Document, payload: dict) -> None:
        """Professional cover page with device info and readiness score."""
        # Title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\n\n\n")
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("510(k) Substantial Equivalence Comparison")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Predicate Intelligence Matrix with Evidence Linkage")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

        doc.add_paragraph("")

        # Device info table
        info_table = doc.add_table(rows=5, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("Subject Device", payload.get("device_name", "—")),
            ("Predicate Device", payload.get("predicate_device_name", "—")),
            ("Predicate K-Number", payload.get("predicate_k_number", "—")),
            ("Defense Readiness Score",
             f'{payload.get("defense_readiness_score", 0):.1f} / 100'),
            ("Generated",
             payload.get("generation_timestamp", datetime.now(timezone.utc).isoformat())[:10]),
        ]
        for idx, (label, value) in enumerate(info_data):
            row = info_table.rows[idx]
            cell_label = row.cells[0]
            cell_value = row.cells[1]
            cell_label.text = label
            cell_value.text = str(value)
            # Style label cell
            for para in cell_label.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            for para in cell_value.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
            _set_cell_borders(cell_label)
            _set_cell_borders(cell_value)
            _set_cell_shading(cell_label, COLORS["light_gray"])

        # Readiness color highlight
        readiness = payload.get("defense_readiness_score", 0)
        readiness_cell = info_table.rows[3].cells[1]
        if readiness >= 75:
            _set_cell_shading(readiness_cell, COLORS["green"])
        elif readiness >= 50:
            _set_cell_shading(readiness_cell, COLORS["yellow"])
        else:
            _set_cell_shading(readiness_cell, COLORS["red"])

        # Page break
        doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────
    # 2. Executive summary
    # ─────────────────────────────────────────────────────────────────

    def _render_executive_summary(self, doc: Document, payload: dict) -> None:
        """Summary stats: equivalence counts, evidence coverage."""
        doc.add_heading("Executive Summary", level=1)

        rows = payload.get("comparison_rows", [])
        total = len(rows)
        equiv = sum(1 for r in rows if r.get("equivalence_status") == "EQUIVALENT")
        disc = sum(1 for r in rows if r.get("equivalence_status") == "DISCUSSION_REQUIRED")
        not_eq = sum(1 for r in rows if r.get("equivalence_status") == "NOT_EQUIVALENT")
        pending = total - equiv - disc - not_eq

        # Evidence coverage
        with_evidence = sum(
            1 for r in rows
            if r.get("subject_value", {}).get("evidence_ids")
        )

        # Summary table
        summary = doc.add_table(rows=6, cols=2)
        summary.alignment = WD_TABLE_ALIGNMENT.CENTER
        stats = [
            ("Total Characteristics Compared", str(total)),
            ("Equivalent", str(equiv)),
            ("Discussion Required", str(disc)),
            ("Not Equivalent", str(not_eq)),
            ("Pending Review", str(pending)),
            ("Evidence Coverage", f"{with_evidence}/{total} ({100 * with_evidence // max(total, 1)}%)"),
        ]
        stat_colors = [
            COLORS["light_gray"],
            COLORS["green"],
            COLORS["yellow"],
            COLORS["red"],
            COLORS["light_gray"],
            COLORS["green"] if with_evidence == total else COLORS["yellow"],
        ]
        for idx, ((label, value), bg) in enumerate(zip(stats, stat_colors)):
            row = summary.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
            for c in row.cells:
                _set_cell_borders(c)
                for para in c.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
            row.cells[0].paragraphs[0].runs[0].bold = True
            _set_cell_shading(row.cells[1], bg)

        doc.add_paragraph("")

        # Readiness interpretation
        readiness = payload.get("defense_readiness_score", 0)
        if readiness >= 80:
            assessment = (
                "This submission demonstrates strong substantial equivalence "
                "with comprehensive evidence linkage. The predicate selection "
                "is well-supported and no critical gaps were identified."
            )
        elif readiness >= 60:
            assessment = (
                "The submission shows moderate substantial equivalence. "
                "Some characteristics require additional discussion or evidence "
                "before the defense packet is fully ready for FDA review."
            )
        elif readiness >= 40:
            assessment = (
                "Significant gaps exist in the substantial equivalence argument. "
                "Multiple characteristics require discussion with supporting evidence. "
                "Address the items flagged below before submitting."
            )
        else:
            assessment = (
                "The substantial equivalence argument is incomplete. "
                "Critical differences and missing evidence must be resolved. "
                "Consider re-evaluating the predicate selection."
            )

        p = doc.add_paragraph()
        run = p.add_run("Assessment: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(assessment)
        run.font.size = Pt(10)

        doc.add_paragraph("")

    # ─────────────────────────────────────────────────────────────────
    # 3. SE comparison matrix table
    # ─────────────────────────────────────────────────────────────────

    def _render_matrix_table(self, doc: Document, payload: dict) -> None:
        """The main comparison table with color-coded cells and bookmarks."""
        doc.add_heading("Substantial Equivalence Comparison Matrix", level=1)

        rows = payload.get("comparison_rows", [])
        device_name = payload.get("device_name", "Subject Device")
        pred_name = payload.get("predicate_device_name", "Predicate Device")
        pred_k = payload.get("predicate_k_number", "")

        # Header columns
        headers = [
            "#",
            "Characteristic",
            f"Subject: {device_name}",
            f"Predicate: {pred_name}\n({pred_k})",
            "SE Determination",
            "Severity",
        ]

        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Style header row
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = header_text
            _set_cell_shading(cell, COLORS["header_bg"])
            _set_cell_borders(cell, "002060")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Data rows
        for row_idx, row_data in enumerate(rows):
            table_row = table.rows[row_idx + 1]

            subj_val = row_data.get("subject_value", {})
            pred_val = row_data.get("predicate_value", {})
            eq_status = row_data.get("equivalence_status", "PENDING")
            diff_sev = row_data.get("diff_severity", "none")

            subj_text = subj_val.get("value", "N/A") if isinstance(subj_val, dict) else str(subj_val)
            pred_text = pred_val.get("value", "N/A") if isinstance(pred_val, dict) else str(pred_val)
            subj_evidence = subj_val.get("evidence_ids", []) if isinstance(subj_val, dict) else []
            pred_evidence = pred_val.get("evidence_ids", []) if isinstance(pred_val, dict) else []

            # Column 0: Row number
            cell_num = table_row.cells[0]
            cell_num.text = str(row_data.get("sort_order", row_idx + 1))
            cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Column 1: Characteristic name
            cell_char = table_row.cells[1]
            cell_char.text = row_data.get("characteristic", "")
            cell_char.paragraphs[0].runs[0].bold = True if cell_char.paragraphs[0].runs else False

            # Column 2: Subject value (with evidence bookmarks)
            cell_subj = table_row.cells[2]
            p = cell_subj.paragraphs[0]
            p.clear()
            run = p.add_run(subj_text)
            run.font.size = Pt(9)
            if subj_evidence:
                ev_run = p.add_run(f"\n[{', '.join(subj_evidence)}]")
                ev_run.font.size = Pt(7)
                ev_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                for ev_id in subj_evidence:
                    _add_bookmark(p, f"EV_{ev_id}")
            _set_cell_shading(cell_subj, _cell_color_for_status(
                eq_status, bool(subj_evidence), diff_sev))

            # Column 3: Predicate value
            cell_pred = table_row.cells[3]
            p = cell_pred.paragraphs[0]
            p.clear()
            run = p.add_run(pred_text)
            run.font.size = Pt(9)
            if pred_evidence:
                ev_run = p.add_run(f"\n[{', '.join(pred_evidence)}]")
                ev_run.font.size = Pt(7)
                ev_run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
                for ev_id in pred_evidence:
                    _add_bookmark(p, f"EV_{ev_id}")
            _set_cell_shading(cell_pred, _cell_color_for_status(
                eq_status, bool(pred_evidence), diff_sev))

            # Column 4: SE Determination
            cell_status = table_row.cells[4]
            cell_status.text = EQUIV_LABELS.get(eq_status, eq_status)
            cell_status.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_shading(cell_status, _status_color(eq_status))
            for run in cell_status.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

            # Column 5: Severity
            cell_sev = table_row.cells[5]
            cell_sev.text = SEVERITY_LABELS.get(diff_sev, diff_sev)
            cell_sev.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if diff_sev in ("critical", "high"):
                _set_cell_shading(cell_sev, COLORS["red"])
            elif diff_sev == "medium":
                _set_cell_shading(cell_sev, COLORS["yellow"])

            # Apply borders to all cells in row
            for cell in table_row.cells:
                _set_cell_borders(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Alternating row shading for non-highlighted cells
            if row_idx % 2 == 1:
                for ci in (0, 1):
                    if table_row.cells[ci]._tc.find(qn("w:tcPr/w:shd")) is None:
                        _set_cell_shading(table_row.cells[ci], COLORS["light_gray"])

        doc.add_paragraph("")

        # Color legend
        legend = doc.add_paragraph()
        legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = legend.add_run("Legend: ")
        run.bold = True
        run.font.size = Pt(8)
        for label, color in [
            ("■ Equivalent + Evidence", COLORS["green"]),
            ("  ■ Discussion Required", COLORS["yellow"]),
            ("  ■ Not Equivalent / Critical", COLORS["red"]),
        ]:
            run = legend.add_run(label)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(
                int(color[:2], 16), int(color[2:4], 16), int(color[4:6], 16),
            )

        doc.add_paragraph("")

    # ─────────────────────────────────────────────────────────────────
    # 4. Per-characteristic discussion
    # ─────────────────────────────────────────────────────────────────

    def _render_discussions(self, doc: Document, payload: dict) -> None:
        """Expanded discussion for each non-equivalent characteristic."""
        rows_needing_discussion = [
            r for r in payload.get("comparison_rows", [])
            if r.get("equivalence_status") in ("DISCUSSION_REQUIRED", "NOT_EQUIVALENT")
            or r.get("requires_citation")
        ]

        if not rows_needing_discussion:
            return

        doc.add_heading("Detailed Discussion by Characteristic", level=1)

        for row in rows_needing_discussion:
            char_name = row.get("characteristic", "Unknown")
            eq_status = row.get("equivalence_status", "PENDING")
            discussion = row.get("discussion_text", "")
            suggested_tests = row.get("suggested_tests", [])

            # Characteristic heading
            p = doc.add_heading(f"{char_name}", level=2)

            # Status badge line
            p = doc.add_paragraph()
            run = p.add_run(f"Status: {EQUIV_LABELS.get(eq_status, eq_status)}")
            run.bold = True
            run.font.size = Pt(10)
            if eq_status == "NOT_EQUIVALENT":
                run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
            elif eq_status == "DISCUSSION_REQUIRED":
                run.font.color.rgb = RGBColor(0x9C, 0x65, 0x00)

            # Discussion text
            if discussion:
                p = doc.add_paragraph()
                run = p.add_run(discussion)
                run.font.size = Pt(10)

            # Suggested tests
            if suggested_tests:
                p = doc.add_paragraph()
                run = p.add_run("Recommended Testing/Evidence:")
                run.bold = True
                run.font.size = Pt(10)
                for test in suggested_tests:
                    p = doc.add_paragraph(f"  • {test}", style="List Bullet")

            doc.add_paragraph("")

    # ─────────────────────────────────────────────────────────────────
    # 5. Reviewer questions appendix
    # ─────────────────────────────────────────────────────────────────

    def _render_reviewer_questions(
        self, doc: Document, questions: list[dict],
    ) -> None:
        """Appendix: Anticipated FDA Reviewer Questions."""
        doc.add_page_break()
        doc.add_heading("Appendix A: Anticipated FDA Reviewer Questions", level=1)

        p = doc.add_paragraph()
        run = p.add_run(
            "The following questions are predicted based on deterministic "
            "regulatory analysis of the differences between the subject "
            "device and the selected predicate. Each question includes "
            "the applicable regulatory citation and recommended evidence."
        )
        run.font.size = Pt(9)
        run.italic = True

        doc.add_paragraph("")

        # Questions table
        table = doc.add_table(rows=1 + len(questions), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        header_texts = ["Severity", "Question", "Citation", "Required Evidence"]
        for idx, text in enumerate(header_texts):
            cell = table.rows[0].cells[idx]
            cell.text = text
            _set_cell_shading(cell, COLORS["header_bg"])
            _set_cell_borders(cell, "002060")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, q in enumerate(questions):
            row = table.rows[idx + 1]

            # Severity
            sev = q.get("severity", "medium")
            row.cells[0].text = sev.upper()
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            sev_color = {
                "critical": COLORS["red"],
                "high": COLORS["yellow"],
                "medium": COLORS["light_gray"],
                "low": COLORS["white"],
            }.get(sev, COLORS["white"])
            _set_cell_shading(row.cells[0], sev_color)
            for run in row.cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(9)

            # Question
            row.cells[1].text = q.get("question", "")
            for run in row.cells[1].paragraphs[0].runs:
                run.font.size = Pt(9)

            # Citation
            row.cells[2].text = q.get("citation", "")
            for run in row.cells[2].paragraphs[0].runs:
                run.font.size = Pt(8)
                run.italic = True

            # Required evidence
            evidence = q.get("required_evidence", [])
            row.cells[3].text = "\n".join(evidence) if evidence else "—"
            for run in row.cells[3].paragraphs[0].runs:
                run.font.size = Pt(8)

            for cell in row.cells:
                _set_cell_borders(cell)

    # ─────────────────────────────────────────────────────────────────
    # 6. Toxicity warnings appendix
    # ─────────────────────────────────────────────────────────────────

    def _render_toxicity_warnings(
        self, doc: Document, warnings: list[dict],
    ) -> None:
        """Appendix: Predicate Toxicity Warnings."""
        doc.add_page_break()
        doc.add_heading("Appendix B: Predicate Toxicity Warnings", level=1)

        p = doc.add_paragraph()
        run = p.add_run(
            "⚠ The selected predicate has safety signals that may "
            "impact the 510(k) submission strategy. Review each "
            "warning and prepare mitigations."
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x9C, 0x00, 0x06)
        run.italic = True

        doc.add_paragraph("")

        table = doc.add_table(rows=1 + len(warnings), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["K-Number", "Signal Type", "Date", "Severity", "Description"]
        for idx, text in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = text
            _set_cell_shading(cell, COLORS["red"])
            _set_cell_borders(cell)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, w in enumerate(warnings):
            row = table.rows[idx + 1]
            row.cells[0].text = w.get("k_number", "")
            row.cells[1].text = w.get("signal_type", "")
            row.cells[2].text = str(w.get("signal_date", "—"))
            row.cells[3].text = f'{w.get("severity_score", 0):.2f}'
            row.cells[4].text = (w.get("description", ""))[:300]

            for cell in row.cells:
                _set_cell_borders(cell)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

            # Color severity cell
            sev = w.get("severity_score", 0)
            if sev >= 0.8:
                _set_cell_shading(row.cells[3], COLORS["red"])
            elif sev >= 0.5:
                _set_cell_shading(row.cells[3], COLORS["yellow"])

    # ─────────────────────────────────────────────────────────────────
    # 7. Evidence traceability index
    # ─────────────────────────────────────────────────────────────────

    def _render_evidence_index(self, doc: Document, payload: dict) -> None:
        """Index of all evidence IDs referenced in the matrix."""
        all_evidence: list[dict[str, str]] = []

        for row in payload.get("comparison_rows", []):
            char_name = row.get("characteristic", "")
            subj = row.get("subject_value", {})
            pred = row.get("predicate_value", {})

            for ev_id in (subj.get("evidence_ids", []) if isinstance(subj, dict) else []):
                all_evidence.append({
                    "evidence_id": ev_id,
                    "characteristic": char_name,
                    "source": "Subject",
                    "bookmark": f"EV_{ev_id}",
                })
            for ev_id in (pred.get("evidence_ids", []) if isinstance(pred, dict) else []):
                all_evidence.append({
                    "evidence_id": ev_id,
                    "characteristic": char_name,
                    "source": "Predicate",
                    "bookmark": f"EV_{ev_id}",
                })

        if not all_evidence:
            return

        doc.add_page_break()
        doc.add_heading("Evidence Traceability Index", level=1)

        p = doc.add_paragraph()
        run = p.add_run(
            "Every cell in the SE matrix is linked to evidence sources. "
            "Each entry below corresponds to a bookmark (EV_*) in the "
            "comparison table that can be extracted for defense packet assembly."
        )
        run.font.size = Pt(9)
        run.italic = True

        doc.add_paragraph("")

        table = doc.add_table(rows=1 + len(all_evidence), cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, text in enumerate(["Evidence ID", "Characteristic", "Source", "Bookmark"]):
            cell = table.rows[0].cells[idx]
            cell.text = text
            _set_cell_shading(cell, COLORS["header_bg"])
            _set_cell_borders(cell, "002060")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, ev in enumerate(all_evidence):
            row = table.rows[idx + 1]
            row.cells[0].text = ev["evidence_id"]
            row.cells[1].text = ev["characteristic"]
            row.cells[2].text = ev["source"]
            row.cells[3].text = ev["bookmark"]
            for cell in row.cells:
                _set_cell_borders(cell)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

    # ─────────────────────────────────────────────────────────────────
    # 8. Footer metadata
    # ─────────────────────────────────────────────────────────────────

    def _render_footer(
        self, doc: Document, payload: dict, manifest_hash: str,
    ) -> None:
        """Document provenance and chain-of-custody metadata."""
        doc.add_paragraph("")
        p = doc.add_paragraph()
        run = p.add_run("─" * 80)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)

        meta_lines = [
            f"Template: {payload.get('template_id', '510k_se_matrix_v2')} v{payload.get('version', '6.6.0')}",
            f"Generated: {payload.get('generation_timestamp', datetime.now(timezone.utc).isoformat())}",
            f"Regulatory Standard: {payload.get('regulatory_standard', 'FDA_510k_SE')}",
            f"Defense Readiness: {payload.get('defense_readiness_score', 0):.1f}/100",
            "Generator: ClinicalSage Predicate Intelligence — Phase 6.6",
        ]
        if manifest_hash:
            meta_lines.append(f"Manifest SHA-256: {manifest_hash}")

        for line in meta_lines:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.italic = True

        p = doc.add_paragraph()
        run = p.add_run(
            "This document was electronically generated. Evidence linkage "
            "bookmarks (EV_*) are embedded for defense packet extraction."
        )
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.italic = True
