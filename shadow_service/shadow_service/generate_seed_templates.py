"""Generate seed DOCX template files with {{ variable }} placeholders.

Run once to create the actual .docx files in demo_templates/.
These are committed to the repo so the seed script can load them.

Usage:
    python -m shadow_service.generate_seed_templates
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


TEMPLATES_DIR = Path(__file__).parent / "demo_templates"


def _add_heading(doc: Document, text: str, level: int = 0) -> None:
    doc.add_heading(text, level=level)


def _add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def _add_table_row(table, cells: list[str]) -> None:
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val


# ---------------------------------------------------------------------------
# 1. eCTD Cover Letter
# ---------------------------------------------------------------------------
def make_ectd_cover_letter() -> None:
    doc = Document()
    _add_heading(doc, "eCTD Cover Letter", level=0)
    _add_para(doc, "Date: {{ submission_date }}")
    _add_para(doc, "")
    _add_para(doc, "{{ regulatory_authority }}")
    _add_para(doc, "{{ authority_address }}")
    _add_para(doc, "")
    _add_para(doc, "Re: {{ submission_type }} Submission for {{ product_name }}")
    _add_para(doc, "Application Number: {{ application_number }}")
    _add_para(doc, "Sequence Number: {{ sequence_number }}")
    _add_para(doc, "")
    _add_para(doc, "Dear {{ reviewer_name }},")
    _add_para(doc, "")
    _add_para(
        doc,
        "{{ sponsor_name }} hereby submits this {{ submission_type }} for "
        "{{ product_name }} ({{ generic_name }}), {{ dosage_form }}, for the "
        "proposed indication of {{ indication }}."
    )
    _add_para(doc, "")
    _add_para(doc, "This submission contains the following modules:", bold=True)
    _add_para(doc, "{{ modules_included }}")
    _add_para(doc, "")
    _add_para(
        doc,
        "Key changes in this sequence include: {{ key_changes }}"
    )
    _add_para(doc, "")
    _add_para(
        doc,
        "Should you have any questions, please contact {{ contact_name }} "
        "at {{ contact_email }} or {{ contact_phone }}."
    )
    _add_para(doc, "")
    _add_para(doc, "Sincerely,")
    _add_para(doc, "{{ signatory_name }}")
    _add_para(doc, "{{ signatory_title }}")
    _add_para(doc, "{{ sponsor_name }}")
    doc.save(str(TEMPLATES_DIR / "ectd_cover_letter.docx"))


# ---------------------------------------------------------------------------
# 2. Form FDA 1571 Narrative Summary
# ---------------------------------------------------------------------------
def make_fda_1571_narrative() -> None:
    doc = Document()
    _add_heading(doc, "IND Application — Form FDA 1571 Narrative Summary", level=0)

    _add_para(doc, "1. Sponsor Information", bold=True)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for label, var in [
        ("Sponsor Name", "{{ sponsor_name }}"),
        ("Sponsor Address", "{{ sponsor_address }}"),
        ("IND Number", "{{ ind_number }}"),
        ("Date of Submission", "{{ submission_date }}"),
        ("Product Name", "{{ product_name }}"),
        ("Phase of Investigation", "{{ phase }}"),
    ]:
        _add_table_row(table, [label, var])

    _add_para(doc, "")
    _add_para(doc, "2. Investigational Plan Summary", bold=True)
    _add_para(
        doc,
        "{{ product_name }} ({{ generic_name }}) is a {{ drug_class }} being "
        "developed for the treatment of {{ indication }}. This {{ phase }} "
        "study is designed as {{ study_design }}."
    )

    _add_para(doc, "")
    _add_para(doc, "3. Chemistry, Manufacturing, and Controls Summary", bold=True)
    _add_para(
        doc,
        "The drug substance is manufactured by {{ manufacturer }} at "
        "{{ manufacturing_site }}. The dosage form is {{ dosage_form }} "
        "with a strength of {{ dosage_strength }}."
    )

    _add_para(doc, "")
    _add_para(doc, "4. Pharmacology / Toxicology Summary", bold=True)
    _add_para(doc, "{{ nonclinical_summary }}")

    _add_para(doc, "")
    _add_para(doc, "5. Previous Human Experience", bold=True)
    _add_para(doc, "{{ previous_human_experience }}")

    _add_para(doc, "")
    _add_para(doc, "6. Investigator Information", bold=True)
    _add_para(doc, "Principal Investigator: {{ principal_investigator }}")
    _add_para(doc, "Clinical Site: {{ clinical_site }}")

    doc.save(str(TEMPLATES_DIR / "fda_1571_narrative.docx"))


# ---------------------------------------------------------------------------
# 3. Investigator Brochure Change Summary
# ---------------------------------------------------------------------------
def make_ib_change_summary() -> None:
    doc = Document()
    _add_heading(doc, "Investigator's Brochure — Change Summary", level=0)

    _add_para(doc, "Product: {{ product_name }} ({{ generic_name }})")
    _add_para(doc, "IB Edition: {{ ib_edition }}")
    _add_para(doc, "Effective Date: {{ effective_date }}")
    _add_para(doc, "Sponsor: {{ sponsor_name }}")
    _add_para(doc, "")

    _add_para(doc, "1. Summary of Changes", bold=True)
    _add_para(
        doc,
        "This edition of the Investigator's Brochure for {{ product_name }} "
        "incorporates the following changes from the previous edition "
        "({{ previous_edition }}):"
    )
    _add_para(doc, "")
    _add_para(doc, "{{ changes_summary }}")

    _add_para(doc, "")
    _add_para(doc, "2. New Safety Data", bold=True)
    _add_para(doc, "{{ new_safety_data }}")

    _add_para(doc, "")
    _add_para(doc, "3. Updated Efficacy Information", bold=True)
    _add_para(doc, "{{ updated_efficacy }}")

    _add_para(doc, "")
    _add_para(doc, "4. Revised Dosing Information", bold=True)
    _add_para(doc, "{{ revised_dosing }}")

    _add_para(doc, "")
    _add_para(doc, "5. Impact Assessment", bold=True)
    _add_para(
        doc,
        "The changes described above {{ impact_assessment }}. "
        "The overall benefit-risk profile remains {{ benefit_risk_status }}."
    )

    _add_para(doc, "")
    _add_para(doc, "Prepared by: {{ prepared_by }}")
    _add_para(doc, "Reviewed by: {{ reviewed_by }}")
    _add_para(doc, "Date: {{ review_date }}")

    doc.save(str(TEMPLATES_DIR / "ib_change_summary.docx"))


# ---------------------------------------------------------------------------
# 4. CMC Drug Substance Overview
# ---------------------------------------------------------------------------
def make_cmc_drug_substance() -> None:
    doc = Document()
    _add_heading(doc, "CMC Section: Drug Substance Overview", level=0)
    _add_para(doc, "Module 3.2.S — {{ product_name }} ({{ generic_name }})")
    _add_para(doc, "")

    _add_para(doc, "S.1 General Information", bold=True)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Property"
    table.rows[0].cells[1].text = "Description"
    for label, var in [
        ("INN / Generic Name", "{{ generic_name }}"),
        ("Chemical Name", "{{ chemical_name }}"),
        ("Molecular Formula", "{{ molecular_formula }}"),
        ("Molecular Weight", "{{ molecular_weight }}"),
        ("Structural Class", "{{ structural_class }}"),
        ("CAS Number", "{{ cas_number }}"),
    ]:
        _add_table_row(table, [label, var])

    _add_para(doc, "")
    _add_para(doc, "S.2 Manufacture", bold=True)
    _add_para(doc, "Manufacturer: {{ manufacturer }}")
    _add_para(doc, "Manufacturing Site: {{ manufacturing_site }}")
    _add_para(doc, "Process Description: {{ process_description }}")

    _add_para(doc, "")
    _add_para(doc, "S.3 Characterization", bold=True)
    _add_para(doc, "{{ characterization_summary }}")

    _add_para(doc, "")
    _add_para(doc, "S.4 Control of Drug Substance", bold=True)
    _add_para(doc, "Specification: {{ specification_summary }}")
    _add_para(doc, "Analytical Methods: {{ analytical_methods }}")

    _add_para(doc, "")
    _add_para(doc, "S.5 Reference Standards", bold=True)
    _add_para(doc, "{{ reference_standards }}")

    _add_para(doc, "")
    _add_para(doc, "S.6 Container Closure", bold=True)
    _add_para(doc, "{{ container_closure }}")

    _add_para(doc, "")
    _add_para(doc, "S.7 Stability", bold=True)
    _add_para(doc, "{{ stability_summary }}")

    doc.save(str(TEMPLATES_DIR / "cmc_drug_substance.docx"))


# ---------------------------------------------------------------------------
# 5. Clinical Overview: Benefit/Risk Summary
# ---------------------------------------------------------------------------
def make_clinical_benefit_risk() -> None:
    doc = Document()
    _add_heading(doc, "Clinical Overview — Benefit/Risk Summary", level=0)
    _add_para(doc, "Module 2.5 — {{ product_name }} for {{ indication }}")
    _add_para(doc, "Sponsor: {{ sponsor_name }}")
    _add_para(doc, "Date: {{ report_date }}")
    _add_para(doc, "")

    _add_para(doc, "1. Product Overview", bold=True)
    _add_para(
        doc,
        "{{ product_name }} ({{ generic_name }}) is a {{ drug_class }} "
        "indicated for {{ indication }}. The proposed dosage is "
        "{{ proposed_dosage }}."
    )

    _add_para(doc, "")
    _add_para(doc, "2. Benefits", bold=True)
    _add_para(doc, "2.1 Primary Efficacy Evidence", bold=True)
    _add_para(doc, "{{ primary_efficacy }}")
    _add_para(doc, "2.2 Secondary Endpoints", bold=True)
    _add_para(doc, "{{ secondary_endpoints }}")
    _add_para(doc, "2.3 Patient-Reported Outcomes", bold=True)
    _add_para(doc, "{{ patient_reported_outcomes }}")

    _add_para(doc, "")
    _add_para(doc, "3. Risks", bold=True)
    _add_para(doc, "3.1 Common Adverse Events", bold=True)
    _add_para(doc, "{{ common_adverse_events }}")
    _add_para(doc, "3.2 Serious Adverse Events", bold=True)
    _add_para(doc, "{{ serious_adverse_events }}")
    _add_para(doc, "3.3 Risk Mitigation Measures", bold=True)
    _add_para(doc, "{{ risk_mitigation }}")

    _add_para(doc, "")
    _add_para(doc, "4. Benefit-Risk Conclusion", bold=True)
    _add_para(doc, "{{ benefit_risk_conclusion }}")

    _add_para(doc, "")
    _add_para(doc, "5. Unmet Medical Need", bold=True)
    _add_para(doc, "{{ unmet_medical_need }}")

    doc.save(str(TEMPLATES_DIR / "clinical_benefit_risk.docx"))


# ---------------------------------------------------------------------------
# 6. 510(k) Cover Letter + Administrative Summary
# ---------------------------------------------------------------------------
def make_510k_cover_letter() -> None:
    doc = Document()
    _add_heading(doc, "510(k) Premarket Notification — Cover Letter", level=0)
    _add_para(doc, "Date: {{ submission_date }}")
    _add_para(doc, "")
    _add_para(doc, "U.S. Food and Drug Administration")
    _add_para(doc, "Center for Devices and Radiological Health")
    _add_para(doc, "Document Control Center — WO66-G609")
    _add_para(doc, "10903 New Hampshire Avenue")
    _add_para(doc, "Silver Spring, MD 20993-0002")
    _add_para(doc, "")
    _add_para(doc, "Re: 510(k) Premarket Notification for {{ device_name }}")
    _add_para(doc, "")
    _add_para(doc, "Dear Sir/Madam,")
    _add_para(doc, "")
    _add_para(
        doc,
        "{{ submitter_name }} hereby submits this 510(k) Premarket Notification "
        "for {{ device_name }}, classified under 21 CFR {{ regulation_number }}, "
        "product code {{ product_code }}."
    )
    _add_para(doc, "")

    _add_para(doc, "Device Description", bold=True)
    _add_para(doc, "{{ device_description }}")
    _add_para(doc, "")

    _add_para(doc, "Intended Use", bold=True)
    _add_para(doc, "{{ intended_use }}")
    _add_para(doc, "")

    _add_para(doc, "Predicate Device", bold=True)
    _add_para(doc, "Predicate: {{ predicate_device }} (510(k) Number: {{ predicate_510k }})")
    _add_para(doc, "")

    _add_para(doc, "Substantial Equivalence Summary", bold=True)
    _add_para(doc, "{{ substantial_equivalence }}")
    _add_para(doc, "")

    _add_para(doc, "Administrative Information", bold=True)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for label, var in [
        ("Submitter", "{{ submitter_name }}"),
        ("Contact Person", "{{ contact_name }}"),
        ("Contact Email", "{{ contact_email }}"),
        ("Contact Phone", "{{ contact_phone }}"),
        ("Establishment Registration #", "{{ establishment_number }}"),
        ("Device Class", "{{ device_class }}"),
    ]:
        _add_table_row(table, [label, var])

    _add_para(doc, "")
    _add_para(doc, "Sincerely,")
    _add_para(doc, "{{ signatory_name }}")
    _add_para(doc, "{{ signatory_title }}")
    _add_para(doc, "{{ submitter_name }}")

    doc.save(str(TEMPLATES_DIR / "510k_cover_letter.docx"))


# =============================================================================
# Main
# =============================================================================
def main() -> None:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    generators = [
        ("eCTD Cover Letter", make_ectd_cover_letter),
        ("FDA 1571 Narrative", make_fda_1571_narrative),
        ("IB Change Summary", make_ib_change_summary),
        ("CMC Drug Substance", make_cmc_drug_substance),
        ("Clinical Benefit/Risk", make_clinical_benefit_risk),
        ("510(k) Cover Letter", make_510k_cover_letter),
    ]

    for name, gen in generators:
        gen()
        print(f"  ✓ Generated: {name}")

    print(f"\n{len(generators)} templates written to {TEMPLATES_DIR}")


if __name__ == "__main__":
    main()
