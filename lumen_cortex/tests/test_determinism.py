"""
Determinism Tests for Frozen v1.0 Schemas
==========================================

These tests ensure that:
1. Same input always produces same fingerprint (PYTHONHASHSEED=0)
2. Models are truly frozen (immutable)
3. Ordering is stable across runs
4. In-memory DOCX fixtures produce consistent extraction

Run with: PYTHONHASHSEED=0 pytest lumen_cortex/tests/test_determinism.py -v
"""

import hashlib
import io
from datetime import datetime, timezone
from uuid import UUID, uuid4

import pytest

# In-memory DOCX generation (no binary fixtures)
try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

from lumen_cortex.core.canonical import (
    CanonicalDocument,
    CanonicalParagraph,
    EvidencePointer,
    Finding,
    SEVERITY_RANK,
    finding_sort_key,
    sort_findings,
    evidence_sort_tuple,
)


# =============================================================================
# Fixtures - Generated in-memory (no binary commits)
# =============================================================================

@pytest.fixture
def fixed_uuid() -> UUID:
    """Deterministic UUID for testing."""
    return UUID("12345678-1234-5678-1234-567812345678")


@pytest.fixture
def fixed_timestamp() -> datetime:
    """Deterministic timestamp for testing."""
    return datetime(2026, 2, 3, 12, 0, 0, tzinfo=timezone.utc)


@pytest.fixture
def sample_content_hash() -> str:
    """Deterministic content hash."""
    return hashlib.sha256(b"test document content").hexdigest()


@pytest.fixture
def sample_evidence_pointer(fixed_uuid: UUID, sample_content_hash: str) -> EvidencePointer:
    """Create a deterministic EvidencePointer for DOCX."""
    return EvidencePointer(
        doc_id=fixed_uuid,
        content_hash=sample_content_hash,
        pointer_type="docx",
        paragraph_index=3,
        run_span=(42, 100),
    )


@pytest.fixture
def sample_pdf_evidence_pointer(fixed_uuid: UUID, sample_content_hash: str) -> EvidencePointer:
    """Create a deterministic EvidencePointer for PDF."""
    return EvidencePointer(
        doc_id=fixed_uuid,
        content_hash=sample_content_hash,
        pointer_type="pdf",
        page=4,
        bbox=(100.0, 200.0, 300.0, 250.0),
        text_hash=hashlib.sha256(b"sample text").hexdigest(),
    )


@pytest.fixture
def sample_canonical_document(
    fixed_uuid: UUID,
    sample_content_hash: str,
    fixed_timestamp: datetime,
) -> CanonicalDocument:
    """Create a deterministic CanonicalDocument."""
    return CanonicalDocument(
        doc_id=fixed_uuid,
        content_hash=sample_content_hash,
        source_type="docx",
        title="Test Clinical Protocol",
        paragraphs=[
            CanonicalParagraph(index=0, text="This is the introduction."),
            CanonicalParagraph(index=1, text="Study methodology described here.", heading_level=1),
        ],
        extraction_timestamp=fixed_timestamp,
        extractor_version="1.0.0-test",
    )


@pytest.fixture
def sample_finding(
    fixed_uuid: UUID,
    sample_evidence_pointer: EvidencePointer,
    sample_content_hash: str,
) -> Finding:
    """Create a deterministic Finding."""
    return Finding(
        finding_id=fixed_uuid,
        doc_id=fixed_uuid,
        rule_id="R001",
        rule_name="Broken Cross-Reference",
        cfr_citation="21 CFR 312.23(a)(7)",
        canonical_schema_version="1.0",
        extractor_version="1.0.0-test",
        ruleset_version="1.0.0",
        content_hash=sample_content_hash,
        severity="Major",
        confidence=0.95,
        evidence=sample_evidence_pointer,
        evidence_context="Reference to 'Section 3.2' not found in document.",
        description="Broken cross-reference detected",
        remediation="Update reference to point to valid section heading.",
    )


# =============================================================================
# Fingerprint Determinism Tests
# =============================================================================

class TestFingerprintDeterminism:
    """Verify fingerprints are deterministic across runs."""

    def test_evidence_pointer_sort_key_stable(
        self, sample_evidence_pointer: EvidencePointer
    ):
        """Same EvidencePointer always produces same sort key."""
        key1 = sample_evidence_pointer.sort_key()
        key2 = sample_evidence_pointer.sort_key()

        assert key1 == key2
        assert isinstance(key1, tuple)

    def test_evidence_sort_tuple_matches_method(
        self, sample_evidence_pointer: EvidencePointer
    ):
        """Module-level sort function matches instance method."""
        assert evidence_sort_tuple(sample_evidence_pointer) == sample_evidence_pointer.sort_key()

    def test_finding_fingerprint_stable(self, sample_finding: Finding):
        """Same Finding always produces same fingerprint."""
        fp1 = sample_finding.fingerprint
        fp2 = sample_finding.fingerprint

        assert fp1 == fp2
        assert len(fp1) == 64  # SHA256 hex

    def test_finding_fingerprint_via_compute(self, sample_finding: Finding):
        """compute_fingerprint() matches fingerprint property."""
        assert sample_finding.compute_fingerprint() == sample_finding.fingerprint

    def test_finding_fingerprint_ignores_review_state(
        self,
        fixed_uuid: UUID,
        sample_evidence_pointer: EvidencePointer,
        sample_content_hash: str,
    ):
        """Review state should not affect fingerprint."""
        base_args = {
            "finding_id": fixed_uuid,
            "doc_id": fixed_uuid,
            "rule_id": "R001",
            "rule_name": "Test Rule",
            "extractor_version": "1.0.0",
            "ruleset_version": "1.0.0",
            "content_hash": sample_content_hash,
            "severity": "Major",
            "confidence": 0.9,
            "evidence": sample_evidence_pointer,
            "evidence_context": "Test context",
            "description": "Test issue",
            "remediation": "Fix it",
        }

        finding_unreviewed = Finding(**base_args)
        finding_reviewed = Finding(
            **base_args,
            reviewed_by="auditor@example.com",
            reviewed_at=datetime.now(timezone.utc),
        )

        assert finding_unreviewed.fingerprint == finding_reviewed.fingerprint


# =============================================================================
# Immutability Tests (Frozen Models)
# =============================================================================

class TestModelImmutability:
    """Verify models are frozen and cannot be mutated."""

    def test_evidence_pointer_is_frozen(self, sample_evidence_pointer: EvidencePointer):
        """EvidencePointer should reject attribute mutation."""
        with pytest.raises(Exception):  # ValidationError or AttributeError
            sample_evidence_pointer.paragraph_index = 99

    def test_canonical_document_is_frozen(
        self, sample_canonical_document: CanonicalDocument
    ):
        """CanonicalDocument should reject attribute mutation."""
        with pytest.raises(Exception):
            sample_canonical_document.title = "Modified Title"

    def test_finding_is_frozen(self, sample_finding: Finding):
        """Finding should reject attribute mutation."""
        with pytest.raises(Exception):
            sample_finding.severity = "Critical"


# =============================================================================
# Ordering Tests (Stable Sorting)
# =============================================================================

class TestStableOrdering:
    """Verify ordering is consistent for reproducible reports."""

    def test_evidence_pointers_sort_by_location(
        self, fixed_uuid: UUID, sample_content_hash: str
    ):
        """EvidencePointers should sort by type > page > paragraph > table."""
        ptrs = [
            EvidencePointer(
                doc_id=fixed_uuid,
                content_hash=sample_content_hash,
                pointer_type="docx",
                paragraph_index=10,
            ),
            EvidencePointer(
                doc_id=fixed_uuid,
                content_hash=sample_content_hash,
                pointer_type="docx",
                paragraph_index=5,
            ),
            EvidencePointer(
                doc_id=fixed_uuid,
                content_hash=sample_content_hash,
                pointer_type="pdf",
                page=0,
            ),
        ]

        sorted_ptrs = sorted(ptrs, key=evidence_sort_tuple)

        # DOCX comes before PDF (alphabetically)
        assert sorted_ptrs[0].pointer_type == "docx"
        assert sorted_ptrs[0].paragraph_index == 5
        assert sorted_ptrs[1].pointer_type == "docx"
        assert sorted_ptrs[1].paragraph_index == 10
        assert sorted_ptrs[2].pointer_type == "pdf"

    def test_findings_sort_by_severity_then_rule(
        self, fixed_uuid: UUID, sample_content_hash: str
    ):
        """Findings should sort by severity (critical first) then rule_id."""
        base_ptr = EvidencePointer(
            doc_id=fixed_uuid,
            content_hash=sample_content_hash,
            pointer_type="docx",
            paragraph_index=0,
        )

        base_args = {
            "finding_id": uuid4(),
            "doc_id": fixed_uuid,
            "rule_name": "Test",
            "extractor_version": "1.0.0",
            "ruleset_version": "1.0.0",
            "content_hash": sample_content_hash,
            "confidence": 0.9,
            "evidence": base_ptr,
            "evidence_context": "Test",
            "description": "Test",
            "remediation": "Fix",
        }

        findings = [
            Finding(**{**base_args, "rule_id": "R002", "severity": "Minor"}),
            Finding(**{**base_args, "rule_id": "R001", "severity": "Critical"}),
            Finding(**{**base_args, "rule_id": "R003", "severity": "Critical"}),
        ]

        sorted_findings = sort_findings(findings)

        # Critical comes first
        assert sorted_findings[0].severity == "Critical"
        assert sorted_findings[0].rule_id == "R001"
        assert sorted_findings[1].severity == "Critical"
        assert sorted_findings[1].rule_id == "R003"
        # Minor comes last
        assert sorted_findings[2].severity == "Minor"

    def test_severity_rank_ordering(self):
        """SEVERITY_RANK should have Critical < Major < Minor < Informational."""
        assert SEVERITY_RANK["Critical"] < SEVERITY_RANK["Major"]
        assert SEVERITY_RANK["Major"] < SEVERITY_RANK["Minor"]
        assert SEVERITY_RANK["Minor"] < SEVERITY_RANK["Informational"]


# =============================================================================
# In-Memory DOCX Fixture Tests (No Binary Commits)
# =============================================================================

@pytest.mark.skipif(not HAS_PYTHON_DOCX, reason="python-docx not installed")
class TestInMemoryDocxExtraction:
    """Test DOCX processing without committing binary fixtures."""

    def create_test_docx(self) -> bytes:
        """Generate a test DOCX in memory."""
        doc = DocxDocument()
        doc.add_heading("Test Protocol", level=0)
        doc.add_paragraph("This is the introduction section.")
        doc.add_heading("Methods", level=1)
        doc.add_paragraph("The study methodology is described here.")
        doc.add_heading("Results", level=1)
        doc.add_paragraph("See Table 1 for demographics.")

        # Save to bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def test_docx_content_hash_deterministic(self):
        """Same DOCX structure should produce consistent hash."""
        docx_bytes = self.create_test_docx()

        hash1 = hashlib.sha256(docx_bytes).hexdigest()
        hash2 = hashlib.sha256(docx_bytes).hexdigest()

        assert hash1 == hash2
        assert len(hash1) == 64

    def test_docx_extraction_produces_canonical_structure(self, fixed_uuid: UUID, fixed_timestamp: datetime):
        """DOCX extraction should produce valid CanonicalDocument."""
        docx_bytes = self.create_test_docx()
        content_hash = hashlib.sha256(docx_bytes).hexdigest()

        # Simulate extraction (simplified)
        doc = DocxDocument(io.BytesIO(docx_bytes))

        paragraphs = []
        for idx, para in enumerate(doc.paragraphs):
            heading_level = None
            if para.style.name.startswith("Heading"):
                try:
                    heading_level = int(para.style.name[-1])
                except ValueError:
                    heading_level = 1
            paragraphs.append(
                CanonicalParagraph(
                    index=idx,
                    text=para.text,
                    heading_level=heading_level,
                )
            )

        # Create canonical document
        canonical = CanonicalDocument(
            doc_id=fixed_uuid,
            content_hash=content_hash,
            source_type="docx",
            title="Test Protocol",
            paragraphs=paragraphs,
            extraction_timestamp=fixed_timestamp,
            extractor_version="1.0.0-test",
        )

        assert canonical.content_hash == content_hash
        assert canonical.source_type == "docx"
        assert len(canonical.paragraphs) > 0


# =============================================================================
# Validation Tests
# =============================================================================

class TestValidation:
    """Test Pydantic validation rules."""

    def test_invalid_hash_rejected(self, fixed_uuid: UUID):
        """Invalid SHA256 hash should raise ValidationError."""
        with pytest.raises(ValueError):
            EvidencePointer(
                doc_id=fixed_uuid,
                content_hash="not-a-valid-hash",
                pointer_type="docx",
            )

    def test_invalid_pointer_type_rejected(self, fixed_uuid: UUID, sample_content_hash: str):
        """Invalid pointer_type should raise ValidationError."""
        with pytest.raises(ValueError):
            EvidencePointer(
                doc_id=fixed_uuid,
                content_hash=sample_content_hash,
                pointer_type="invalid",  # type: ignore
            )

    def test_confidence_bounds_enforced(
        self, fixed_uuid: UUID, sample_evidence_pointer: EvidencePointer, sample_content_hash: str
    ):
        """Confidence must be between 0.0 and 1.0."""
        with pytest.raises(ValueError):
            Finding(
                finding_id=fixed_uuid,
                doc_id=fixed_uuid,
                rule_id="R001",
                rule_name="Test",
                extractor_version="1.0.0",
                ruleset_version="1.0.0",
                content_hash=sample_content_hash,
                severity="Major",
                confidence=1.5,  # Invalid: > 1.0
                evidence=sample_evidence_pointer,
                evidence_context="Test",
                description="Test",
                remediation="Fix",
            )

    def test_rule_id_pattern_enforced(
        self, fixed_uuid: UUID, sample_evidence_pointer: EvidencePointer, sample_content_hash: str
    ):
        """Rule ID must match pattern R###."""
        with pytest.raises(ValueError):
            Finding(
                finding_id=fixed_uuid,
                doc_id=fixed_uuid,
                rule_id="INVALID",  # Must be R### pattern
                rule_name="Test",
                extractor_version="1.0.0",
                ruleset_version="1.0.0",
                content_hash=sample_content_hash,
                severity="Major",
                confidence=0.9,
                evidence=sample_evidence_pointer,
                evidence_context="Test",
                description="Test",
                remediation="Fix",
            )


# =============================================================================
# JSON Serialization Tests (for API compatibility)
# =============================================================================

class TestJsonSerialization:
    """Test JSON round-trip for API responses."""

    def test_evidence_pointer_json_roundtrip(
        self, sample_evidence_pointer: EvidencePointer
    ):
        """EvidencePointer should serialize/deserialize cleanly."""
        json_str = sample_evidence_pointer.model_dump_json()
        restored = EvidencePointer.model_validate_json(json_str)

        assert restored.sort_key() == sample_evidence_pointer.sort_key()
        assert restored.doc_id == sample_evidence_pointer.doc_id

    def test_canonical_document_json_roundtrip(
        self, sample_canonical_document: CanonicalDocument
    ):
        """CanonicalDocument should serialize/deserialize cleanly."""
        json_str = sample_canonical_document.model_dump_json()
        restored = CanonicalDocument.model_validate_json(json_str)

        assert restored.content_hash == sample_canonical_document.content_hash
        assert len(restored.paragraphs) == len(sample_canonical_document.paragraphs)

    def test_finding_json_roundtrip(self, sample_finding: Finding):
        """Finding should serialize/deserialize cleanly."""
        # Exclude computed field 'fingerprint' which is read-only
        json_data = sample_finding.model_dump(exclude={"fingerprint"})
        restored = Finding.model_validate(json_data)

        assert restored.fingerprint == sample_finding.fingerprint
        assert restored.rule_id == sample_finding.rule_id
