"""
Text extraction helpers for DOCX/PDF bytes.

These functions use optional dependencies and return empty text when
libraries are unavailable to keep the batch pipeline stable.
"""

from __future__ import annotations

import re
import warnings
from io import BytesIO
from typing import Optional


def normalize_text(text: str) -> str:
    """Normalize text for stable hashing across extractors."""
    if not text:
        return ""
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\s+", " ", normalized)
    normalized = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]", "", normalized)
    return normalized.strip()


def extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract text from DOCX bytes using python-docx if available."""
    try:
        import docx  # type: ignore
    except Exception:
        return ""

    try:
        document = docx.Document(BytesIO(data))
        chunks = []

        for paragraph in document.paragraphs:
            if paragraph.text:
                chunks.append(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        chunks.append(cell.text)

        return "\n\n".join(chunks)
    except Exception:
        return ""


def extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF bytes using pypdf/PyPDF2 if available."""
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        try:
            with warnings.catch_warnings():
                warnings.filterwarnings("ignore", category=DeprecationWarning, module="PyPDF2")
                from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            return ""

    try:
        reader = PdfReader(BytesIO(data))
        chunks = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text:
                chunks.append(text)
        return "\n\n".join(chunks)
    except Exception:
        return ""
