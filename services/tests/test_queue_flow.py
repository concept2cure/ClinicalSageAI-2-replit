import json
import os
from fastapi.testclient import TestClient
from services.api import app, store
from services.celery_app import generate_docx_task
from services.job_store import get_store
from pathlib import Path
import pytest

client = TestClient(app)


@pytest.mark.skipif(not (os.getenv('CI') or Path('/var/run/docker.sock').exists()), reason="Integration requires Docker or CI")
def test_full_flow(tmp_path):
    # Prepare template and data
    tpl = tmp_path / "tpl.docx"
    from docx import Document
    d = Document()
    d.sections[0].header.paragraphs[0].text = "HDR"
    d.sections[0].footer.paragraphs[0].text = "FTR"
    d.add_paragraph('This is a variable: {{var}}')
    d.add_paragraph('{{efficacy_table}}')
    d.save(str(tpl))

    data = {
        "metadata": {"var": "Real"},
        "tables": [{"name": "efficacy_table", "headers": ["A","B"], "rows": [["x","y"]]}]
    }

    resp = client.post('/api/ectd/generate', json={'data': data, 'template_path': str(tpl)})
    assert resp.status_code == 202
    job_id = resp.json()['job_id']

    # Verify PENDING state
    job = store.get(job_id)
    assert job and job.get('status') == 'PENDING'

    # Run worker synchronously by invoking the task directly
    data_path = job.get('input')
    res = generate_docx_task.run(job_id, data_path, str(tpl))
    assert res['status'] == 'COMPLETED'
    out = res['output']
    assert Path(out).exists()

    # Check final store status
    final = store.get(job_id)
    assert final.get('status') == 'COMPLETED'
    assert final.get('output') == out
