import json
import os
from docx import Document
from services.ectd_generator import generate_docx_from_json


def test_table_headers_render(tmp_path):
    """Test A: Input JSON with 0 rows -> headers render and table has only header row."""
    inp = tmp_path / "input_a.json"
    # Allow CI to collect outputs by setting TEST_OUTPUT_DIR env var
    out_dir = os.getenv("TEST_OUTPUT_DIR")
    if out_dir:
        out = Path(out_dir) / "out_a.docx"
        Path(out_dir).mkdir(parents=True, exist_ok=True)
    else:
        out = tmp_path / "out_a.docx"

    data = {"tables": [{"title": "Table 1", "headers": ["Col A", "Col B"], "rows": []}]}
    inp.write_text(json.dumps(data), encoding="utf-8")

    generate_docx_from_json(str(inp), str(out))

    assert out.exists()
    assert out.stat().st_size > 0

    doc = Document(str(out))
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    # only header row
    assert len(tbl.rows) == 1
    hdr = [c.text.strip() for c in tbl.rows[0].cells]
    assert hdr == ["Col A", "Col B"]


def test_special_characters_and_ampersand(tmp_path):
    """Test B: Special characters like '<' and '&' should not corrupt output and round-trip as text."""
    inp = tmp_path / "input_b.json"
    out_dir = os.getenv("TEST_OUTPUT_DIR")
    if out_dir:
        out = Path(out_dir) / "out_b.docx"
        Path(out_dir).mkdir(parents=True, exist_ok=True)
    else:
        out = tmp_path / "out_b.docx"

    rows = [["Subject 001 & 002", "p < 0.001"]]
    data = {"tables": [{"title": "Special", "headers": ["Subject", "p-value"], "rows": rows}]}
    inp.write_text(json.dumps(data), encoding="utf-8")

    generate_docx_from_json(str(inp), str(out))

    assert out.exists() and out.stat().st_size > 0
    doc = Document(str(out))
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    # header + one row
    assert len(tbl.rows) == 2
    row_cells = [c.text for c in tbl.rows[1].cells]
    assert row_cells[0].strip() == "Subject 001 & 002"
    assert row_cells[1].strip() == "p < 0.001"


def test_missing_values_are_na(tmp_path):
    """Test C: Missing data in a row (shorter lists) should be rendered as 'N/A'."""
    inp = tmp_path / "input_c.json"
    out_dir = os.getenv("TEST_OUTPUT_DIR")
    if out_dir:
        out = Path(out_dir) / "out_c.docx"
        Path(out_dir).mkdir(parents=True, exist_ok=True)
    else:
        out = tmp_path / "out_c.docx"

    # header has three columns, but row provides only one
    data = {"tables": [{"title": "Missing", "headers": ["A", "B", "C"], "rows": [["only one"]]}]}
    inp.write_text(json.dumps(data), encoding="utf-8")

    generate_docx_from_json(str(inp), str(out))

    assert out.exists() and out.stat().st_size > 0
    doc = Document(str(out))
    tbl = doc.tables[0]
    assert len(tbl.rows) == 2
    cells = [c.text.strip() for c in tbl.rows[1].cells]
    assert cells == ["only one", "N/A", "N/A"]


def test_template_injection_and_header_footer(tmp_path):
    """Template test: placeholders replaced and header/footer preserved."""
    # Create mock template
    tpl = tmp_path / "mock_template.docx"
    doc = Document()
    # Add header and footer
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "TEMPLATE HEADER"
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "TEMPLATE FOOTER"

    # Body placeholder for var and table
    doc.add_paragraph("This is a variable: {{var}}")
    doc.add_paragraph("{{efficacy_table}}")
    doc.save(str(tpl))

    inp = tmp_path / "input_tpl.json"
    out = tmp_path / "out_tpl.docx"

    data = {
        "metadata": {"var": "Real Value"},
        "tables": [
            {"name": "efficacy_table", "headers": ["Endpoint", "T1", "T2"], "rows": [["E1", "1", "2"]]}
        ]
    }
    inp.write_text(json.dumps(data), encoding="utf-8")

    generate_docx_from_json(str(inp), str(out), template_path=str(tpl))

    assert out.exists() and out.stat().st_size > 0
    gen = Document(str(out))

    # Header/footer preserved
    sec = gen.sections[0]
    assert any("TEMPLATE HEADER" in p.text for p in sec.header.paragraphs)
    assert any("TEMPLATE FOOTER" in p.text for p in sec.footer.paragraphs)

    # Variable replaced
    assert any("Real Value" in p.text for p in gen.paragraphs)

    # Table injected (should be present)
    assert len(gen.tables) >= 1
    found = False
    for tbl in gen.tables:
        hdr = [c.text.strip() for c in tbl.rows[0].cells]
        if hdr == ["Endpoint", "T1", "T2"]:
            found = True
            assert len(tbl.rows) == 2
            assert [c.text.strip() for c in tbl.rows[1].cells] == ["E1", "1", "2"]
    assert found
