"""
IND Automation Templates Module

This module provides functions for generating FDA IND application forms and related documents
using python-docx instead of docxtpl.
"""

import os
import io
from typing import Dict, Any
from io import BytesIO
from datetime import datetime

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _create_module_document(title: str, module_code: str, data: Dict[str, Any]) -> docx.Document:
    doc = docx.Document()
    heading = doc.add_heading(f"IND Module {module_code}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(title)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"Sponsor: {data.get('sponsor_name', 'Not Provided')}")
    doc.add_paragraph(f"Drug Name: {data.get('drug_name', 'Not Provided')}")
    doc.add_paragraph(f"IND Number: {data.get('ind_number', 'Pending')}")
    return doc


def _save_document(doc: docx.Document) -> BytesIO:
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _add_section(doc: docx.Document, title: str, content: Any) -> None:
    doc.add_heading(title, 2)
    if isinstance(content, list):
        for item in content:
            doc.add_paragraph(str(item), style='List Bullet')
        return

    if isinstance(content, dict):
        for key, value in content.items():
            doc.add_paragraph(f"{key}: {value}")
        return

    doc.add_paragraph(str(content or "Not Provided"))

def _get_template_path(template_name: str) -> str:
    """Get the absolute path to a template file"""
    base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, "templates", "forms", template_name)

def render_form1571(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 1571 (Investigational New Drug Application)

    Args:
        data: Dictionary containing form data

    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()

    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_title = doc.add_heading("INVESTIGATIONAL NEW DRUG APPLICATION (IND)", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_number = doc.add_paragraph("FORM FDA 1571")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add form sections
    doc.add_heading("1. NAME OF SPONSOR", 2)
    doc.add_paragraph(data.get("sponsor_name", ""))

    doc.add_heading("2. DATE OF SUBMISSION", 2)
    doc.add_paragraph(data.get("submission_date", datetime.now().strftime("%Y-%m-%d")))

    doc.add_heading("3. NAME OF SPONSOR CONTACT", 2)
    doc.add_paragraph(data.get("authorizer_name", ""))

    doc.add_heading("4. TELEPHONE NUMBER", 2)
    doc.add_paragraph(data.get("sponsor_phone", ""))

    doc.add_heading("5. NAME OF DRUG", 2)
    doc.add_paragraph(data.get("drug_name", ""))

    doc.add_heading("6. IND NUMBER (if previously assigned)", 2)
    doc.add_paragraph(data.get("ind_number", ""))

    doc.add_heading("7. INDICATION", 2)
    doc.add_paragraph(data.get("indication", ""))

    doc.add_heading("8. PHASE OF CLINICAL INVESTIGATION", 2)
    doc.add_paragraph(data.get("phase", ""))

    doc.add_heading("9. PROTOCOL NUMBER AND TITLE", 2)
    doc.add_paragraph(f"Protocol Number: {data.get('protocol_number', '')}")
    doc.add_paragraph(f"Title: {data.get('protocol_title', '')}")

    doc.add_heading("10. SERIAL NUMBER", 2)
    doc.add_paragraph(data.get("serial_number", "001"))

    # Add authorization section
    doc.add_heading("AUTHORIZATION", 2)
    auth_text = f"I agree to update this statement in accordance with 21 CFR §312.23. I agree to comply with all other requirements regarding the obligations of clinical investigators and all other pertinent requirements in 21 CFR Part 312."
    doc.add_paragraph(auth_text)

    signature = doc.add_paragraph("\n\n___________________________________")
    signature.add_run(f"\n{data.get('authorizer_name', '')}, {data.get('authorizer_title', '')}")
    signature.add_run(f"\nDate: {data.get('submission_date', datetime.now().strftime('%Y-%m-%d'))}")

    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def render_form1572(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 1572 (Statement of Investigator)

    Args:
        data: Dictionary containing form data

    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()

    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_title = doc.add_heading("STATEMENT OF INVESTIGATOR", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_number = doc.add_paragraph("FORM FDA 1572")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add form sections
    doc.add_heading("1. NAME AND ADDRESS OF INVESTIGATOR", 2)
    doc.add_paragraph(data.get("principal_investigator_name", ""))
    doc.add_paragraph(data.get("investigator_address", ""))

    doc.add_heading("2. EDUCATION, TRAINING, AND EXPERIENCE", 2)
    doc.add_paragraph("See attached curriculum vitae")

    doc.add_heading("3. NAME AND ADDRESS OF ANY MEDICAL SCHOOL, HOSPITAL OR OTHER RESEARCH FACILITY WHERE THE CLINICAL INVESTIGATION(S) WILL BE CONDUCTED", 2)
    doc.add_paragraph(data.get("research_facility_name", ""))
    doc.add_paragraph(data.get("research_facility_address", ""))

    doc.add_heading("4. NAME AND ADDRESS OF ANY CLINICAL LABORATORY FACILITIES TO BE USED IN THE STUDY", 2)
    doc.add_paragraph(data.get("clinical_lab_name", ""))
    doc.add_paragraph(data.get("clinical_lab_address", ""))

    doc.add_heading("5. NAME AND ADDRESS OF INSTITUTIONAL REVIEW BOARD (IRB) RESPONSIBLE FOR REVIEW AND APPROVAL OF THE STUDY(IES)", 2)
    doc.add_paragraph(data.get("irb_name", ""))
    doc.add_paragraph(data.get("irb_address", ""))

    doc.add_heading("6. NAMES OF SUBINVESTIGATORS WHO WILL BE ASSISTING IN THE CONDUCT OF THE INVESTIGATION(S)", 2)
    doc.add_paragraph(data.get("subinvestigators", ""))

    doc.add_heading("7. NAME AND CODE NUMBER, IF ANY, OF THE PROTOCOL(S) IN THE IND FOR THE STUDY(IES) TO BE CONDUCTED BY THE INVESTIGATOR", 2)
    doc.add_paragraph(f"Protocol Number: {data.get('protocol_number', '')}")
    doc.add_paragraph(f"Title: {data.get('protocol_title', '')}")

    doc.add_heading("8. STUDY INFORMATION", 2)
    doc.add_paragraph(f"Drug: {data.get('drug_name', '')}")
    doc.add_paragraph(f"Indication: {data.get('indication', '')}")
    doc.add_paragraph(f"Phase: {data.get('phase', '')}")

    # Add commitment statement
    doc.add_heading("COMMITMENT", 2)
    commitment_text = "I agree to conduct the study(ies) in accordance with the relevant, current protocol(s) and will only make changes in a protocol after notifying the sponsor, except when necessary to protect the safety, rights, or welfare of subjects. I agree to personally conduct or supervise the described investigation(s). I agree to inform any patients, or any persons used as controls, that the drugs are being used for investigational purposes and I will ensure that the requirements relating to obtaining informed consent in 21 CFR Part 50 and institutional review board (IRB) review and approval in 21 CFR Part 56 are met."
    doc.add_paragraph(commitment_text)

    # Add signature block
    signature = doc.add_paragraph("\n\n___________________________________")
    signature.add_run(f"\n{data.get('principal_investigator_name', '')}")
    signature.add_run(f"\nDate: {data.get('submission_date', datetime.now().strftime('%Y-%m-%d'))}")

    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def render_form3674(data: Dict[str, Any]) -> BytesIO:
    """
    Generate FDA Form 3674 (Certification of Compliance with ClinicalTrials.gov)

    Args:
        data: Dictionary containing form data

    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()

    # Add title
    title = doc.add_heading("DEPARTMENT OF HEALTH AND HUMAN SERVICES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("FOOD AND DRUG ADMINISTRATION", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_title = doc.add_heading("CERTIFICATION OF COMPLIANCE WITH CLINICALTRIALS.GOV REQUIREMENTS", 1)
    form_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    form_number = doc.add_paragraph("FORM FDA 3674")
    form_number.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add form sections
    doc.add_heading("1. DRUG/PRODUCT INFORMATION", 2)
    doc.add_paragraph(f"Drug Name: {data.get('drug_name', '')}")
    doc.add_paragraph(f"Indication: {data.get('indication', '')}")

    doc.add_heading("2. PROTOCOL INFORMATION", 2)
    doc.add_paragraph(f"Protocol Number: {data.get('protocol_number', '')}")
    doc.add_paragraph(f"Title: {data.get('protocol_title', '')}")
    doc.add_paragraph(f"Phase: {data.get('phase', '')}")

    doc.add_heading("3. APPLICABLE CLINICAL TRIAL INFORMATION", 2)
    doc.add_paragraph(f"NCT Number: {data.get('nct_number', '')}")

    doc.add_heading("4. CERTIFICATION", 2)
    certification_text = "I certify that the requirements of 42 U.S.C. § 282(j), Section 402(j) of the Public Health Service Act, enacted by Title VIII of Public Law 110-85, have been met for the applicable clinical trial identified in Section 3 above. I further certify that appropriate subject consent disclosures regarding trial registration, results information submission, and other trial information disclosure requirements have been provided to each trial subject in accordance with FDA regulations."
    doc.add_paragraph(certification_text)

    # Signature section
    doc.add_heading("5. CERTIFIER INFORMATION", 2)
    doc.add_paragraph(f"Name: {data.get('certifier_name', '')}")
    doc.add_paragraph(f"Title: {data.get('certifier_title', '')}")
    doc.add_paragraph(f"Address: {data.get('certifier_address', '')}")
    doc.add_paragraph(f"Email: {data.get('certifier_email', '')}")
    doc.add_paragraph(f"Phone: {data.get('certifier_phone', '')}")
    if data.get('certifier_fax'):
        doc.add_paragraph(f"Fax: {data.get('certifier_fax', '')}")

    # Add signature block
    signature = doc.add_paragraph("\n\n___________________________________")
    signature.add_run(f"\n{data.get('certifier_name', '')}")
    signature.add_run(f"\nDate: {data.get('submission_date', datetime.now().strftime('%Y-%m-%d'))}")

    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generate_cover_letter(data: Dict[str, Any]) -> BytesIO:
    """
    Generate a cover letter for IND submission

    Args:
        data: Dictionary containing cover letter data

    Returns:
        BytesIO object containing the generated document
    """
    # Create a new document
    doc = docx.Document()

    # Add header content
    doc.add_paragraph(f"Date: {data.get('submission_date', datetime.now().strftime('%Y-%m-%d'))}")
    doc.add_paragraph("To: Food and Drug Administration\nCenter for Drug Evaluation and Research\nCentral Document Room\n5901-B Ammendale Road\nBeltsville, MD 20705-1266")
    doc.add_paragraph(f"From: {data.get('sponsor_name', '')}\n{data.get('sponsor_address', '')}")

    # Add title
    title = doc.add_heading("INVESTIGATIONAL NEW DRUG APPLICATION COVER LETTER", 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add main content
    doc.add_heading("RE: Initial Investigational New Drug Application (IND)", 2)
    doc.add_paragraph(f"Drug Name: {data.get('drug_name', '')}")
    doc.add_paragraph(f"Indication: {data.get('indication', '')}")
    doc.add_paragraph(f"Phase: {data.get('phase', '')}")

    # Add letter body
    doc.add_paragraph("Dear FDA Review Staff,")

    body_text = f"""
    On behalf of {data.get('sponsor_name', '')}, I am submitting this Investigational New Drug (IND) application to conduct clinical investigations with {data.get('drug_name', '')} for the treatment of {data.get('indication', '')}.

    The enclosed IND includes the following:
    • Form FDA 1571 (Investigational New Drug Application)
    • Form FDA 1572 (Statement of Investigator)
    • Form FDA 3674 (Certification of Compliance with ClinicalTrials.gov Requirements)
    • Protocol: {data.get('protocol_title', '')} (Protocol No. {data.get('protocol_number', '')})
    • Investigator's Brochure
    • Chemistry, Manufacturing, and Controls (CMC) Information
    • Pharmacology and Toxicology Information
    • Previous Human Experience

    This clinical investigation is a {data.get('phase', '')} study in subjects with {data.get('indication', '')}. The protocol has been reviewed and approved by the Institutional Review Board at {data.get('research_facility_name', '')}.

    We look forward to working with the FDA on the development of {data.get('drug_name', '')}. If you have any questions or require additional information, please do not hesitate to contact:

    {data.get('contact_name', '')}
    {data.get('contact_email', '')}
    {data.get('contact_phone', '')}
    """

    for paragraph in body_text.strip().split('\n'):
        doc.add_paragraph(paragraph.strip())

    # Add closing
    doc.add_paragraph("\nSincerely,")

    # Add signature block
    doc.add_paragraph("\n\n___________________________________")
    doc.add_paragraph(f"{data.get('authorizer_name', '')}")
    doc.add_paragraph(f"{data.get('authorizer_title', '')}")
    doc.add_paragraph(f"{data.get('sponsor_name', '')}")

    # Save the document to a BytesIO object
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def render_module1(data: Dict[str, Any]) -> BytesIO:
    doc = _create_module_document(
        "Administrative and Prescribing Information",
        "1",
        data,
    )

    _add_section(
        doc,
        "Submission Components",
        [
            "Form FDA 1571",
            "Form FDA 1572",
            "Form FDA 3674",
            "Cover Letter",
            "Investigator's Brochure Reference",
        ],
    )
    _add_section(doc, "Sponsor Contact", {
        "Name": data.get("contact_name", ""),
        "Email": data.get("contact_email", ""),
        "Phone": data.get("contact_phone", ""),
    })
    _add_section(doc, "Regulatory Notes", data.get("module1_notes", "Initial IND administrative package."))
    return _save_document(doc)


def render_module2(data: Dict[str, Any]) -> BytesIO:
    doc = _create_module_document("CTD Summaries", "2", data)
    _add_section(doc, "2.3 Quality Overall Summary", data.get("module2_quality_summary", data.get("quality_summary", "Pending quality summary.")))
    _add_section(doc, "2.4 Nonclinical Overview", data.get("module2_nonclinical_overview", data.get("nonclinical_overview", "Pending nonclinical overview.")))
    _add_section(doc, "2.5 Clinical Overview", data.get("module2_clinical_overview", data.get("clinical_overview", "Pending clinical overview.")))
    return _save_document(doc)


def render_module3(data: Dict[str, Any]) -> BytesIO:
    doc = _create_module_document("Quality (CMC)", "3", data)
    _add_section(doc, "Drug Substance", data.get("drug_substance", "Not Provided"))
    _add_section(doc, "Drug Product", data.get("drug_product", "Not Provided"))
    _add_section(doc, "Manufacturing Site", data.get("manufacturing_site", "Not Provided"))
    _add_section(doc, "Batch Number", data.get("batch_number", "Not Provided"))
    _add_section(doc, "Specifications", data.get("specifications", []))
    _add_section(doc, "Stability Data", data.get("stability_data", []))
    return _save_document(doc)


def render_module4(data: Dict[str, Any]) -> BytesIO:
    doc = _create_module_document("Nonclinical Study Reports", "4", data)
    _add_section(doc, "Pharmacology", data.get("module4_pharmacology", data.get("pharmacology_summary", "Pending pharmacology summary.")))
    _add_section(doc, "Pharmacokinetics", data.get("module4_pk", data.get("pk_summary", "Pending pharmacokinetics summary.")))
    _add_section(doc, "Toxicology", data.get("module4_toxicology", data.get("toxicology_summary", "Pending toxicology summary.")))
    _add_section(doc, "Nonclinical Conclusions", data.get("module4_conclusions", "Risk assessment and no-observed-adverse-effect level to be finalized."))
    return _save_document(doc)


def render_module5(data: Dict[str, Any]) -> BytesIO:
    doc = _create_module_document("Clinical Study Reports", "5", data)
    _add_section(doc, "Protocol Synopsis", data.get("protocol_title", data.get("module5_protocol_synopsis", "Pending protocol synopsis.")))
    _add_section(doc, "Study Design", data.get("module5_study_design", "Randomized, controlled, multi-center study design details pending."))
    _add_section(doc, "Endpoints", data.get("module5_endpoints", data.get("primary_endpoint", "Pending endpoint definition.")))
    _add_section(doc, "Safety Summary", data.get("module5_safety", "Safety profile to be finalized from available data."))
    _add_section(doc, "Efficacy Summary", data.get("module5_efficacy", "Efficacy narrative to be completed once data is locked."))
    return _save_document(doc)


def render_ind_module(module_number: int, data: Dict[str, Any]) -> BytesIO:
    renderers = {
        1: render_module1,
        2: render_module2,
        3: render_module3,
        4: render_module4,
        5: render_module5,
    }

    renderer = renderers.get(module_number)
    if renderer is None:
        raise ValueError("Module must be between 1 and 5")

    return renderer(data)
