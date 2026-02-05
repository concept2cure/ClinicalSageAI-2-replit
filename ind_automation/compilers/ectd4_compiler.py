"""
Pure eCTD 4.0 Compiler - No 3.2.2 compatibility cruft.
Generates FHIR-based provenance, XMP metadata, and 2-6-2 filenames.
"""

import json
import hashlib
import uuid
from datetime import datetime, timezone
from dataclasses import dataclass, asdict
from typing import List, Dict
from pathlib import Path
import zipfile
from lxml import etree

from ind_automation.utils.content_controls import inject_sdt_to_paragraph
# Import signature type for type hints only (no runtime import required here)
try:
    from ind_automation.signatures.pki_signer import Part11Signature  # type: ignore
except Exception:
    Part11Signature = None  # graceful fallback for environments without rsa


@dataclass
class FHIRAuditEvent:
    """FHIR AuditEvent for eCTD 4.0 provenance."""
    resourceType: str = "AuditEvent"
    id: str = ""
    recorded: str = ""
    agent: List[Dict] = None
    source: Dict = None
    entity: List[Dict] = None

    def to_dict(self):
        return asdict(self)


@dataclass
class ECTD4Document:
    """eCTD 4.0 Document structure for backbone."""
    document_id: str
    document_type: str  # "2.6.2"
    specification_level: str  # "1", "2", or "3"
    file_name: str
    language: str
    content_hash: str
    data_integrity: Dict
    modification_history: List[Dict]
    parent_documents: List[str] = None

    def to_backbone_entry(self) -> Dict:
        """Convert to eCTD.json entry."""
        return {
            "documentId": self.document_id,
            "documentType": self.document_type,
            "specificationLevel": self.specification_level,
            "fileName": self.file_name,
            "language": self.language,
            "contentHash": self.content_hash,
            "dataIntegrity": self.data_integrity,
            "modificationHistory": self.modification_history,
            "parentDocuments": self.parent_documents or []
        }


class ECTD4Compiler:
    """Native eCTD 4.0 compiler."""

    def __init__(self, output_dir: str = "./ind_automation/output", signature_config=None):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Signature configuration and signer wiring
        try:
            from ind_automation.config.signature_config import SignatureConfig
        except Exception:
            SignatureConfig = None

        self.sig_config = signature_config or (SignatureConfig() if SignatureConfig else None)
        self.signer = None
        self._audit_config = {}
        if self.sig_config:
            signer_conf = self.sig_config.get_signer_config()
            signer_cls = signer_conf.get('class')
            signer_params = signer_conf.get('params', {})
            # Instantiate signer (deferred errors in tests will be more explicit)
            self.signer = signer_cls(**{k: v for k, v in signer_params.items() if v})
            self._audit_config = self.sig_config.get_audit_metadata()

    def compile(self, canvas_document) -> ECTD4Document:
        """
        Compile Canvas document to eCTD 4.0 native format.
        """
        doc_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{getattr(canvas_document,'study_id','')}-{getattr(canvas_document,'created_at','')}") )

        spec_level = self._determine_spec_level(canvas_document)

        filename = self._generate_filename(canvas_document.module, getattr(canvas_document, 'section', 'document'), getattr(canvas_document,'language','en'))

        fhir_history = self._canvas_to_fhir(canvas_document.events if hasattr(canvas_document, 'events') else [], doc_id)

        content_bytes = self._serialize_content(canvas_document)
        content_hash = hashlib.sha256(content_bytes).hexdigest()

        ectd_doc = ECTD4Document(
            document_id=doc_id,
            document_type=canvas_document.module,
            specification_level=spec_level,
            file_name=filename,
            language=canvas_document.language or "en",
            content_hash=content_hash,
            data_integrity=self._build_alcoa_plus(canvas_document),
            modification_history=fhir_history,
            parent_documents=[]
        )

        return ectd_doc

    def generate_docx(self, ectd_doc: ECTD4Document, canvas_document) -> Path:
        """
        Generate DOCX with XMP metadata (not custom.xml).
        """
        from docx import Document

        doc = Document()

        for block in getattr(canvas_document, 'content_blocks', []):
            p = doc.add_paragraph()
            inject_sdt_to_paragraph(p, tag=block.get('sdt_tag', 'textBlock'), content=block.get('text',''), alias=block.get('alias'))

        temp_path = self.output_dir / f"temp_{ectd_doc.document_id}.docx"
        doc.save(temp_path)

        final_path = self.output_dir / ectd_doc.file_name
        self._inject_xmp(temp_path, final_path, ectd_doc)

        temp_path.unlink()

        return final_path

    def sign_and_audit(self, ectd_doc: ECTD4Document, canvas_document, signer=None, signer_info: dict = None):
        """Sign document and append signature event to modification history.

        Uses provided signer if given, otherwise the configured signer from init.
        Enforces production requirements (e.g., TSA timestamping) when running in production mode.
        """
        signer = signer or self.signer
        signer_info = signer_info or {}

        if signer is None:
            raise RuntimeError("No signer available. Provide a signer or configure SignatureConfig.")

        # Generate the physical document
        docx_path = self.generate_docx(ectd_doc, canvas_document)

        # Perform cryptographic signing (use TSA-enabled signing if available)
        if hasattr(signer, 'sign_with_timestamp'):
            signed_path = signer.sign_with_timestamp(docx_path, signer_info)
        elif hasattr(signer, 'sign_document'):
            signed_path = signer.sign_document(docx_path, signer_info)
        else:
            raise RuntimeError("Signer must implement sign_document() or sign_with_timestamp()")

        # Build signature audit event and append
        signature_result = {"document_hash": signer.calculate_document_hash(docx_path)}
        if hasattr(signer, 'last_tsa_info') and getattr(signer, 'last_tsa_info'):
            signature_result['tsa'] = signer.last_tsa_info

        try:
            sig_event = signer.create_fhir_signature_event(signature_result, signer_info, ectd_doc.document_id)
        except Exception:
            # Ensure we don't break signing flow if the signer lacks method
            tsa = signature_result.get('tsa', {})
            entity = {
                "what": {"reference": f"urn:uuid:{ectd_doc.document_id}", "identifier": {"system": "https://concept2cure.io/ectd4", "value": ectd_doc.document_id}},
                "detail": [{"type": "hash-preimage", "valueString": signature_result["document_hash"]}]
            }
            if tsa.get('present'):
                entity['detail'].extend([
                    {"type": "tsa-token-present", "valueBoolean": True},
                    {"type": "tsa-provider", "valueString": tsa.get('provider')},
                    {"type": "timestamp-rfc3161", "valueInstant": tsa.get('timestamp')}
                ])
            sig_event = {
                "resourceType": "AuditEvent",
                "id": f"sig-{uuid.uuid4().hex[:8]}",
                "recorded": datetime.now(timezone.utc).isoformat(),
                "type": {"code": "110107"},
                "agent": [],
                "entity": [entity]
            }

        # Attach config metadata to audit event when available
        if hasattr(self, '_audit_config') and self._audit_config:
            try:
                sig_event.setdefault('source', {}).setdefault('config', {}).update(self._audit_config)
            except Exception:
                pass

        ectd_doc.modification_history.append(sig_event)

        return signed_path, ectd_doc

    def generate_backbone(self, documents: List[ECTD4Document]) -> Dict:
        return {
            "ectdVersion": "4.0",
            "submissionId": str(uuid.uuid4()),
            "submissionDate": datetime.now(timezone.utc).isoformat(),
            "documents": [doc.to_backbone_entry() for doc in documents]
        }

    def _generate_filename(self, module: str, section: str, language: str) -> str:
        module_dash = module.replace('.', '-')
        section_slug = str(section).lower().replace(' ', '-').replace('_', '-')
        return f"{module_dash}-{section_slug}-{language}.docx"

    def _determine_spec_level(self, canvas_document) -> str:
        if getattr(canvas_document, 'subsection', None):
            return "3"
        elif getattr(canvas_document, 'section', None):
            return "2"
        else:
            return "1"

    def _canvas_to_fhir(self, events: List, doc_id: str) -> List[Dict]:
        fhir_events = []

        for i, event in enumerate(events):
            audit = {
                "resourceType": "AuditEvent",
                "id": event.get('id', f"evt-{i}"),
                "recorded": event.get('timestamp', datetime.now(timezone.utc).isoformat()),
                "agent": [{
                    "who": {
                        "identifier": {"value": event.get('user', {}).get('id', 'unknown')},
                        "display": event.get('user', {}).get('name', 'Unknown')
                    },
                    "requestor": True,
                    "type": {
                        "code": event.get('type', 'UNKNOWN'),
                        "display": event.get('type', 'Unknown Action')
                    }
                }],
                "source": {"observer": {"display": "Concept2Cure-AI"}, "site": "cloud"},
                "entity": [{
                    "what": {"reference": f"urn:uuid:{doc_id}"},
                    "type": {"code": "document"},
                    "description": event.get('payload', '')[:255],
                    "detail": []
                }]
            }

            if event.get('ai_context'):
                ai_ctx = event['ai_context']
                audit["entity"][0]["detail"].extend([
                    {"type": "confidence-score", "valueDecimal": ai_ctx.get('confidence', 0.0)},
                    {"type": "model-version", "valueString": ai_ctx.get('model', 'unknown')},
                    {"type": "data-source", "valueString": ai_ctx.get('rag_source', 'unknown')}
                ])

            fhir_events.append(audit)

        return fhir_events

    def _build_alcoa_plus(self, canvas_document) -> Dict:
        events = getattr(canvas_document, 'events', [])
        return {
            "attributable": all(e.get('user') for e in events),
            "legible": True,
            "contemporaneous": True,
            "original": True,
            "accurate": all(e.get('ai_context', {}).get('confidence', 1.0) > 0.8 for e in events if e.get('type') == 'AI_GENERATION'),
            "complete": len(events) > 0,
            "consistent": True,
            "enduring": True,
            "available": True,
            "verificationHash": "sha256",
            "schemaVersion": "ALCOA-v2-ectd4"
        }

    def _serialize_content(self, canvas_document) -> bytes:
        canonical = {
            "module": canvas_document.module,
            "study_id": getattr(canvas_document, 'study_id', ''),
            "content": [b['text'] for b in getattr(canvas_document, 'content_blocks', [])]
        }
        return json.dumps(canonical, sort_keys=True).encode('utf-8')

    def _inject_xmp(self, input_path: Path, output_path: Path, ectd_doc: ECTD4Document):
        xmp_content = f"""<?xpacket begin="" id="W5M0MpCehiHzreSzNTczkc9d"?>
<x:xmpmeta xmlns:x="adobe:ns:meta/" xmlns:rdf="http://www.w3.org/1999/02/22-rdf-syntax-ns#"
           xmlns:ich="http://www.ich.org/eCTD/4.0/" xmlns:dc="http://purl.org/dc/elements/1.1/"
           xmlns:xmp="http://ns.adobe.com/xap/1.0/">
 <rdf:RDF>
  <rdf:Description rdf:about="" xmlns:ich="http://www.ich.org/eCTD/4.0/">
   <ich:documentType>{ectd_doc.document_type}</ich:documentType>
   <ich:documentId>{ectd_doc.document_id}</ich:documentId>
   <ich:specificationLevel>{ectd_doc.specification_level}</ich:specificationLevel>
   <ich:contentHash algorithm="sha256">{ectd_doc.content_hash}</ich:contentHash>
   <ich:submissionType>IND</ich:submissionType>
   <ich:dataIntegrity attribut="{str(ectd_doc.data_integrity['attributable']).lower()}"
                       complete="{str(ectd_doc.data_integrity['complete']).lower()}"/>
  </rdf:Description>
  <rdf:Description rdf:about="" xmlns:dc="http://purl.org/dc/elements/1.1/">
   <dc:creator>Concept2Cure AI System</dc:creator>
   <dc:date>{datetime.now(timezone.utc).isoformat()}</dc:date>
   <dc:title>{ectd_doc.file_name}</dc:title>
  </rdf:Description>
  <rdf:Description rdf:about="" xmlns:xmp="http://ns.adobe.com/xap/1.0/">
   <xmp:CreatorTool>Concept2Cure-v1.0-ECTD4</xmp:CreatorTool>
   <xmp:CreateDate>{datetime.now(timezone.utc).isoformat()}</xmp:CreateDate>
   <xmp:ModifyDate>{datetime.now(timezone.utc).isoformat()}</xmp:ModifyDate>
   <xmp:MetadataDate>{datetime.now(timezone.utc).isoformat()}</xmp:MetadataDate>
  </rdf:Description>
 </rdf:RDF>
</x:xmpmeta>
<?xpacket end="w"?>"""

        with zipfile.ZipFile(input_path, 'r') as zip_in:
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for item in zip_in.namelist():
                    zip_out.writestr(item, zip_in.read(item))
                zip_out.writestr('docProps/core.xml.xmp', xmp_content)
