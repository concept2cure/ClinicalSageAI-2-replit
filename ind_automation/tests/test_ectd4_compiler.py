"""
Pure eCTD 4.0 compliance tests.
Validates: FHIR AuditEvents, XMP metadata, 2-6-2 filenames, JSON backbone.
"""

import pytest
import json
from datetime import datetime
from ind_automation.compilers.ectd4_compiler import ECTD4Compiler


@pytest.fixture
def sample_canvas_doc():
    """Mock Canvas document for testing."""
    class MockCanvas:
        def __init__(self):
            self.study_id = "TEST-001"
            self.module = "2.6.2"
            self.section = "Pharmacokinetics"
            self.subsection = None
            self.language = "en"
            self.created_at = "2026-01-01T00:00:00Z"
            self.content_blocks = [
                {"text": "Cmax was 125.5 ng/mL", "sdt_tag": "pkCmaxValue"}
            ]
            self.events = [
                {
                    "id": "evt-001",
                    "timestamp": "2026-01-01T12:00:00Z",
                    "type": "AI_GENERATION",
                    "user": {"id": "ai_001", "name": "Concept2Cure AI"},
                    "ai_context": {
                        "confidence": 0.94,
                        "model": "gpt-4-turbo",
                        "rag_source": "study_report_001.pdf#page=10"
                    }
                }
            ]
    return MockCanvas()


class TestECTD4Compiler:
    """eCTD 4.0 native compliance."""

    def test_filename_no_m_prefix(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        result = compiler.compile(sample_canvas_doc)

        assert not result.file_name.startswith("m"), f"Legacy 3.2.2 'm' prefix detected: {result.file_name}"
        assert result.file_name.startswith("2-6-2"), f"Wrong format, expected 2-6-2-*: {result.file_name}"
        assert result.file_name.endswith("-en.docx")

    def test_fhir_audit_event_structure(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        result = compiler.compile(sample_canvas_doc)

        assert len(result.modification_history) == 1
        event = result.modification_history[0]

        assert event["resourceType"] == "AuditEvent"
        assert "agent" in event
        assert event["agent"][0]["who"]["identifier"]["value"] == "ai_001"
        assert "entity" in event

        details = event["entity"][0]["detail"]
        confidence = [d for d in details if d["type"] == "confidence-score"]
        assert len(confidence) == 1
        assert confidence[0]["valueDecimal"] == 0.94

    def test_xmp_not_custom_xml(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        ectd_doc = compiler.compile(sample_canvas_doc)
        docx_path = compiler.generate_docx(ectd_doc, sample_canvas_doc)

        import zipfile
        with zipfile.ZipFile(docx_path) as z:
            assert 'docProps/custom.xml' not in z.namelist(), "Legacy custom.xml found (eCTD 3.2.2 artifact)"
            assert 'docProps/core.xml.xmp' in z.namelist(), "Missing XMP metadata (eCTD 4.0 required)"
            xmp = z.read('docProps/core.xml.xmp').decode('utf-8')
            assert '2.6.2' in xmp
            assert 'sha256' in xmp

    def test_json_backbone_not_xml(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        ectd_doc = compiler.compile(sample_canvas_doc)
        backbone = compiler.generate_backbone([ectd_doc])

        assert backbone["ectdVersion"] == "4.0"
        assert "documents" in backbone
        assert isinstance(backbone["documents"], list)
        assert backbone["documents"][0]["documentType"] == "2.6.2"
        assert "modificationHistory" in backbone["documents"][0]

    def test_content_hash_integrity(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        result = compiler.compile(sample_canvas_doc)

        assert len(result.content_hash) == 64
        assert all(c in '0123456789abcdef' for c in result.content_hash)

    def test_alcoa_plus_flags(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        result = compiler.compile(sample_canvas_doc)

        assert result.data_integrity["attributable"] is True
        assert result.data_integrity["accurate"] is True
        assert result.data_integrity["complete"] is True
        assert result.data_integrity["schemaVersion"] == "ALCOA-v2-ectd4"

    def test_specification_level_2(self, sample_canvas_doc, tmp_path):
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        result = compiler.compile(sample_canvas_doc)

        assert result.specification_level == "2"


import zipfile
import pytest
from ind_automation.signatures.pki_signer import Part11Signature


class TestDigitalSignatures:
    """21 CFR Part 11 compliance tests."""
    
    def test_document_signature_creates_xades(self, tmp_path):
        """Signed documents must contain XAdES signature XML."""
        from ind_automation.compilers.ectd4_compiler import ECTD4Compiler
        
        # Use existing sample data pattern from earlier tests
        class MockCanvas:
            def __init__(self):
                self.study_id = "SIG-TEST-001"
                self.module = "2.6.2"
                self.section = "Pharmacokinetics"
                self.subsection = None
                self.language = "en"
                self.created_at = "2026-01-01T00:00:00Z"
                self.content_blocks = [
                    {"text": "Cmax was 125.5 ng/mL", "sdt_tag": "pkCmaxValue"}
                ]
                self.events = []
        
        canvas_doc = MockCanvas()
        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        doc = compiler.compile(canvas_doc)
        docx = compiler.generate_docx(doc, canvas_doc)
        
        signer = Part11Signature()  # Test keys
        signed = signer.sign_document(docx, {
            "name": "Dr. Jane Smith",
            "role": "Medical Monitor",
            "user_id": "mon_001",
            "meaning": "Approved for IND submission"
        })
        
        with zipfile.ZipFile(signed) as z:
            assert '_signatures/signatures.xml' in z.namelist()
            sig_xml = z.read('_signatures/signatures.xml')
            assert b'XAdES' in sig_xml or b'XAdESSignatures' in sig_xml
            assert b'DigestValue' in sig_xml
            assert b'SignatureValue' in sig_xml
    
    def test_signature_tamper_detection(self, tmp_path):
        """Changing document after signing must invalidate signature."""
        signer = Part11Signature()
        
        # Create minimal docx for testing
        from docx import Document
        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(docx_path)
        
        signed = signer.sign_document(docx_path, {"name": "Test", "role": "QA"})
        
        # Verify initially valid
        initial_check = signer.verify_signature(signed)
        assert initial_check["document_hash_match"] is True
        
        # Tamper with document
        tampered_path = tmp_path / "tampered.docx"
        with zipfile.ZipFile(signed, 'r') as zin:
            with zipfile.ZipFile(tampered_path, 'w') as zout:
                for item in zin.namelist():
                    content = zin.read(item)
                    if item.endswith('.xml') and 'document' in item:
                        content = content.replace(b'Test content', b'TAMPERED')
                    zout.writestr(item, content)
        
        tamper_check = signer.verify_signature(tampered_path)
        assert tamper_check["document_hash_match"] is False
        assert tamper_check["valid"] is False

    def test_signature_creates_fhir_audit_event(self, tmp_path, sample_canvas_doc):
        """Signature must create traceable FHIR AuditEvent in backbone."""
        from ind_automation.compilers.ectd4_compiler import ECTD4Compiler

        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        signer = Part11Signature()

        ectd_doc = compiler.compile(sample_canvas_doc)
        signed_path, audited_doc = compiler.sign_and_audit(ectd_doc, sample_canvas_doc, signer, {
            "name": "Dr. Smith",
            "role": "Principal Investigator",
            "user_id": "pi_001"
        })

        sig_events = [e for e in audited_doc.modification_history if isinstance(e.get('type'), dict) and e['type'].get('code') == '110107']
        assert len(sig_events) == 1

        event = sig_events[0]
        details = event["entity"][0]["detail"]
        cert_fp = [d for d in details if d["type"] == "certificate-fingerprint"]
        assert len(cert_fp) == 1
        assert len(cert_fp[0]["valueString"]) == 64  # SHA-256 hex

        hash_pre = [d for d in details if d["type"] == "hash-preimage"]
        assert len(hash_pre) == 1
        assert len(hash_pre[0]["valueString"]) == 64

    def test_signature_audit_references_correct_document(self, tmp_path, sample_canvas_doc):
        """Signature AuditEvent must reference actual document UUID."""
        from ind_automation.compilers.ectd4_compiler import ECTD4Compiler

        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        signer = Part11Signature()

        ectd_doc = compiler.compile(sample_canvas_doc)
        signed_path, audited_doc = compiler.sign_and_audit(ectd_doc, sample_canvas_doc, signer, {
            "name": "Dr. Smith", "role": "PI", "user_id": "pi_001"
        })

        sig_event = [e for e in audited_doc.modification_history if isinstance(e.get('type'), dict) and e['type'].get('code') == '110107'][0]
        assert sig_event["entity"][0]["what"]["reference"] == f"urn:uuid:{audited_doc.document_id}"

    def test_signature_includes_tsa_timestamp(self, tmp_path, monkeypatch):
        """TSA token must be embedded in XAdES and AuditEvent must record TSA info."""
        from ind_automation.compilers.ectd4_compiler import ECTD4Compiler
        from ind_automation.signatures.pki_signer import TSASignature

        # Mock TSA response (return a valid TimeStampResp DER with granted status)
        class FakeResp:
            status_code = 200
            def __init__(self, content):
                self.content = content
            def raise_for_status(self):
                return None
        from pyasn1.codec.der import encoder
        from pyasn1_modules import rfc3161, rfc5652
        def fake_post(url, data=None, headers=None, timeout=None):
            # Construct a minimal TimeStampResp with granted status
            resp = rfc3161.TimeStampResp()
            pki = rfc3161.PKIStatusInfo()
            pki.setComponentByName('status', 0)
            resp.setComponentByName('status', pki)
            # minimal empty ContentInfo
            resp.setComponentByName('timeStampToken', rfc5652.ContentInfo())
            return FakeResp(encoder.encode(resp))
        monkeypatch.setattr('requests.post', fake_post)

        compiler = ECTD4Compiler(output_dir=str(tmp_path))
        signer = TSASignature(tsa_url='http://timestamp.digicert.com')

        # Create a minimal canvas doc
        class Canvas:
            def __init__(self):
                self.study_id = 'TSA-001'
                self.module = '2.6.2'
                self.section = 'Pharmacokinetics'
                self.language = 'en'
                self.content_blocks = [{'text': 'TSA test', 'sdt_tag': 't'}]
                self.events = []
                self.created_at = '2026-01-01T00:00:00Z'
        canvas = Canvas()

        ectd_doc = compiler.compile(canvas)

        # Sign with TSA
        signed_path, audited_doc = compiler.sign_and_audit(ectd_doc, canvas, signer, {
            'name': 'TSA User', 'role': 'QA', 'user_id': 'tsa_01'
        })

        # Inspect signatures.xml for timestamp element
        import zipfile
        with zipfile.ZipFile(signed_path) as z:
            sig_xml = z.read('_signatures/signatures.xml')
            assert b'SignatureTimeStamp' in sig_xml or b'EncapsulatedTimeStamp' in sig_xml

        # Verify AuditEvent includes TSA details
        sig_events = [e for e in audited_doc.modification_history if isinstance(e.get('type'), dict) and e['type'].get('code') == '110107']
        assert len(sig_events) == 1
        event = sig_events[0]
        details = event['entity'][0]['detail']
        assert any(d.get('type') == 'tsa-token-present' and d.get('valueBoolean') for d in details)
        assert any(d.get('type') == 'tsa-provider' for d in details)
        assert any(d.get('type') == 'timestamp-rfc3161' for d in details)

    def test_tsa_request_is_valid_asn1(self):
        """TSA query must be valid RFC3161 ASN.1 TimeStampReq."""
        from ind_automation.signatures.pki_signer import TSASignature
        from pyasn1.codec.der import decoder
        from pyasn1_modules import rfc3161
        from pyasn1.type import univ

        signer = TSASignature()
        test_data = b"test signature xml"
        tsq = signer._create_tsa_query(test_data)

        decoded, _ = decoder.decode(tsq, asn1Spec=rfc3161.TimeStampReq())
        assert isinstance(decoded, univ.Sequence)
        assert int(decoded.getComponentByName('version')) == 1
        msg_imprint = decoded.getComponentByName('messageImprint')
        algo_oid = str(msg_imprint.getComponentByName('hashAlgorithm').getComponentByName('algorithm'))
        assert algo_oid == '2.16.840.1.101.3.4.2.1'

    def test_tsa_response_parsing(self):
        """Must parse a minimal TimeStampResp ASN.1 and extract token."""
        from ind_automation.signatures.pki_signer import TSASignature
        from pyasn1_modules import rfc3161, rfc5652

        signer = TSASignature()
        # Build minimal granted TimeStampResp
        from pyasn1.codec.der import encoder
        resp = rfc3161.TimeStampResp()
        pki = rfc3161.PKIStatusInfo()
        pki.setComponentByName('status', 0)
        resp.setComponentByName('status', pki)
        resp.setComponentByName('timeStampToken', rfc5652.ContentInfo())

        der = encoder.encode(resp)
        result = signer._parse_tsa_response(der)
        assert result['status'] == 'granted'
        assert 'token' in result
        # parsed may be False for minimal/partial ContentInfo fixtures
        assert result.get('parsed') in (True, False)


class TestSignatureConfiguration:
    """Validate FDA Part 11 configuration requirements."""

    def test_production_requires_kms_config(self):
        """Production mode must have KMS configuration."""
        import os
        os.environ['CONCEPT2CURE_SIGNER_MODE'] = 'hsm_kms'
        os.environ.pop('KMS_KEY_ID', None)  # Ensure missing

        with pytest.raises(ValueError, match="KMS_KEY_ID"):
            from ind_automation.config.signature_config import SignatureConfig
            SignatureConfig()

        # Cleanup
        os.environ.pop('CONCEPT2CURE_SIGNER_MODE', None)

    def test_dev_mode_allows_missing_config(self):
        """Development mode works without KMS."""
        import os
        os.environ['CONCEPT2CURE_SIGNER_MODE'] = 'dev'

        from ind_automation.config.signature_config import SignatureConfig, SignerMode
        config = SignatureConfig()
        assert config.mode == SignerMode.DEV
        assert not config.is_production()

        os.environ.pop('CONCEPT2CURE_SIGNER_MODE', None)

    def test_config_provides_audit_metadata(self):
        """Configuration must expose audit metadata for FHIR."""
        import os
        os.environ['CONCEPT2CURE_SIGNER_MODE'] = 'hsm_kms'
        os.environ['KMS_KEY_ID'] = 'arn:aws:kms:us-east-1:123:key/abc'
        os.environ['AWS_REGION'] = 'us-west-2'
        os.environ['KMS_KEY_CUSTODY_SOP'] = 'SOP-SEC-001'

        from ind_automation.config.signature_config import SignatureConfig
        config = SignatureConfig()
        meta = config.get_audit_metadata()

        assert meta['signer_mode'] == 'hsm_kms'
        assert meta['environment'] == 'production'
        assert meta['key_custody'] == 'SOP-SEC-001'
        assert meta['hosted_in'] == 'us-west-2'

        os.environ.pop('CONCEPT2CURE_SIGNER_MODE', None)
        os.environ.pop('KMS_KEY_ID', None)
        os.environ.pop('AWS_REGION', None)
        os.environ.pop('KMS_KEY_CUSTODY_SOP', None)


class TestHSMSignatures:
    """Requires AWS credentials (mocked for CI)."""

    def test_hsm_signature_uses_kms(self, tmp_path, monkeypatch):
        """Verify HSM signing invokes KMS (not local key)."""
        from ind_automation.signatures.hsm_signer import HSMSignature
        from unittest.mock import MagicMock, patch
        from docx import Document

        mock_kms = MagicMock()
        mock_kms.get_public_key.return_value = {'PublicKey': b'fakepub', 'SigningAlgorithms': ['RSASSA_PKCS1_V1_5_SHA_256']}
        mock_kms.sign.return_value = {'Signature': b'mocked_hsm_signature'}

        with patch('boto3.client', return_value=mock_kms):
            signer = HSMSignature(kms_key_id='alias/fda-signing-key')

            # Create test doc
            docx_path = tmp_path / "test.docx"
            doc = Document()
            doc.add_paragraph("Test")
            doc.save(docx_path)

            signed = signer.sign_document(docx_path, {
                "name": "Dr. Smith", "role": "PI"
            })

            # Verify KMS was called
            mock_kms.sign.assert_called_once()
            called = mock_kms.sign.call_args.kwargs
            assert called['KeyId'] == 'alias/fda-signing-key'
            assert called['MessageType'] == 'DIGEST'

    def test_hsm_includes_key_id_in_audit(self, tmp_path, monkeypatch):
        """Signature AuditEvent must reference HSM key for traceability."""
        from ind_automation.signatures.hsm_signer import HSMSignature
        from ind_automation.compilers.ectd4_compiler import ECTD4Compiler
        from unittest.mock import MagicMock, patch

        mock_kms = MagicMock()
        mock_kms.get_public_key.return_value = {'PublicKey': b'fakepub', 'SigningAlgorithms': ['RSASSA_PKCS1_V1_5_SHA_256']}
        mock_kms.sign.return_value = {'Signature': b'mocked_hsm_signature'}

        with patch('boto3.client', return_value=mock_kms):
            signer = HSMSignature(kms_key_id='alias/fda-signing-key')
            compiler = ECTD4Compiler(output_dir=str(tmp_path))

            class Canvas:
                def __init__(self):
                    self.study_id = 'HSM-001'
                    self.module = '2.6.2'
                    self.section = 'Pharmacokinetics'
                    self.language = 'en'
                    self.content_blocks = [{'text': 'HSM test', 'sdt_tag': 'h'}]
                    self.events = []
                    self.created_at = '2026-01-01T00:00:00Z'
            canvas = Canvas()

            ectd_doc = compiler.compile(canvas)
            signed_path, audited_doc = compiler.sign_and_audit(ectd_doc, canvas, signer, {
                'name': 'HSM User', 'role': 'QA', 'user_id': 'hsm_01'
            })

            sig_events = [e for e in audited_doc.modification_history if isinstance(e.get('type'), dict) and e['type'].get('code') == '110107']
            assert len(sig_events) == 1
            details = sig_events[0]['entity'][0]['detail']
            assert any(d.get('type') == 'certificate-fingerprint' for d in details)
            # HSM key id should be discoverable somewhere in the audit event payload
            details_json = json.dumps(details).lower()
            assert 'hsm' in details_json or 'kms' in details_json

