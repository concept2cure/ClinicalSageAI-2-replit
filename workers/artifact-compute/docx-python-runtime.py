"""
docx-python isolated runtime contract.
Input: JSON file path argument with title/content/output_path.
Output: writes JSON payload to output_path with base64-encoded DOCX.

Uses python-docx to generate real Word documents with:
- Headers and footers
- Styled headings and body text
- Bullet and numbered lists (via markdown-style prefixes)
- Tables (via pipe-delimited rows)
- Page breaks (via --- marker)
- Regulatory-compliant formatting
"""

import base64
import io
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_docx(payload: dict) -> dict:
    title = payload.get("title", "Untitled")
    content = payload.get("content", "")
    # Images dict: key → base64 encoded image data
    # Used by ![alt](key) syntax in content
    images: dict = payload.get("images", {}) or {}
    output_type = "invalid" if payload.get("force_invalid_output") else "docx"

    if output_type == "invalid":
        blob = f"{title}\n\n{content}".encode("utf-8")
        return {
            "output_type": output_type,
            "file_name": f"{_safe_filename(title)}.docx.txt",
            "mime_type": "text/plain",
            "content_base64": base64.b64encode(blob).decode("ascii"),
            "generated_at": datetime.utcnow().isoformat() + "Z",
            "network": _network_status(),
        }

    doc = Document()

    # Configure default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Header
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(2.54)
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = title
    header_run = header_para.runs[0] if header_para.runs else header_para.add_run()
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer with page number
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Confidential")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Title
    doc.add_heading(title, level=0)

    # Parse and render content
    _render_content(doc, content, images)

    # Serialize to buffer
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    safe_name = _safe_filename(title)
    return {
        "output_type": "docx",
        "file_name": f"{safe_name}.docx",
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "content_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "network": _network_status(),
    }


def _render_content(doc: Document, content: str, images: dict | None = None) -> None:
    """Parse content string and render as structured DOCX elements.

    Supports:
    - Lines starting with # / ## / ### → headings
    - Lines starting with - or * → bullet list items
    - Lines starting with 1. / 2. etc → numbered list items
    - Lines of --- → page breaks
    - Lines starting with | → table rows (pipe-delimited)
    - Lines matching ![alt](key) → inline images from images dict
    - Lines matching ![alt](data:image/...) → inline base64 images
    - Everything else → body paragraphs
    """
    lines = content.split("\n")
    table_buffer: list[str] = []
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Flush pending table if current line is not a table row
        if table_buffer and not stripped.startswith("|"):
            _render_table(doc, table_buffer)
            table_buffer = []

        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)

        # Page break
        elif stripped == "---":
            doc.add_page_break()

        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")

        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s*", "", stripped)
            doc.add_paragraph(text, style="List Number")

        # Inline image: ![alt text](key_or_data_uri)
        elif stripped.startswith("!["):
            _render_image(doc, stripped, images)

        # Table row
        elif stripped.startswith("|"):
            table_buffer.append(stripped)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            # Handle inline bold (**text**) and italic (*text*)
            _render_inline_formatting(para, stripped)

        i += 1

    # Flush remaining table
    if table_buffer:
        _render_table(doc, table_buffer)


def _render_image(doc: Document, line: str, images: dict | None) -> None:
    """Render an image from markdown ![alt](src) syntax.

    Supports three source types:
    1. Key lookup: ![alt](chart_1) — looks up 'chart_1' in images dict (base64)
    2. Data URI:   ![alt](data:image/png;base64,...) — inline base64
    3. File path:  ![alt](/path/to/file.png) — reads from disk

    Images are embedded inline using python-docx add_picture with
    proper width constraints (max 6 inches to fit page margins).
    """
    match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
    if not match:
        doc.add_paragraph(line)
        return

    alt_text = match.group(1)
    src = match.group(2).strip()
    img_stream: io.BytesIO | None = None

    try:
        # Data URI: data:image/png;base64,iVBOR...
        if src.startswith("data:image/"):
            b64_part = src.split(",", 1)
            if len(b64_part) == 2:
                img_stream = io.BytesIO(base64.b64decode(b64_part[1]))

        # Key lookup from images dict
        elif images and src in images:
            raw = images[src]
            if isinstance(raw, str):
                img_stream = io.BytesIO(base64.b64decode(raw))
            elif isinstance(raw, bytes):
                img_stream = io.BytesIO(raw)

        # File path
        elif os.path.isfile(src):
            with open(src, "rb") as f:
                img_stream = io.BytesIO(f.read())

        if img_stream:
            doc.add_picture(img_stream, width=Inches(5.5))
            # Add caption below image
            if alt_text:
                caption = doc.add_paragraph()
                run = caption.add_run(alt_text)
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Source not found — render as text placeholder
            para = doc.add_paragraph()
            run = para.add_run(f"[Image: {alt_text or src}]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    except Exception:
        # Image rendering failed — degrade to text
        para = doc.add_paragraph()
        run = para.add_run(f"[Image could not be rendered: {alt_text or src}]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _render_inline_formatting(para, text: str) -> None:
    """Render inline **bold** and *italic* formatting."""
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def _render_table(doc: Document, rows: list[str]) -> None:
    """Render pipe-delimited rows as a Word table."""
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        # Skip separator rows (e.g., |---|---|)
        if all(re.match(r"^-+$", c) for c in cells):
            continue
        parsed.append(cells)

    if len(parsed) < 2:
        return

    col_count = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(parsed):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < col_count:
                table.rows[row_idx].cells[col_idx].text = cell_text

    # Bold the header row
    if parsed:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True


def _safe_filename(title: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_\- ]", "", title).replace(" ", "_").lower()


def _network_status() -> str:
    return "disabled" if os.getenv("ARTIFACT_COMPUTE_NO_NETWORK") == "1" else "unknown"


def main() -> int:
    if len(sys.argv) < 2:
        print("missing input path", file=sys.stderr)
        return 2

    input_path = sys.argv[1]
    with open(input_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    result = render_docx(payload)
    output_path = payload.get("output_path")
    if not output_path:
        print("missing output path", file=sys.stderr)
        return 2

    with open(output_path, "w", encoding="utf-8") as out:
        json.dump(result, out)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
