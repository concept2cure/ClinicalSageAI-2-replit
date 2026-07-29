"""
docx-insert isolated runtime contract.

Precise, surgical insertion into an existing Word document using python-docx.
This is AnA's "targeted insertions" path — the governed, document-aware
equivalent of writing a one-off script to patch a file at exact locations.

Input:  JSON file path argument with:
          {
            "source_docx_base64": "<base64 of the .docx to edit>",
            "insertions": [
              {
                "anchor_type":  "heading_text" | "placeholder" | "paragraph_index" | "start" | "end",
                "anchor_value": "<heading text | token | index>",   # not needed for start/end
                "position":     "before" | "after" | "replace",     # default "after"
                "content":      "<markdown-ish lines to insert>",
                "match":        "exact" | "contains"                 # default "contains"
              },
              ...
            ],
            "output_path": "/abs/path/output.json"
          }
Output: writes JSON payload to output_path with base64-encoded DOCX:
          {
            "output_type": "docx",
            "file_name":   "edited.docx",
            "mime_type":   "...",
            "content_base64": "...",
            "applied":     [ { "index": 0, "anchor": "...", "status": "applied"/"anchor_not_found", "paragraphs": 3 } ],
            "network":     "disabled"/"unknown",
            "generated_at":"<iso>"
          }

Supported content syntax (paragraph-level): '# / ## / ###' headings,
'- ' / '* ' bullets, '1. ' numbered, blank lines ignored, everything else a
body paragraph. Tables and images are intentionally out of scope here — use
author_docx_native for full document authoring; this tool is for precise
edits to documents that already exist.
"""

import base64
import io
import json
import os
import re
import sys
from datetime import datetime

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement


def _network_status() -> str:
    return "disabled" if os.getenv("ARTIFACT_COMPUTE_NO_NETWORK") == "1" else "unknown"


def _new_paragraph_after(ref_para: Paragraph) -> Paragraph:
    """Create an empty paragraph immediately after ref_para, in document order."""
    new_p = OxmlElement("w:p")
    ref_para._p.addnext(new_p)
    return Paragraph(new_p, ref_para._parent)


def _new_paragraph_before(ref_para: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    ref_para._p.addprevious(new_p)
    return Paragraph(new_p, ref_para._parent)


def _style_for_line(line: str):
    """Return (text, style_name_or_None, heading_level_or_None) for a content line."""
    s = line.strip()
    if s.startswith("### "):
        return s[4:], None, 3
    if s.startswith("## "):
        return s[3:], None, 2
    if s.startswith("# "):
        return s[2:], None, 1
    if s.startswith("- ") or s.startswith("* "):
        return s[2:], "List Bullet", None
    if re.match(r"^\d+\.\s", s):
        return re.sub(r"^\d+\.\s*", "", s), "List Number", None
    return s, None, None


def _apply_line_to_paragraph(doc: Document, para: Paragraph, line: str) -> None:
    """Render a single content line into an (empty) paragraph."""
    text, style, level = _style_for_line(line)
    if level is not None:
        # Heading: borrow the document's heading style by name.
        para.style = doc.styles["Heading %d" % level]
        para.add_run(text)
        return
    if style is not None:
        try:
            para.style = doc.styles[style]
        except KeyError:
            pass
    para.add_run(text)


def _content_lines(content: str) -> list:
    return [ln for ln in (content or "").split("\n") if ln.strip()]


def _insert_block(doc: Document, ref_para: Paragraph, content: str, after: bool) -> int:
    """Insert each content line as a paragraph relative to ref_para.

    Returns the number of paragraphs inserted. When after=True the block is
    placed immediately after ref_para in document order; when after=False it
    is placed immediately before. The reference advances so multi-line blocks
    keep their order.
    """
    lines = _content_lines(content)
    count = 0
    cursor = ref_para
    if after:
        for line in lines:
            para = _new_paragraph_after(cursor)
            _apply_line_to_paragraph(doc, para, line)
            cursor = para
            count += 1
    else:
        # Insert before: build top-to-bottom by always inserting before the anchor.
        for line in lines:
            para = _new_paragraph_before(ref_para)
            _apply_line_to_paragraph(doc, para, line)
            count += 1
    return count


def _find_anchor(doc: Document, anchor_type: str, anchor_value, match: str):
    """Locate the anchor paragraph. Returns (paragraph, error_or_None)."""
    paras = doc.paragraphs

    if anchor_type == "paragraph_index":
        try:
            idx = int(anchor_value)
        except (TypeError, ValueError):
            return None, "paragraph_index requires an integer anchor_value"
        if idx < 0 or idx >= len(paras):
            return None, f"paragraph_index {idx} out of range (0..{len(paras) - 1})"
        return paras[idx], None

    if anchor_type in ("heading_text", "placeholder"):
        needle = str(anchor_value or "")
        for p in paras:
            text = p.text
            if match == "exact":
                if text.strip() == needle.strip():
                    return p, None
            else:
                if needle in text:
                    return p, None
        return None, "anchor_not_found"

    return None, f"unknown anchor_type: {anchor_type}"


def _replace_placeholder(doc: Document, para: Paragraph, token: str, content: str) -> int:
    """Replace `token` inside para with content.

    Single-line content → inline token replacement (preserves surrounding
    run text). Multi-line content → the token's paragraph is cleared to the
    pre-token text and the content block is inserted after it.
    """
    lines = _content_lines(content)
    full_text = para.text
    if not lines:
        replacement = ""
    elif len(lines) == 1:
        replacement = lines[0]
    else:
        replacement = ""

    # Inline replacement across the paragraph's runs, rebuilt as a single run
    # (run-boundary-spanning tokens are why we rebuild rather than edit in place).
    new_text = full_text.replace(token, replacement)
    for run in list(para.runs):
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

    if len(lines) > 1:
        return _insert_block(doc, para, content, after=True)
    return 1


def edit_docx(payload: dict) -> dict:
    src_b64 = payload.get("source_docx_base64")
    if not src_b64:
        raise ValueError("source_docx_base64 is required")
    doc = Document(io.BytesIO(base64.b64decode(src_b64)))

    insertions = payload.get("insertions", []) or []
    applied = []

    for i, ins in enumerate(insertions):
        anchor_type = ins.get("anchor_type", "end")
        anchor_value = ins.get("anchor_value")
        position = ins.get("position", "after")
        match = ins.get("match", "contains")
        content = ins.get("content", "")

        if anchor_type == "end":
            count = 0
            for line in _content_lines(content):
                text, style, level = _style_for_line(line)
                if level is not None:
                    doc.add_heading(text, level=level)
                elif style is not None:
                    doc.add_paragraph(text, style=style)
                else:
                    doc.add_paragraph(text)
                count += 1
            applied.append({"index": i, "anchor": "end", "status": "applied", "paragraphs": count})
            continue

        if anchor_type == "start":
            if not doc.paragraphs:
                doc.add_paragraph("")
            count = _insert_block(doc, doc.paragraphs[0], content, after=False)
            applied.append({"index": i, "anchor": "start", "status": "applied", "paragraphs": count})
            continue

        ref, err = _find_anchor(doc, anchor_type, anchor_value, match)
        if err:
            applied.append({"index": i, "anchor": str(anchor_value), "status": err, "paragraphs": 0})
            continue

        if position == "replace" and anchor_type == "placeholder":
            count = _replace_placeholder(doc, ref, str(anchor_value), content)
            applied.append({"index": i, "anchor": str(anchor_value), "status": "applied", "paragraphs": count})
        elif position == "replace":
            # Replace a located paragraph: insert after, then delete the anchor.
            count = _insert_block(doc, ref, content, after=True)
            ref._p.getparent().remove(ref._p)
            applied.append({"index": i, "anchor": str(anchor_value), "status": "applied", "paragraphs": count})
        elif position == "before":
            count = _insert_block(doc, ref, content, after=False)
            applied.append({"index": i, "anchor": str(anchor_value), "status": "applied", "paragraphs": count})
        else:  # after (default)
            count = _insert_block(doc, ref, content, after=True)
            applied.append({"index": i, "anchor": str(anchor_value), "status": "applied", "paragraphs": count})

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    return {
        "output_type": "docx",
        "file_name": payload.get("file_name", "edited.docx"),
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "content_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "applied": applied,
        "network": _network_status(),
        "generated_at": datetime.utcnow().isoformat() + "Z",
    }


def main() -> int:
    if len(sys.argv) < 2:
        print("missing input path", file=sys.stderr)
        return 2

    input_path = sys.argv[1]
    with open(input_path, "r", encoding="utf-8") as f:
        payload = json.load(f)

    output_path = payload.get("output_path")
    if not output_path:
        print("missing output path", file=sys.stderr)
        return 2

    result = edit_docx(payload)
    with open(output_path, "w", encoding="utf-8") as out:
        json.dump(result, out)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
