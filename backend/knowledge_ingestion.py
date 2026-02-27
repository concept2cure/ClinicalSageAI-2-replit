
import os
import sys
import json
import logging
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional, List, Dict, Any

# Configure logging
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger("knowledge_ingestion")

# Allow imports from project root
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Define directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
CONTENT_DIR = os.path.join(DATA_DIR, "processed_knowledge")
METADATA_DIR = os.path.join(DATA_DIR, "knowledge_metadata")
PDF_DIR = os.path.join(BASE_DIR, "attached_assets")
PROJECTS_DIR = os.path.join(DATA_DIR, "projects")

# Create directories if they don't exist
os.makedirs(CONTENT_DIR, exist_ok=True)
os.makedirs(METADATA_DIR, exist_ok=True)
os.makedirs(PROJECTS_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# MULTI-FORMAT TEXT EXTRACTION
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """
    Extract text from a PDF using the multi-strategy cascade:
    pymupdf → pypdf2 → pdfminer → ocr → ocr_enhanced.
    Falls back to raw PyPDF2 if the ingestion module is unavailable.
    """
    try:
        from ingestion.pdf_extractor import extract_pdf_text
        result = extract_pdf_text(pdf_path)
        if result.get("success") and result.get("text", "").strip():
            logger.info(
                f"PDF extracted via '{result['strategy']}': {pdf_path} "
                f"({result.get('page_count', '?')} pages, {result.get('char_count', 0)} chars)"
            )
            return result["text"]
        logger.warning(f"Multi-strategy extractor returned empty text for {pdf_path}")
    except ImportError:
        logger.warning("ingestion.pdf_extractor not available — falling back to PyPDF2")
    except Exception as e:
        logger.error(f"Multi-strategy PDF extraction failed for {pdf_path}: {e}")

    # Fallback: raw PyPDF2
    try:
        import PyPDF2
        text = ""
        with open(pdf_path, 'rb') as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n\n"
        logger.info(f"PDF extracted via PyPDF2 fallback: {pdf_path}")
        return text or None
    except Exception as e:
        logger.error(f"PyPDF2 fallback failed for {pdf_path}: {e}")
        return None


def extract_text_from_docx(docx_path: str) -> Optional[str]:
    """Extract text from a .docx file preserving heading and paragraph structure."""
    try:
        from docx import Document
        doc = Document(docx_path)
        parts: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Prefix headings so downstream AI has document structure
                if para.style.name.startswith("Heading"):
                    parts.append(f"\n## {para.text.strip()}\n")
                else:
                    parts.append(para.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        text = "\n".join(parts)
        logger.info(f"DOCX extracted: {docx_path} ({len(doc.paragraphs)} paragraphs)")
        return text or None
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return None
    except Exception as e:
        logger.error(f"DOCX extraction failed for {docx_path}: {e}")
        return None


def extract_text_from_xlsx(xlsx_path: str) -> Optional[str]:
    """Extract text from an .xlsx/.xls file, one row per line, sheets separated."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
        parts: List[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"\n### Sheet: {sheet_name}\n")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        text = "\n".join(parts)
        logger.info(f"XLSX extracted: {xlsx_path}")
        return text or None
    except ImportError:
        logger.error("openpyxl not installed. Run: pip install openpyxl")
        return None
    except Exception as e:
        logger.error(f"XLSX extraction failed for {xlsx_path}: {e}")
        return None


def extract_text_from_txt(txt_path: str) -> Optional[str]:
    """Read a plain-text file, trying common encodings."""
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(txt_path, "r", encoding=enc) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"TXT read failed for {txt_path}: {e}")
            return None
    return None


def extract_text_from_file(file_path: str) -> Optional[str]:
    """
    Unified dispatcher: routes to the correct extractor based on file extension.
    Supported: .pdf, .docx, .doc, .xlsx, .xls, .txt, .md, .csv
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    if ext in (".xlsx", ".xls"):
        return extract_text_from_xlsx(file_path)
    if ext in (".txt", ".md", ".csv"):
        return extract_text_from_txt(file_path)
    logger.warning(f"Unsupported file type: {ext}  ({file_path})")
    return None

def detect_ich_guideline_type(text, filename):
    """Detect ICH guideline type and specific attributes."""
    # Defaults
    guideline_type = "unknown"
    document_category = "regulatory"
    document_class = "ich_guideline"
    region = "global"
    version = "unknown"
    status = "unknown"

    # Identify guideline based on content and filename
    if "Q1" in filename:
        guideline_type = "stability"
        if "Q1A" in filename:
            document_class = "ich_q1a_stability_testing"
        elif "Q1B" in filename:
            document_class = "ich_q1b_photostability"
    elif "Q2" in filename:
        guideline_type = "analytical_validation"
        document_class = "ich_q2_validation"
    elif "Q3" in filename:
        guideline_type = "impurities"
        if "Q3A" in filename:
            document_class = "ich_q3a_impurities"
        elif "Q3C" in filename:
            document_class = "ich_q3c_residual_solvents"
        elif "Q3D" in filename:
            document_class = "ich_q3d_elemental_impurities"
    elif "Q12" in filename:
        guideline_type = "lifecycle_management"
        document_class = "ich_q12_lifecycle"

    # Detect status from content
    if "step 4" in text.lower():
        status = "adopted"
    elif "step 2" in text.lower():
        status = "draft"

    # Detect version info
    if "(R2)" in filename or "R2" in filename:
        version = "R2"
    elif "(R1)" in filename or "R1" in filename:
        version = "R1"

    return {
        "guideline_type": guideline_type,
        "document_category": document_category,
        "document_class": document_class,
        "region": region,
        "version": version,
        "status": status
    }

def process_pdf(pdf_path):
    """Process a PDF file and extract structured data. Delegates to process_file()."""
    return process_file(pdf_path)


def process_file(file_path: str, project_id: Optional[str] = None,
                  extra_metadata: Optional[Dict[str, Any]] = None) -> Optional[str]:
    """
    Process any supported file (PDF, DOCX, XLSX, TXT) and persist its
    extracted text + metadata to the knowledge store.

    Returns the document_id (MD5 hash of content) or None on failure.
    """
    filename = os.path.basename(file_path)
    ext = Path(file_path).suffix.lower()
    logger.info(f"Processing file: {filename} (ext={ext})")

    text = extract_text_from_file(file_path)
    if not text or not text.strip():
        logger.error(f"Failed to extract text from {filename}")
        return None

    # Generate deterministic document ID
    doc_id = hashlib.md5(text.encode("utf-8")).hexdigest()

    # Detect ICH guideline metadata for PDF/DOCX
    guideline_info: Dict[str, Any] = {}
    if ext in (".pdf", ".docx", ".doc"):
        guideline_info = detect_ich_guideline_type(text, filename)

    # Detect document category from filename/text heuristics
    doc_type = _detect_document_type(filename, text)

    metadata: Dict[str, Any] = {
        "document_id": doc_id,
        "filename": filename,
        "source_path": file_path,
        "file_extension": ext,
        "extraction_date": datetime.now().isoformat(),
        "document_type": doc_type,
        "document_title": Path(filename).stem.replace("_", " ").replace("-", " ").title(),
        "char_count": len(text),
        "word_count": len(text.split()),
        **guideline_info,
    }

    if project_id:
        metadata["project_id"] = project_id

    if extra_metadata:
        metadata.update(extra_metadata)

    # Persist
    content_path = os.path.join(CONTENT_DIR, f"{doc_id}.txt")
    with open(content_path, "w", encoding="utf-8") as fh:
        fh.write(text)

    meta_path = os.path.join(METADATA_DIR, f"{doc_id}.json")
    with open(meta_path, "w", encoding="utf-8") as fh:
        json.dump(metadata, fh, indent=2)

    # If project-scoped, update project index
    if project_id:
        _add_to_project_index(project_id, doc_id, metadata)

    logger.info(f"Stored {filename} → doc_id={doc_id}")
    return doc_id

def process_directory(directory=PDF_DIR):
    """Process all supported files in a directory (PDFs, DOCX, XLSX, TXT)."""
    processed_files = []
    supported_exts = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt", ".md"}

    all_files = [
        f for f in os.listdir(directory)
        if Path(f).suffix.lower() in supported_exts
    ]

    # Prioritise ICH PDFs first, then everything else
    ich_files  = [f for f in all_files if any(x in f for x in ["ICH", "Q1", "Q2", "Q3", "Q12"])]
    other_files = [f for f in all_files if f not in ich_files]

    for filename in ich_files + other_files:
        doc_id = process_file(os.path.join(directory, filename))
        if doc_id:
            processed_files.append({"filename": filename, "document_id": doc_id})

    summary = {
        "timestamp": datetime.now().isoformat(),
        "processed_files": len(processed_files),
        "files": processed_files,
    }
    summary_path = os.path.join(METADATA_DIR, "ingestion_summary.json")
    with open(summary_path, "w", encoding="utf-8") as fh:
        json.dump(summary, fh, indent=2)

    logger.info(f"Processed {len(processed_files)} files.")
    return summary


# ─────────────────────────────────────────────────────────────────────────────
# PROJECT-SCOPED INGESTION & CONTEXT SYNTHESIS
# ─────────────────────────────────────────────────────────────────────────────

def _project_index_path(project_id: str) -> str:
    return os.path.join(PROJECTS_DIR, f"{project_id}_index.json")


def _add_to_project_index(project_id: str, doc_id: str, metadata: Dict[str, Any]) -> None:
    """Add a document entry to the project's index file."""
    path = _project_index_path(project_id)
    index: Dict[str, Any] = {"project_id": project_id, "documents": {}}
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as fh:
                index = json.load(fh)
        except Exception:
            pass
    index["documents"][doc_id] = {
        "filename": metadata.get("filename"),
        "document_title": metadata.get("document_title"),
        "document_type": metadata.get("document_type"),
        "char_count": metadata.get("char_count", 0),
        "word_count": metadata.get("word_count", 0),
        "ingested_at": metadata.get("extraction_date"),
    }
    index["last_updated"] = datetime.now().isoformat()
    with open(path, "w", encoding="utf-8") as fh:
        json.dump(index, fh, indent=2)


def ingest_project_files(
    file_paths: List[str],
    project_id: str,
    extra_metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """
    Ingest a list of file paths into the knowledge store under a project namespace.
    Supports PDFs, DOCX, XLSX, and TXT.

    Returns a summary dict: { project_id, ingested, failed, documents: [...] }
    """
    ingested: List[Dict[str, Any]] = []
    failed: List[Dict[str, Any]] = []

    for fp in file_paths:
        if not os.path.exists(fp):
            failed.append({"path": fp, "error": "File not found"})
            continue
        doc_id = process_file(fp, project_id=project_id, extra_metadata=extra_metadata)
        if doc_id:
            ingested.append({
                "filename": os.path.basename(fp),
                "document_id": doc_id,
                "path": fp,
            })
        else:
            failed.append({"path": fp, "error": "Extraction failed"})

    return {
        "project_id": project_id,
        "ingested": len(ingested),
        "failed": len(failed),
        "documents": ingested,
        "errors": failed,
        "timestamp": datetime.now().isoformat(),
    }


def ingest_project_bytes(
    files: List[Dict[str, Any]],
    project_id: str,
) -> Dict[str, Any]:
    """
    Ingest in-memory file bytes into the project knowledge store.

    `files` is a list of dicts with keys:
        filename  (str)  — original filename including extension
        content   (bytes) — raw file bytes
        mime_type (str, optional)

    Returns the same summary dict as ingest_project_files().
    """
    import tempfile
    ingested: List[Dict[str, Any]] = []
    failed:   List[Dict[str, Any]] = []

    with tempfile.TemporaryDirectory(prefix="c2c_ingest_") as tmp:
        for f in files:
            filename = f.get("filename", "upload.bin")
            content  = f.get("content", b"")
            if not content:
                failed.append({"filename": filename, "error": "Empty content"})
                continue

            tmp_path = os.path.join(tmp, filename)
            with open(tmp_path, "wb") as fh:
                fh.write(content)

            doc_id = process_file(tmp_path, project_id=project_id)
            if doc_id:
                ingested.append({"filename": filename, "document_id": doc_id})
            else:
                failed.append({"filename": filename, "error": "Extraction failed"})

    return {
        "project_id": project_id,
        "ingested": len(ingested),
        "failed": len(failed),
        "documents": ingested,
        "errors": failed,
        "timestamp": datetime.now().isoformat(),
    }


def get_project_context(
    project_id: str,
    max_chars_per_doc: int = 4000,
    max_total_chars: int = 80_000,
) -> Dict[str, Any]:
    """
    Load and synthesize all documents for a project into a single context object.

    The returned dict contains:
        project_id       — project identifier
        document_count   — number of documents found
        documents        — list of { id, title, type, excerpt }
        combined_context — concatenated text capped at max_total_chars
        summary_prompt   — a ready-to-use AI prompt that includes all context
        keywords         — top extracted keywords across all documents
    """
    index_path = _project_index_path(project_id)
    if not os.path.exists(index_path):
        return {
            "project_id": project_id,
            "document_count": 0,
            "documents": [],
            "combined_context": "",
            "summary_prompt": "",
            "keywords": [],
            "error": "No documents ingested for this project yet.",
        }

    with open(index_path, "r", encoding="utf-8") as fh:
        index = json.load(fh)

    doc_entries = index.get("documents", {})
    documents_detail: List[Dict[str, Any]] = []
    combined_parts: List[str] = []
    total_chars = 0

    for doc_id, doc_meta in doc_entries.items():
        content_path = os.path.join(CONTENT_DIR, f"{doc_id}.txt")
        if not os.path.exists(content_path):
            continue
        with open(content_path, "r", encoding="utf-8") as fh:
            full_text = fh.read()

        # Truncate per-document
        excerpt = full_text[:max_chars_per_doc]
        if total_chars + len(excerpt) > max_total_chars:
            remaining = max_total_chars - total_chars
            if remaining <= 0:
                break
            excerpt = excerpt[:remaining]

        documents_detail.append({
            "id": doc_id,
            "title": doc_meta.get("document_title", doc_meta.get("filename", "Unknown")),
            "type": doc_meta.get("document_type", "unknown"),
            "filename": doc_meta.get("filename", ""),
            "word_count": doc_meta.get("word_count", 0),
            "excerpt": excerpt[:500] + ("..." if len(full_text) > 500 else ""),
        })

        combined_parts.append(
            f"--- Document: {doc_meta.get('document_title', doc_id)} ---\n{excerpt}"
        )
        total_chars += len(excerpt)

    combined_context = "\n\n".join(combined_parts)
    keywords = _extract_keywords(combined_context)

    summary_prompt = (
        f"You are an expert regulatory affairs specialist.\n"
        f"The following documents have been uploaded for project '{project_id}':\n\n"
        + "\n".join(f"- {d['title']} ({d['type']})" for d in documents_detail)
        + f"\n\nCombined source context ({total_chars:,} chars):\n\n{combined_context}"
    )

    return {
        "project_id": project_id,
        "document_count": len(documents_detail),
        "documents": documents_detail,
        "combined_context": combined_context,
        "summary_prompt": summary_prompt,
        "keywords": keywords,
        "last_updated": index.get("last_updated"),
    }


def _extract_keywords(text: str, top_n: int = 30) -> List[str]:
    """
    Simple TF-IDF-style keyword extraction without external dependencies.
    Returns the top_n most distinctive words (min 5 chars, not stop-words).
    """
    STOP = {
        "the", "and", "for", "with", "that", "this", "from", "are", "have",
        "been", "will", "should", "which", "their", "they", "shall", "each",
        "other", "also", "into", "more", "not", "its", "all", "any", "such",
        "may", "must", "used", "data", "study", "studies", "clinical",
    }
    words = [
        w.lower().strip(".,;:()[]\"'!?")
        for w in text.split()
        if len(w) >= 5
    ]
    freq: Dict[str, int] = {}
    for w in words:
        if w not in STOP and w.isalpha():
            freq[w] = freq.get(w, 0) + 1
    sorted_words = sorted(freq.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in sorted_words[:top_n]]


def _detect_document_type(filename: str, text: str) -> str:
    """Heuristic document type detection from filename and content."""
    fn_upper = filename.upper()
    text_lower = text[:2000].lower()

    if any(x in fn_upper for x in ["ICH", "Q1", "Q2", "Q3", "Q4", "Q5", "Q6",
                                     "Q7", "Q8", "Q9", "Q10", "Q11", "Q12",
                                     "S1", "S2", "S3", "S4", "S5", "S6", "S7",
                                     "E1", "E2", "E3", "E4", "E5", "E6", "E7",
                                     "E8", "E8", "E9", "M1", "M2", "M3", "M4",
                                     "M5", "M6", "M7"]):
        return "ich_guideline"
    if any(x in fn_upper for x in ["IND", "INVESTIGATIONAL", "NDA", "BLA", "ANDA"]):
        return "regulatory_submission"
    if any(x in fn_upper for x in ["CSR", "CLINICAL_STUDY", "CLINICAL STUDY"]):
        return "clinical_study_report"
    if any(x in fn_upper for x in ["CMC", "QUALITY", "DRUG_SUBSTANCE", "DRUG_PRODUCT"]):
        return "cmc_document"
    if any(x in fn_upper for x in ["PROTOCOL", "STUDY_PROTOCOL"]):
        return "protocol"
    if any(x in fn_upper for x in ["INVESTIGATOR_BROCHURE", "IB_"]):
        return "investigator_brochure"
    if "clinical trial" in text_lower or "adverse event" in text_lower:
        return "clinical_document"
    if "chemistry" in text_lower or "manufacturing" in text_lower or "specification" in text_lower:
        return "cmc_document"
    return "general_document"


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="TrialSage Knowledge Ingestion")
    subparsers = parser.add_subparsers(dest="command")

    # Ingest a directory
    dir_p = subparsers.add_parser("directory", help="Ingest all files in a directory")
    dir_p.add_argument("--dir", default=PDF_DIR, help="Directory path")

    # Ingest specific files into a project
    proj_p = subparsers.add_parser("project", help="Ingest files into a project")
    proj_p.add_argument("--project-id", required=True, help="Project ID")
    proj_p.add_argument("files", nargs="+", help="File paths to ingest")

    # Dump project context
    ctx_p = subparsers.add_parser("context", help="Print project context summary")
    ctx_p.add_argument("--project-id", required=True, help="Project ID")

    args = parser.parse_args()

    if args.command == "directory":
        logger.info("Starting directory ingestion")
        summary = process_directory(args.dir)
        logger.info(f"Ingestion complete: {summary['processed_files']} files.")
    elif args.command == "project":
        summary = ingest_project_files(args.files, args.project_id)
        print(json.dumps(summary, indent=2))
    elif args.command == "context":
        ctx = get_project_context(args.project_id)
        # Print everything except the full combined_context to keep output readable
        ctx_display = {k: v for k, v in ctx.items() if k != "combined_context"}
        print(json.dumps(ctx_display, indent=2))
    else:
        parser.print_help()
