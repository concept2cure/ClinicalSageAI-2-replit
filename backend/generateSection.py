"""
Enhanced Section Generation Agent for eCTD CoAuthor Module

This agent provides LLM-powered, source-based generation for regulatory document sections
within the existing eCTD CoAuthor workflow. It integrates directly with the document
workspace and maintains full traceability of generated content.

New capabilities (2026-02):
- load_project_documents()     — auto-load all ingested project files as context
- generate_docx_document()     — produce a formatted .docx from generated sections
- generate_ind_package_docx()  — generate a complete CTD-structured IND .docx
- generate_document_from_project() — end-to-end: context → AI → .docx in one call
"""

import os
import io
import sys
import json
import asyncio
from typing import Dict, List, Optional, Any
from datetime import datetime
import openai
from dataclasses import dataclass, asdict
import hashlib
import uuid
from pathlib import Path

# Allow imports from project root (for knowledge_ingestion)
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# Set up OpenAI client
openai.api_key = os.getenv('OPENAI_API_KEY')

@dataclass
class SourceReference:
    """Represents a source document reference used in generation"""
    document_id: str
    document_title: str
    section: str
    page_number: Optional[int]
    excerpt: str
    confidence_score: float
    timestamp: str

@dataclass
class GeneratedSection:
    """Represents a generated document section with full traceability"""
    section_id: str
    section_title: str
    content: str
    source_references: List[SourceReference]
    generation_metadata: Dict[str, Any]
    quality_score: float
    regulatory_compliance_score: float
    timestamp: str

class EnhancedSectionGenerator:
    """
    Enhanced section generator that integrates with the eCTD CoAuthor module
    to provide intelligent, source-based content generation.
    """

    def __init__(self):
        self.model = "gpt-4o"  # Latest OpenAI model
        self.generation_history = []
        self.source_database = {}

    async def generate_regulatory_section(
        self,
        section_type: str,
        context_documents: List[Dict],
        user_requirements: Dict,
        compliance_region: str = "FDA"
    ) -> GeneratedSection:
        """
        Generate a regulatory document section with full source traceability

        Args:
            section_type: Type of section (e.g., "executive_summary", "safety_overview")
            context_documents: List of source documents with content
            user_requirements: User-specified requirements and constraints
            compliance_region: Regulatory region (FDA, EMA, PMDA, etc.)

        Returns:
            GeneratedSection with content and full traceability
        """

        section_id = str(uuid.uuid4())
        timestamp = datetime.utcnow().isoformat()

        # Prepare source context from documents
        source_context = self._prepare_source_context(context_documents)

        # Generate regulatory-compliant content
        generated_content = await self._generate_content_with_sources(
            section_type, source_context, user_requirements, compliance_region
        )

        # Extract and validate source references
        source_references = self._extract_source_references(
            generated_content, context_documents
        )

        # Calculate quality and compliance scores
        quality_score = await self._calculate_quality_score(generated_content)
        compliance_score = await self._calculate_compliance_score(
            generated_content, section_type, compliance_region
        )

        # Create generation metadata
        metadata = {
            "model_used": self.model,
            "generation_timestamp": timestamp,
            "user_id": user_requirements.get("user_id", "unknown"),
            "section_type": section_type,
            "compliance_region": compliance_region,
            "source_document_count": len(context_documents),
            "content_hash": hashlib.sha256(generated_content.encode()).hexdigest(),
            "generation_parameters": user_requirements
        }

        generated_section = GeneratedSection(
            section_id=section_id,
            section_title=user_requirements.get("title", f"Generated {section_type}"),
            content=generated_content,
            source_references=source_references,
            generation_metadata=metadata,
            quality_score=quality_score,
            regulatory_compliance_score=compliance_score,
            timestamp=timestamp
        )

        # Store in generation history for traceability
        self.generation_history.append(generated_section)

        return generated_section

    def _prepare_source_context(self, documents: List[Dict]) -> str:
        """Prepare source document context for generation"""
        context_parts = []

        for doc in documents:
            doc_context = f"Document: {doc.get('title', 'Unknown')}\n"
            doc_context += f"Type: {doc.get('type', 'Unknown')}\n"
            doc_context += f"Content: {doc.get('content', '')[:2000]}...\n"
            doc_context += "---\n"
            context_parts.append(doc_context)

        return "\n".join(context_parts)

    async def _generate_content_with_sources(
        self,
        section_type: str,
        source_context: str,
        requirements: Dict,
        region: str
    ) -> str:
        """Generate content using OpenAI with source-based prompting"""

        # Build regulatory-specific prompt
        prompt = self._build_regulatory_prompt(section_type, source_context, requirements, region)

        try:
            client = openai.OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            response = client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": f"You are an expert regulatory writer specializing in {region} submissions. "
                                   "Generate content that is accurate, compliant, and fully traceable to source documents. "
                                   "Always cite specific sources and maintain regulatory standards. "
                                   "Return your response as JSON with 'content' field."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.3,  # Lower temperature for more consistent regulatory content
                max_tokens=4000,
                response_format={"type": "json_object"}
            )

            content_data = json.loads(response.choices[0].message.content)
            return content_data.get("content", "")

        except Exception as e:
            print(f"Error generating content: {e}")
            return f"Error generating {section_type}: {str(e)}"

    def _build_regulatory_prompt(
        self,
        section_type: str,
        source_context: str,
        requirements: Dict,
        region: str
    ) -> str:
        """Build a regulatory-specific prompt for content generation"""

        prompt = f"""
        Generate a {section_type} section for an {region} regulatory submission based on the provided source documents.

        Requirements:
        - Section Type: {section_type}
        - Regulatory Region: {region}
        - Length: {requirements.get('length', 'standard')}
        - Focus Areas: {', '.join(requirements.get('focus_areas', []))}

        Source Documents:
        {source_context}

        Instructions:
        1. Generate content that is accurate and compliant with {region} regulatory standards
        2. Base all statements on the provided source documents
        3. Include specific citations to source documents
        4. Maintain professional regulatory writing style
        5. Ensure logical flow and clear structure
        6. Include any required regulatory disclosures

        Return your response as JSON with the following structure:
        {{
            "content": "Generated section content with proper citations",
            "key_points": ["List of key points covered"],
            "compliance_notes": ["Any regulatory compliance considerations"],
            "source_citations": ["List of sources referenced"]
        }}
        """

        return prompt.strip()

    def _extract_source_references(
        self,
        content: str,
        documents: List[Dict]
    ) -> List[SourceReference]:
        """Extract and validate source references from generated content"""

        references = []
        timestamp = datetime.utcnow().isoformat()

        for doc in documents:
            # Simple extraction based on document mentions in content
            if doc.get('title', '').lower() in content.lower():
                reference = SourceReference(
                    document_id=doc.get('id', str(uuid.uuid4())),
                    document_title=doc.get('title', 'Unknown'),
                    section=doc.get('section', 'Unknown'),
                    page_number=doc.get('page', None),
                    excerpt=doc.get('content', '')[:200] + "...",
                    confidence_score=0.85,  # Placeholder - would use actual similarity scoring
                    timestamp=timestamp
                )
                references.append(reference)

        return references

    async def _calculate_quality_score(self, content: str) -> float:
        """Calculate content quality score based on various metrics"""

        # Simple quality metrics (would be enhanced with proper NLP analysis)
        score = 0.0

        # Length appropriateness
        if 500 <= len(content) <= 5000:
            score += 0.3

        # Professional language indicators
        professional_terms = ["clinical", "regulatory", "compliance", "safety", "efficacy"]
        term_count = sum(1 for term in professional_terms if term in content.lower())
        score += min(term_count * 0.1, 0.4)

        # Structure indicators
        if content.count('\n') > 3:  # Has paragraphs
            score += 0.2

        # Citation indicators
        if "[" in content and "]" in content:  # Has citations
            score += 0.1

        return min(score, 1.0)

    async def _calculate_compliance_score(
        self,
        content: str,
        section_type: str,
        region: str
    ) -> float:
        """Calculate regulatory compliance score"""

        # Placeholder compliance scoring (would integrate with regulatory knowledge base)
        base_score = 0.7

        # Region-specific compliance checks
        if region == "FDA":
            if "FDA" in content or "CFR" in content:
                base_score += 0.1
        elif region == "EMA":
            if "EMA" in content or "EU" in content:
                base_score += 0.1

        # Section-specific compliance
        if section_type == "safety_overview" and "safety" in content.lower():
            base_score += 0.1

        return min(base_score, 1.0)

    def get_generation_history(self) -> List[Dict]:
        """Get history of all generated sections"""
        return [asdict(section) for section in self.generation_history]

    def get_section_by_id(self, section_id: str) -> Optional[Any]:
        """Retrieve a specific generated section by ID"""
        for section in self.generation_history:
            if section.section_id == section_id:
                return section
        return None

    # ── Project-context integration ──────────────────────────────────────────

    def load_project_documents(self, project_id: str,
                                max_chars_per_doc: int = 4000,
                                max_total_chars: int = 80_000) -> List[Dict]:
        """
        Load all ingested documents for *project_id* from the knowledge store
        and return them as the context_documents list expected by
        generate_regulatory_section().

        Falls back to [] if knowledge_ingestion is unavailable.
        """
        try:
            from knowledge_ingestion import get_project_context
            ctx = get_project_context(
                project_id,
                max_chars_per_doc=max_chars_per_doc,
                max_total_chars=max_total_chars,
            )
            documents = []
            for doc in ctx.get("documents", []):
                documents.append({
                    "id":      doc["id"],
                    "title":   doc["title"],
                    "type":    doc["type"],
                    "content": doc["excerpt"],
                })
            return documents
        except ImportError:
            print("[generateSection] knowledge_ingestion not available")
            return []

    # ── DOCX generation ──────────────────────────────────────────────────────

    def generate_docx_document(
        self,
        sections: List[Dict[str, Any]],
        document_title: str = "Regulatory Document",
        organization: str = "",
        author: str = "",
        output_path: Optional[str] = None,
    ) -> bytes:
        """
        Produce a formatted .docx from a list of generated section dicts.

        Each section dict must have at least:
            section_title (str)
            content       (str)

        Optional fields used when present:
            citations     (list of str)
            validation    (dict with score)

        Returns raw docx bytes. If output_path is given, also saves to disk.
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # ── Page setup
        section_prop = doc.sections[0]
        section_prop.page_width  = Inches(8.5)
        section_prop.page_height = Inches(11)
        section_prop.left_margin   = Inches(1.25)
        section_prop.right_margin  = Inches(1.25)
        section_prop.top_margin    = Inches(1)
        section_prop.bottom_margin = Inches(1)

        # ── Styles
        styles = doc.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)

        # ── Cover page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(document_title)
        run.bold = True
        run.font.size = Pt(18)

        if organization:
            org_para = doc.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            org_para.add_run(organization).font.size = Pt(12)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(
            f"Generated: {datetime.now().strftime('%B %d, %Y')}"
        ).font.size = Pt(11)

        if author:
            auth_para = doc.add_paragraph()
            auth_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            auth_para.add_run(f"Prepared by: {author}").font.size = Pt(11)

        doc.add_page_break()

        # ── Table of Contents placeholder
        toc_heading = doc.add_heading("Table of Contents", level=1)
        for s in sections:
            toc_item = doc.add_paragraph(style="List Number")
            toc_item.add_run(s.get("section_title", "Untitled"))
        doc.add_page_break()

        # ── Document sections
        for s in sections:
            title_text = s.get("section_title", "Untitled Section")
            content    = s.get("content", "")
            citations  = s.get("citations", [])
            validation = s.get("validation", {})

            # Section heading
            doc.add_heading(title_text, level=2)

            # Compliance / quality badge
            score = validation.get("score") if validation else None
            if score is not None:
                badge = doc.add_paragraph()
                run   = badge.add_run(f"Quality Score: {score:.0%}  |  FDA Compliance Check")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            # Content — split on newlines to preserve paragraphs
            for para_text in content.split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                if para_text.startswith("##"):
                    doc.add_heading(para_text.lstrip("#").strip(), level=3)
                else:
                    doc.add_paragraph(para_text)

            # Citations block
            if citations:
                doc.add_heading("References", level=3)
                for cite in citations:
                    ref_para = doc.add_paragraph(style="List Bullet")
                    if isinstance(cite, dict):
                        ref_para.add_run(
                            f"{cite.get('title', '')} — {cite.get('source', '')}"
                        )
                    else:
                        ref_para.add_run(str(cite))

            # Separator
            doc.add_paragraph("─" * 60).runs[0].font.size = Pt(8)

        # ── Footer: confidentiality notice
        for sec in doc.sections:
            footer_para = sec.footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run(
                "CONFIDENTIAL — For Regulatory Submission Purposes Only"
            ).font.size = Pt(8)

        # ── Serialise
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        if output_path:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, "wb") as fh:
                fh.write(docx_bytes)

        return docx_bytes

    async def generate_ind_package_docx(
        self,
        project_id: str,
        drug_name: str = "Investigational Drug",
        sponsor: str = "",
        indication: str = "",
        ctd_sections: Optional[List[str]] = None,
        compliance_region: str = "FDA",
        output_path: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        End-to-end IND package generation:
        1. Load all project documents from the knowledge store
        2. Generate AI content for each requested CTD section
        3. Assemble a formatted .docx with all sections
        4. Return { docx_bytes, sections, metadata }

        ctd_sections defaults to the IND-minimum set (Modules 1–5 summaries).
        """
        if ctd_sections is None:
            ctd_sections = [
                "1.0 - Cover Letter & Form FDA 1571",
                "2.1 - Comprehensive Table of Contents",
                "2.2 - Introduction to the Summary",
                "2.3 - Quality Overall Summary (QOS)",
                "2.4 - Nonclinical Overview",
                "2.5 - Clinical Overview",
                "3.2.S - Drug Substance (CMC)",
                "3.2.P - Drug Product (CMC)",
                "4.1 - Nonclinical Study Reports Table of Contents",
                "5.1 - Clinical Study Reports Table of Contents",
            ]

        # Load project context
        project_docs = self.load_project_documents(project_id)

        user_base_requirements = {
            "drug_name":  drug_name,
            "sponsor":    sponsor,
            "indication": indication,
            "user_id":    f"project:{project_id}",
        }

        generated_sections: List[Dict[str, Any]] = []
        errors: List[str] = []

        for ctd_section in ctd_sections:
            try:
                requirements = {
                    **user_base_requirements,
                    "title":       ctd_section,
                    "focus_areas": [ctd_section],
                    "length":      "standard",
                }
                result = await self.generate_regulatory_section(
                    section_type=ctd_section,
                    context_documents=project_docs,
                    user_requirements=requirements,
                    compliance_region=compliance_region,
                )
                generated_sections.append({
                    "section_title": ctd_section,
                    "content":       result.content,
                    "citations":     [asdict(r) for r in result.source_references],
                    "validation":    {
                        "score":              result.regulatory_compliance_score,
                        "quality_score":      result.quality_score,
                        "section_id":         result.section_id,
                    },
                })
            except Exception as e:
                errors.append(f"{ctd_section}: {str(e)}")

        # Build DOCX
        doc_title = f"IND Application — {drug_name}"
        docx_bytes = self.generate_docx_document(
            sections=generated_sections,
            document_title=doc_title,
            organization=sponsor,
            output_path=output_path,
        )

        return {
            "success":            len(generated_sections) > 0,
            "project_id":         project_id,
            "drug_name":          drug_name,
            "document_title":     doc_title,
            "sections_generated": len(generated_sections),
            "sections_failed":    len(errors),
            "errors":             errors,
            "docx_bytes":         docx_bytes,
            "docx_size_bytes":    len(docx_bytes),
            "generated_at":       datetime.now().isoformat(),
        }

    async def generate_document_from_project(
        self,
        project_id: str,
        document_type: str,
        document_config: Optional[Dict[str, Any]] = None,
        output_path: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        High-level entry point: synthesise project knowledge → AI draft → .docx.

        document_type examples:
            "ind_package"         — full IND CTD sections
            "clinical_overview"   — Module 2.5 clinical overview only
            "quality_summary"     — Module 2.3 QOS only
            "safety_summary"      — Module 2.4 nonclinical overview
            "investigator_brochure" — complete IB

        document_config may include: drug_name, sponsor, indication, sections,
            compliance_region, max_chars_per_doc, max_total_chars
        """
        cfg = document_config or {}
        drug_name  = cfg.get("drug_name",  "Investigational Drug")
        sponsor    = cfg.get("sponsor",    "")
        indication = cfg.get("indication", "")
        region     = cfg.get("compliance_region", "FDA")

        DOCUMENT_SECTION_MAP: Dict[str, List[str]] = {
            "ind_package": [
                "1.0 - Cover Letter & Form FDA 1571",
                "2.3 - Quality Overall Summary",
                "2.4 - Nonclinical Overview",
                "2.5 - Clinical Overview",
                "3.2.S - Drug Substance",
                "3.2.P - Drug Product",
                "4.1 - Nonclinical Study Table of Contents",
                "5.1 - Clinical Study Table of Contents",
            ],
            "clinical_overview": ["2.5 - Clinical Overview"],
            "quality_summary":   ["2.3 - Quality Overall Summary (QOS)"],
            "safety_summary":    ["2.4 - Nonclinical Overview"],
            "investigator_brochure": [
                "IB Section 1 - Introduction & Summary",
                "IB Section 2 - Physical, Chemical & Pharmaceutical Properties",
                "IB Section 3 - Nonclinical Studies",
                "IB Section 4 - Effects in Humans",
                "IB Section 5 - Summary",
            ],
        }

        sections = cfg.get("sections") or DOCUMENT_SECTION_MAP.get(document_type, [
            f"{document_type} - Main Section"
        ])

        return await self.generate_ind_package_docx(
            project_id=project_id,
            drug_name=drug_name,
            sponsor=sponsor,
            indication=indication,
            ctd_sections=sections,
            compliance_region=region,
            output_path=output_path,
        )

# Global instance for use in the eCTD CoAuthor module
section_generator = EnhancedSectionGenerator()

async def generate_ectd_section(
    section_type: str,
    documents: List[Dict],
    requirements: Dict,
    region: str = "FDA"
) -> Dict:
    """
    Main function to generate eCTD sections from the CoAuthor module

    Returns:
        Dictionary containing generated section data
    """

    generated_section = await section_generator.generate_regulatory_section(
        section_type, documents, requirements, region
    )

    return asdict(generated_section)

if __name__ == "__main__":
    import sys
    import json
    import argparse
    import asyncio

    parser = argparse.ArgumentParser(description='Enhanced Section Generation Agent')
    parser.add_argument('--input', type=str, required=True, help='JSON input data')
    args = parser.parse_args()

    try:
        input_data = json.loads(args.input)

        async def main():
            result = await generate_ectd_section(
                section_type=input_data.get('section_type', ''),
                documents=input_data.get('documents', []),
                requirements=input_data.get('requirements', {}),
                region=input_data.get('compliance_region', 'FDA')
            )
            print(json.dumps(result))

        asyncio.run(main())

    except Exception as e:
        error_result = {
            "success": False,
            "error": f"Command line execution error: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }
        print(json.dumps(error_result))
